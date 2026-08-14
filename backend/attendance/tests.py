from datetime import date
from io import BytesIO
from io import StringIO

from django.contrib.auth import get_user_model
from django.conf import settings
from django.contrib.sessions.models import Session as DjangoSession
from django.test import TestCase
from django.core.management import call_command
from django.urls import reverse
from openpyxl import Workbook, load_workbook
from docx import Document
from rest_framework.test import APIClient

from .models import (
    AcademicYear,
    AccountantProfile,
    AttendanceRecord,
    ClassLevel,
    Module,
    PaymentCategory,
    Semester,
    Session,
    Student,
    StudentFinanceClearance,
    StudentFinanceObligation,
    StudentPayment,
    StudentResult,
    TeacherProfile,
    EstateOfficerProfile, InventoryLocation, AssetCategory, InventoryItemType, Asset, AssetTransfer,
)
from .serializers import StudentResultSerializer
from .grading import gpa_classification


User = get_user_model()


class InventoryManagementTests(TestCase):
    def setUp(self):
        self.admin = User.objects.create_superuser('inventory-admin', 'admin@example.com', 'password123')
        self.officer = User.objects.create_user('estate', password='password123')
        EstateOfficerProfile.objects.create(user=self.officer, full_name='Estate Officer')
        self.location = InventoryLocation.objects.create(name='Test Inventory Office')
        self.category = AssetCategory.objects.create(name='Test Equipment')
        self.item_type = InventoryItemType.objects.create(
            name='Office Chair', category=self.category, description='Standard chair',
            default_tag_prefix='BPCH/CH',
        )
        self.client = APIClient()

    def test_default_inventory_catalog_is_available(self):
        expected = {'Office Chair', 'Student Chair', 'CPU', 'Monitor', 'Fire Extinguisher', 'Dustbin', 'Extension Cable'}
        self.assertTrue(expected.issubset(set(InventoryItemType.objects.values_list('name', flat=True))))

    def test_admin_can_create_estate_account_but_cannot_access_inventory(self):
        self.client.force_authenticate(self.admin)
        created = self.client.post('/api/staff-accounts/', {
            'role': 'estate_officer', 'full_name': 'Property Officer',
            'username': 'property', 'password': 'secret12',
        }, format='json')
        self.assertEqual(created.status_code, 201)
        self.assertTrue(EstateOfficerProfile.objects.filter(user__username='property').exists())
        self.assertEqual(self.client.get('/api/assets/').status_code, 403)
        self.assertEqual(self.client.get('/api/inventory/template/').status_code, 403)

    def test_all_manual_asset_fields_are_required_and_tag_is_case_insensitive_unique(self):
        self.client.force_authenticate(self.officer)
        missing = self.client.post('/api/assets/', {'asset_tag': 'TAG-1'}, format='json')
        self.assertEqual(missing.status_code, 400)
        payload = {
            'asset_tag': 'TAG-1', 'name': 'Desktop', 'description': 'Office desktop',
            'category': self.category.id, 'location': self.location.id,
            'responsible_office': 'ICT Office', 'quantity': 1, 'condition': 'good',
        }
        self.assertEqual(self.client.post('/api/assets/', payload, format='json').status_code, 201)
        payload['asset_tag'] = 'tag-1'
        self.assertEqual(self.client.post('/api/assets/', payload, format='json').status_code, 400)

    def test_description_is_optional_and_transfer_updates_current_assignment(self):
        self.client.force_authenticate(self.officer)
        payload = {'asset_tag': 'NO-DESC', 'name': 'Projector', 'category': self.category.id,
                   'location': self.location.id, 'responsible_office': 'ICT Office',
                   'quantity': 1, 'condition': 'good'}
        created = self.client.post('/api/assets/', payload, format='json')
        self.assertEqual(created.status_code, 201)
        destination = InventoryLocation.objects.create(name='Destination Office')
        moved = self.client.post('/api/asset-transfers/', {
            'asset': created.data['id'], 'to_location': destination.id,
            'new_responsible_office': 'Principal Office', 'reason': 'Official allocation',
            'transferred_at': '2026-08-04',
        }, format='json')
        self.assertEqual(moved.status_code, 201)
        asset = Asset.objects.get(pk=created.data['id'])
        self.assertEqual(asset.location, destination)
        self.assertEqual(asset.responsible_office, 'Principal Office')
        self.assertEqual(AssetTransfer.objects.get(asset=asset).from_location, self.location)

    def test_partial_batch_transfer_creates_traceable_split_tag(self):
        self.client.force_authenticate(self.officer)
        batch = Asset.objects.create(
            asset_tag='BPHACOH/L4/CHAIR/001', name='Class chairs', description='',
            category=self.category, location=self.location, responsible_office='Level 4',
            quantity=30, condition='good', created_by=self.officer, updated_by=self.officer,
        )
        destination = InventoryLocation.objects.create(name='Level 5 Test Classroom')
        response = self.client.post('/api/asset-transfers/', {
            'asset': batch.id, 'quantity': 1, 'to_location': destination.id,
            'new_responsible_office': 'Level 5', 'reason': 'Replacement chair',
            'transferred_at': '2026-08-04',
        }, format='json')
        self.assertEqual(response.status_code, 201)
        batch.refresh_from_db()
        self.assertEqual(batch.quantity, 29)
        split = Asset.objects.get(asset_tag='BPHACOH/L4/CHAIR/001/001')
        self.assertEqual(split.quantity, 1)
        self.assertEqual(split.location, destination)
        self.assertEqual(response.data['resulting_asset_tag'], split.asset_tag)

    def test_template_contains_required_headers_and_valid_file_imports_atomically(self):
        self.client.force_authenticate(self.officer)
        template = self.client.get('/api/inventory/template/')
        self.assertEqual(template.status_code, 200)
        wb = load_workbook(BytesIO(template.content))
        self.assertEqual(wb['Assets']['A1'].value, 'Office/Location *')
        ws = wb['Assets']
        ws.append([self.location.name, 'Level 4', f'{self.item_type.name} — {self.category.name}',
                   'TEST/CH', 1, 20, 'Good', 'Student chairs'])
        output = BytesIO(); wb.save(output); output.seek(0)
        output.name = 'inventory.xlsx'
        checked = self.client.post('/api/inventory/import/', {'file': output, 'confirm': 'false'}, format='multipart')
        self.assertEqual(checked.status_code, 200)
        self.assertFalse(Asset.objects.filter(asset_tag='TAG-2').exists())
        output.seek(0)
        imported = self.client.post('/api/inventory/import/', {'file': output, 'confirm': 'true'}, format='multipart')
        self.assertEqual(imported.status_code, 201)
        self.assertEqual(Asset.objects.filter(asset_tag__startswith='TEST/CH/').count(), 20)
        self.assertEqual(Asset.objects.get(asset_tag='TEST/CH/1').quantity, 1)

    def test_invalid_excel_row_imports_nothing(self):
        self.client.force_authenticate(self.officer)
        wb = Workbook(); ws = wb.active; ws.title = 'Assets'
        ws.append(['Office/Location *', 'Responsible Person/Office *', 'Item Type *',
                   'Tag Prefix *', 'Starting Number *', 'Quantity *', 'Condition *', 'Description'])
        item_label = f'{self.item_type.name} — {self.category.name}'
        ws.append([self.location.name, 'Office', item_label, 'TAG-3', 1, 1, 'Good', 'Description'])
        ws.append([self.location.name, '', item_label, 'TAG-4', 1, 1, 'Good', 'Description'])
        output = BytesIO(); wb.save(output); output.seek(0); output.name = 'invalid.xlsx'
        response = self.client.post('/api/inventory/import/', {'file': output, 'confirm': 'true'}, format='multipart')
        self.assertEqual(response.status_code, 400)
        self.assertFalse(Asset.objects.filter(asset_tag__in=['TAG-3/1', 'TAG-4/1']).exists())

    def test_office_first_registration_expands_multiple_item_quantities(self):
        desk = InventoryItemType.objects.create(
            name='Office Desk', category=self.category, default_tag_prefix='BPCH/DK',
        )
        self.client.force_authenticate(self.officer)
        response = self.client.post('/api/inventory/office-register/', {
            'location': self.location.id, 'responsible_office': 'Registry',
            'items': [
                {'item_type': self.item_type.id, 'tag_prefix': 'BPCH/CH', 'start_number': 1, 'quantity': 3, 'condition': 'good'},
                {'item_type': desk.id, 'tag_prefix': 'BPCH/DK', 'start_number': 10, 'quantity': 2, 'condition': 'fair'},
            ],
        }, format='json')
        self.assertEqual(response.status_code, 201)
        self.assertEqual(response.data['created'], 5)
        self.assertTrue(Asset.objects.filter(asset_tag='BPCH/CH/3', item_type=self.item_type, quantity=1).exists())
        self.assertTrue(Asset.objects.filter(asset_tag='BPCH/DK/2', item_type=desk, location=self.location).exists())

    def test_office_registration_automatically_continues_item_tag_sequence(self):
        other_office = InventoryLocation.objects.create(name='Second Office')
        Asset.objects.create(
            asset_tag='BPCH/CH/3', name=self.item_type.name, category=self.category,
            item_type=self.item_type, location=self.location, responsible_office='First Office',
            quantity=1, condition='good', created_by=self.officer, updated_by=self.officer,
        )
        self.client.force_authenticate(self.officer)
        response = self.client.post('/api/inventory/office-register/', {
            'location': other_office.id, 'responsible_office': 'Second Office',
            'items': [{'item_type': self.item_type.id, 'tag_prefix': 'BPCH/CH', 'quantity': 3, 'condition': 'good'}],
        }, format='json')
        self.assertEqual(response.status_code, 201)
        self.assertEqual(response.data['first_tag'], 'BPCH/CH/4')
        self.assertEqual(response.data['last_tag'], 'BPCH/CH/6')
        self.assertEqual(
            list(Asset.objects.filter(location=other_office).order_by('asset_tag').values_list('asset_tag', flat=True)),
            ['BPCH/CH/4', 'BPCH/CH/5', 'BPCH/CH/6'],
        )

    def test_auto_numbered_bulk_registration_creates_individual_assets(self):
        self.client.force_authenticate(self.officer)
        response = self.client.post('/api/inventory/bulk-create/', {
            'asset_tag_prefix': 'BPCH/CH', 'start_number': 1, 'count': 30,
            'name': 'Chair', 'description': 'Office chair',
            'category': self.category.id, 'location': self.location.id,
            'responsible_office': 'Administration Office', 'condition': 'good',
        }, format='json')
        self.assertEqual(response.status_code, 201)
        self.assertEqual(response.data['created'], 30)
        self.assertTrue(Asset.objects.filter(asset_tag='BPCH/CH/1', quantity=1).exists())
        self.assertTrue(Asset.objects.filter(asset_tag='BPCH/CH/30', quantity=1).exists())
        self.assertEqual(Asset.objects.filter(asset_tag__startswith='BPCH/CH/').count(), 30)

    def test_auto_numbered_bulk_registration_is_atomic_on_tag_conflict(self):
        self.client.force_authenticate(self.officer)
        Asset.objects.create(
            asset_tag='BPCH/CH/2', name='Existing Chair', category=self.category,
            location=self.location, responsible_office='Office', quantity=1,
            condition='good', created_by=self.officer, updated_by=self.officer,
        )
        response = self.client.post('/api/inventory/bulk-create/', {
            'asset_tag_prefix': 'BPCH/CH', 'start_number': 1, 'count': 3,
            'name': 'Chair', 'category': self.category.id, 'location': self.location.id,
            'responsible_office': 'Administration Office', 'condition': 'good',
        }, format='json')
        self.assertEqual(response.status_code, 400)
        self.assertEqual(Asset.objects.filter(asset_tag__in=['BPCH/CH/1', 'BPCH/CH/3']).count(), 0)


class AttendanceSecurityTests(TestCase):
    def setUp(self):
        self.year = AcademicYear.objects.create(name='2025/2026', is_active=True)
        self.semester = Semester.objects.create(
            academic_year=self.year, number=1, is_active=True
        )
        self.level = ClassLevel.objects.create(name='NTA Level 4', order=4)
        self.teacher = User.objects.create_user('teacher', password='safe-password')
        self.other_teacher = User.objects.create_user('other', password='safe-password')
        self.admin = User.objects.create_superuser(
            'admin', 'admin@example.com', 'safe-password'
        )
        self.module = Module.objects.create(
            name='Business Mathematics',
            code='BM401',
            teacher='Teacher One',
            class_level=self.level,
            semester=self.semester,
        )
        self.other_module = Module.objects.create(
            name='Communication',
            code='CS401',
            teacher='Teacher Two',
            class_level=self.level,
            semester=self.semester,
        )
        self.module.teachers.add(self.teacher)
        self.other_module.teachers.add(self.other_teacher)
        self.student = Student.objects.create(
            nactvet_reg_no='REG-001',
            name='Asha Mollel',
            module=self.module,
        )
        self.student.set_portal_pin('482913')
        self.student.must_change_portal_password = False
        self.student.save(update_fields=['portal_pin_hash', 'must_change_portal_password'])
        self.client = APIClient()

    def test_student_dashboard_redirects_without_student_session(self):
        response = self.client.get(reverse('student-dashboard'))
        self.assertRedirects(response, reverse('login'))

    def test_sessions_expire_when_browser_session_closes(self):
        self.assertTrue(settings.SESSION_EXPIRE_AT_BROWSER_CLOSE)

        response = self.client.post(reverse('login'), {
            'identifier': self.teacher.username,
            'secret': 'safe-password',
        })
        cookie = response.cookies[settings.SESSION_COOKIE_NAME]

        self.assertIn(cookie['expires'], ('', None))
        self.assertIn(cookie['max-age'], ('', None))
        self.assertTrue(cookie['httponly'])
        self.assertEqual(cookie['samesite'], 'Lax')

    def test_staff_logout_invalidates_old_session_key(self):
        self.client.force_login(self.teacher)
        old_session_key = self.client.session.session_key
        self.assertTrue(DjangoSession.objects.filter(session_key=old_session_key).exists())

        response = self.client.get(reverse('logout'))

        self.assertRedirects(response, reverse('login'))
        self.assertFalse(DjangoSession.objects.filter(session_key=old_session_key).exists())
        deletion_cookie = response.cookies[settings.SESSION_COOKIE_NAME]
        self.assertEqual(deletion_cookie.value, '')
        self.assertEqual(deletion_cookie['max-age'], 0)

    def test_student_logout_invalidates_old_session_key(self):
        session = self.client.session
        session['student_id'] = self.student.id
        session['student_reg_no'] = self.student.nactvet_reg_no
        session.save()
        old_session_key = session.session_key

        response = self.client.get(reverse('student-logout'))

        self.assertRedirects(response, reverse('login'))
        self.assertFalse(DjangoSession.objects.filter(session_key=old_session_key).exists())
        self.assertRedirects(
            self.client.get(reverse('student-dashboard')),
            reverse('login'),
        )

    def test_student_can_login_with_portal_pin_not_surname(self):
        self.student.must_change_portal_password = True
        self.student.save(update_fields=['must_change_portal_password'])
        surname_response = self.client.post(reverse('login'), {
            'identifier': self.student.nactvet_reg_no,
            'secret': 'MOLLEL',
        })
        self.assertEqual(surname_response.status_code, 200)
        self.assertContains(surname_response, 'Invalid credentials')

        pin_response = self.client.post(reverse('login'), {
            'identifier': self.student.nactvet_reg_no,
            'secret': '482913',
        })
        self.assertRedirects(pin_response, reverse('student-dashboard'))
        dashboard_response = self.client.get(reverse('student-dashboard'))
        self.assertEqual(dashboard_response.status_code, 200)
        self.assertContains(dashboard_response, 'Create a secure password')

        weak_response = self.client.post('/api/change-password/', {
            'current_password': '482913', 'new_password': 'weakpassword',
        }, format='json')
        self.assertEqual(weak_response.status_code, 400)

        strong_response = self.client.post('/api/change-password/', {
            'current_password': '482913', 'new_password': 'Strong@123',
        }, format='json')
        self.assertEqual(strong_response.status_code, 200)
        self.student.refresh_from_db()
        self.assertFalse(self.student.must_change_portal_password)
        self.assertContains(self.client.get(reverse('student-dashboard')), 'CA Results · Semester 1')

    def test_approved_ca_shows_incomplete_for_unfilled_component(self):
        StudentResult.objects.create(
            student=self.student, ca_approved=True,
            assign1=60, assign2=60, cat1_theory=60,
        )
        session = self.client.session
        session['student_id'] = self.student.id
        session.save()

        response = self.client.get(reverse('student-dashboard'))

        self.assertContains(response, 'INC')

    def test_abscond_is_a_completed_zero_mark_and_teacher_cannot_set_end_absence(self):
        result = StudentResult.objects.create(
            student=self.student, assign1_absent=True,
            assign2=60, cat1_theory=60, cat2_theory=60,
        )
        serialized = StudentResultSerializer(result).data
        self.assertEqual(serialized['assign1_w'], 0.0)
        self.assertIsNotNone(serialized['total_ca'])

        self.client.force_authenticate(self.teacher)
        allowed = self.client.patch(
            f'/api/results/{result.id}/', {'cat1_theory_absent': True}, format='json'
        )
        denied = self.client.patch(
            f'/api/results/{result.id}/', {'end_theory_absent': True}, format='json'
        )
        self.assertEqual(allowed.status_code, 200)
        self.assertEqual(denied.status_code, 403)

    def test_teacher_can_download_module_ca_signoff_docx(self):
        StudentResult.objects.create(
            student=self.student, assign1=60, assign2=60,
            cat1_theory=60, cat2_theory=60,
        )
        self.client.force_login(self.teacher)

        response = self.client.get(
            reverse('results-download-ca-signoff'), {'module_id': self.module.id}
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        document = Document(BytesIO(response.content))
        self.assertIn('CONTINUOUS ASSESSMENT', document.paragraphs[0].text)
        headers = [cell.text for cell in document.tables[0].rows[0].cells]
        self.assertIn('Student Signature', headers)
        self.assertIn('Total CA Average /40', headers)
        self.assertNotIn('Assignment Avg /100', headers)
        self.assertNotIn('CAT Avg /100', headers)
        self.assertNotIn('Practical 1 /100', headers)
        self.assertNotIn('Practical 2 /100', headers)
        self.assertNotIn('Verification / Date', headers)
        self.assertEqual(document.tables[0].rows[1].cells[1].text, self.student.nactvet_reg_no)

    def test_teacher_cannot_download_other_module_ca_signoff(self):
        self.client.force_login(self.teacher)
        response = self.client.get(
            reverse('results-download-ca-signoff'), {'module_id': self.other_module.id}
        )
        self.assertEqual(response.status_code, 403)

    def test_practical_ca_signoff_includes_practical_columns(self):
        self.module.has_practical = True
        self.module.save(update_fields=['has_practical'])
        StudentResult.objects.create(
            student=self.student, assign1=60, assign2=60,
            cat1_theory=60, cat2_theory=60,
            cat1_practical=70, cat2_practical=80,
        )
        self.client.force_login(self.teacher)

        response = self.client.get(
            reverse('results-download-ca-signoff'), {'module_id': self.module.id}
        )

        document = Document(BytesIO(response.content))
        headers = [cell.text for cell in document.tables[0].rows[0].cells]
        self.assertIn('Practical 1 /100', headers)
        self.assertIn('Practical 2 /100', headers)
        self.assertNotIn('Practical Avg /100', headers)

    def test_login_page_does_not_offer_public_registration(self):
        response = self.client.get(reverse('login'))
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'Register here')
        self.assertNotContains(response, 'Accounts are created by the administrator.')

    def test_register_page_is_admin_only(self):
        anonymous = self.client.get(reverse('register'))
        self.assertRedirects(anonymous, reverse('login'))

        self.client.force_login(self.teacher)
        tutor_response = self.client.get(reverse('register'))
        self.assertRedirects(tutor_response, reverse('frontend'))

        self.client.force_login(self.admin)
        admin_response = self.client.get(reverse('register'))
        self.assertEqual(admin_response.status_code, 200)

    def test_admin_can_create_tutor_and_accountant_accounts(self):
        self.client.force_authenticate(self.admin)

        tutor_response = self.client.post('/api/staff-accounts/', {
            'role': 'tutor',
            'full_name': 'Tutor Three',
            'username': 'tutor3',
            'password': 'safe-password',
            'module_ids': [self.module.id],
        }, format='json')
        self.assertEqual(tutor_response.status_code, 201)
        tutor = User.objects.get(username='tutor3')
        self.assertTrue(TeacherProfile.objects.filter(user=tutor).exists())
        self.assertTrue(self.module.teachers.filter(id=tutor.id).exists())

        accountant_response = self.client.post('/api/staff-accounts/', {
            'role': 'accountant',
            'full_name': 'Finance Officer',
            'username': 'finance1',
            'password': 'safe-password',
        }, format='json')
        self.assertEqual(accountant_response.status_code, 201)
        accountant = User.objects.get(username='finance1')
        self.assertTrue(AccountantProfile.objects.filter(user=accountant).exists())

        self.client.force_authenticate(self.teacher)
        denied = self.client.post('/api/staff-accounts/', {
            'role': 'accountant',
            'full_name': 'Blocked User',
            'username': 'blocked',
            'password': 'safe-password',
        }, format='json')
        self.assertEqual(denied.status_code, 403)

    def test_student_dashboard_includes_both_semesters_in_active_year(self):
        semester_two = Semester.objects.create(
            academic_year=self.year, number=2, is_active=False
        )
        semester_two_module = Module.objects.create(
            name='Entrepreneurship',
            code='ENT402',
            teacher='Teacher One',
            class_level=self.level,
            semester=semester_two,
        )
        Student.objects.create(
            nactvet_reg_no=self.student.nactvet_reg_no,
            name=self.student.name,
            module=semester_two_module,
        )
        session = self.client.session
        session['student_id'] = self.student.id
        session.save()

        response = self.client.get(reverse('student-dashboard'))

        self.assertContains(response, 'Business Mathematics')
        self.assertContains(response, 'Entrepreneurship')

    def test_teacher_cannot_create_session_for_unassigned_module(self):
        self.client.force_authenticate(self.teacher)
        response = self.client.post('/api/sessions/', {
            'module': self.other_module.id,
            'session_type': Session.THEORY,
            'exam_period': Session.GENERAL,
            'date': str(date.today()),
            'label': 'Week 1',
            'topic': 'Introduction',
            'records': [],
        }, format='json')
        self.assertEqual(response.status_code, 403)

    def test_hims_fill_command_creates_sheet_sessions_and_preserves_existing_records(self):
        self.module.code = 'PST05209'
        self.module.name = 'Health Information Management'
        self.module.save(update_fields=['code', 'name'])
        existing = Session.objects.create(
            module=self.module,
            session_type=Session.THEORY,
            exam_period=Session.GENERAL,
            date=date(2026, 4, 7),
            label='01',
        )
        AttendanceRecord.objects.create(
            session=existing,
            student=self.student,
            status=AttendanceRecord.ABSENT,
        )

        dry_run_output = StringIO()
        call_command('fill_hims_attendance', stdout=dry_run_output)
        self.assertEqual(Session.objects.filter(module=self.module).count(), 1)

        call_command('fill_hims_attendance', '--confirm', stdout=StringIO())

        self.assertEqual(Session.objects.filter(module=self.module).count(), 17)
        self.assertFalse(Session.objects.filter(module=self.module, date=date(2026, 5, 20)).exists())
        self.assertEqual(
            AttendanceRecord.objects.get(session=existing, student=self.student).status,
            AttendanceRecord.ABSENT,
        )
        self.assertEqual(
            AttendanceRecord.objects.filter(
                session__module=self.module,
                status=AttendanceRecord.PRESENT,
            ).count(),
            16,
        )

        call_command('fill_hims_attendance', '--confirm', stdout=StringIO())
        self.assertEqual(Session.objects.filter(module=self.module).count(), 17)

    def test_pst04209_fill_command_loads_level_four_sheet_idempotently(self):
        self.module.code = 'PST04209'
        self.module.name = 'Level Four Module'
        self.module.save(update_fields=['code', 'name'])
        existing = Session.objects.create(
            module=self.module,
            session_type=Session.THEORY,
            exam_period=Session.GENERAL,
            date=date(2026, 4, 10),
            label='01',
        )
        AttendanceRecord.objects.create(
            session=existing,
            student=self.student,
            status=AttendanceRecord.ABSENT,
        )

        call_command('fill_pst04209_attendance', stdout=StringIO())
        self.assertEqual(Session.objects.filter(module=self.module).count(), 1)

        call_command('fill_pst04209_attendance', '--confirm', stdout=StringIO())

        self.assertEqual(Session.objects.filter(module=self.module).count(), 15)
        self.assertEqual(
            AttendanceRecord.objects.get(session=existing, student=self.student).status,
            AttendanceRecord.ABSENT,
        )
        self.assertEqual(
            AttendanceRecord.objects.filter(
                session__module=self.module,
                status=AttendanceRecord.PRESENT,
            ).count(),
            14,
        )
        self.assertEqual(
            Session.objects.filter(
                module=self.module, label='09', date__in=(date(2026, 4, 20), date(2026, 4, 25))
            ).count(),
            2,
        )

        call_command('fill_pst04209_attendance', '--confirm', stdout=StringIO())
        self.assertEqual(Session.objects.filter(module=self.module).count(), 15)

    def test_pst05208_fill_command_loads_level_five_sheet_idempotently(self):
        level_five = ClassLevel.objects.create(name='NTA Level 5', order=5)
        self.module.code = 'PST05208'
        self.module.name = 'Level Five Module'
        self.module.class_level = level_five
        self.module.save(update_fields=['code', 'name', 'class_level'])
        existing = Session.objects.create(
            module=self.module,
            session_type=Session.THEORY,
            exam_period=Session.GENERAL,
            date=date(2026, 4, 10),
            label='01',
        )
        AttendanceRecord.objects.create(
            session=existing,
            student=self.student,
            status=AttendanceRecord.ABSENT,
        )

        call_command('fill_pst05208_attendance', stdout=StringIO())
        self.assertEqual(Session.objects.filter(module=self.module).count(), 1)

        call_command('fill_pst05208_attendance', '--confirm', stdout=StringIO())

        self.assertEqual(Session.objects.filter(module=self.module).count(), 17)
        self.assertEqual(
            AttendanceRecord.objects.get(session=existing, student=self.student).status,
            AttendanceRecord.ABSENT,
        )
        self.assertEqual(
            AttendanceRecord.objects.filter(
                session__module=self.module,
                status=AttendanceRecord.PRESENT,
            ).count(),
            16,
        )
        self.assertEqual(
            Session.objects.filter(
                module=self.module,
                date=date(2026, 4, 25),
                label__in=('08', '08 (2)'),
            ).count(),
            2,
        )

        call_command('fill_pst05208_attendance', '--confirm', stdout=StringIO())
        self.assertEqual(Session.objects.filter(module=self.module).count(), 17)

    def test_duplicate_attendance_session_is_rejected_and_can_be_deleted(self):
        self.client.force_authenticate(self.teacher)
        payload = {
            'module': self.module.id,
            'session_type': Session.THEORY,
            'exam_period': Session.GENERAL,
            'date': str(date.today()),
            'label': 'Week 1',
            'topic': 'Introduction',
            'records': [{'nactvet_reg_no': self.student.nactvet_reg_no, 'status': 'P'}],
        }

        created = self.client.post('/api/sessions/', payload, format='json')
        duplicate = self.client.post(
            '/api/sessions/', {**payload, 'label': ' week 1 '}, format='json'
        )

        self.assertEqual(created.status_code, 201)
        self.assertEqual(duplicate.status_code, 400)
        self.assertIn('already been recorded', str(duplicate.data))
        deleted = self.client.delete(f'/api/sessions/{created.data["id"]}/')
        self.assertEqual(deleted.status_code, 204)
        self.assertFalse(Session.objects.filter(id=created.data['id']).exists())

    def test_teacher_can_correct_attendance_for_assigned_session(self):
        session = Session.objects.create(
            module=self.module, session_type=Session.THEORY,
            exam_period=Session.GENERAL, date=date.today(), label='Week 2'
        )
        record = AttendanceRecord.objects.create(
            session=session, student=self.student, status=AttendanceRecord.PRESENT
        )
        self.client.force_authenticate(self.teacher)

        response = self.client.patch(f'/api/sessions/{session.id}/', {
            'records': [{
                'nactvet_reg_no': self.student.nactvet_reg_no,
                'status': AttendanceRecord.SICK,
                'sick_note': 'Clinic visit',
            }],
        }, format='json')

        self.assertEqual(response.status_code, 200)
        record.refresh_from_db()
        self.assertEqual(record.status, AttendanceRecord.SICK)
        self.assertEqual(record.sick_note, 'Clinic visit')

    def test_teacher_cannot_correct_attendance_for_unassigned_session(self):
        session = Session.objects.create(
            module=self.other_module, session_type=Session.THEORY,
            exam_period=Session.GENERAL, date=date.today(), label='Week 2'
        )
        self.client.force_authenticate(self.teacher)

        response = self.client.patch(f'/api/sessions/{session.id}/', {
            'records': [],
        }, format='json')

        self.assertEqual(response.status_code, 404)

    def test_end_exam_below_50_requires_supplementary_only_for_end_exam(self):
        result = StudentResult.objects.create(
            student=self.student,
            assign1=20, assign2=20, cat1_theory=20, cat2_theory=20,
            end_theory=49,
        )

        data = StudentResultSerializer(result).data

        self.assertEqual(data['result_status'], 'SUPP')
        self.assertTrue(data['supplementary_required'])
        self.assertEqual(data['end_exam_mark'], 49.0)

    def test_passed_supplementary_is_capped_at_c_and_two_points(self):
        result = StudentResult.objects.create(
            student=self.student,
            assign1=100, assign2=100, cat1_theory=100, cat2_theory=100,
            end_theory=49, supplementary_mark=95,
        )

        data = StudentResultSerializer(result).data

        self.assertEqual(data['grade'], 'C')
        self.assertEqual(data['grade_point'], 2)
        self.assertEqual(data['result_status'], 'PASS')

    def test_practical_module_requires_supp_if_either_end_component_is_below_half(self):
        self.module.has_practical = True
        self.module.save(update_fields=['has_practical'])
        result = StudentResult.objects.create(
            student=self.student,
            assign1=100, assign2=100,
            cat1_theory=100, cat2_theory=100,
            cat1_practical=100, cat2_practical=100,
            end_theory=100, end_practical=49,
        )

        data = StudentResultSerializer(result).data

        self.assertEqual(data['final_total'], 84.7)
        self.assertEqual(data['end_exam_mark'], 74.5)
        self.assertEqual(data['result_status'], 'SUPP')
        self.assertTrue(data['supplementary_required'])
        self.assertEqual(data['failed_end_components'], ['end_practical'])
        self.assertIsNone(data['grade'])

    def test_practical_module_passes_when_both_end_components_reach_half(self):
        self.module.has_practical = True
        self.module.save(update_fields=['has_practical'])
        result = StudentResult.objects.create(
            student=self.student,
            assign1=100, assign2=100,
            cat1_theory=100, cat2_theory=100,
            cat1_practical=100, cat2_practical=100,
            end_theory=50, end_practical=50,
        )

        data = StudentResultSerializer(result).data

        self.assertFalse(data['supplementary_required'])
        self.assertEqual(data['result_status'], 'PASS')

    def test_failed_supplementary_requires_repeat(self):
        result = StudentResult.objects.create(
            student=self.student,
            assign1=100, assign2=100, cat1_theory=100, cat2_theory=100,
            end_theory=49, supplementary_mark=49,
        )

        data = StudentResultSerializer(result).data

        self.assertEqual(data['result_status'], 'REPEAT')
        self.assertEqual(data['grade_point'], 0)

    def test_level_six_uses_five_point_grading_scale(self):
        level_six = ClassLevel.objects.create(name='NTA Level 6', order=6)
        module = Module.objects.create(
            name='Level Six Module', code='L6M', teacher='Teacher',
            class_level=level_six, semester=self.semester, credits=10,
        )
        student = Student.objects.create(
            nactvet_reg_no='L6-001', name='Level Six Student', module=module,
        )
        result = StudentResult.objects.create(
            student=student,
            assign1=100, assign2=100, cat1_theory=100, cat2_theory=100,
            end_theory=60,
        )

        data = StudentResultSerializer(result).data

        self.assertEqual(data['final_total'], 76.0)
        self.assertEqual(data['grade'], 'A')
        self.assertEqual(data['grade_point'], 5)
        self.assertEqual(gpa_classification(4.6, level_six), 'First Class')

    def test_final_grade_uses_ca_40_plus_ese_60_total(self):
        result = StudentResult.objects.create(
            student=self.student,
            assign1=100, assign2=100,
            cat1_theory=100, cat2_theory=100,
            end_theory=60,
        )

        data = StudentResultSerializer(result).data

        self.assertEqual(data['total_ca'], 40.0)
        self.assertEqual(data['end_exam_total'], 36.0)
        self.assertEqual(data['final_total'], 76.0)
        self.assertEqual(data['grade'], 'B')
        self.assertEqual(data['grade_point'], 3)

    def test_student_portal_displays_module_credits_and_weighted_result_parts(self):
        self.module.credits = 12
        self.module.save(update_fields=['credits'])
        StudentResult.objects.create(
            student=self.student,
            assign1=100, assign2=100,
            cat1_theory=100, cat2_theory=100,
            end_theory=60, ca_approved=True, final_approved=True,
        )
        session = self.client.session
        session['student_id'] = self.student.id
        session.save()

        response = self.client.get(reverse('student-dashboard'))

        self.assertContains(response, '12 credits')
        self.assertContains(response, 'ESE /60')
        self.assertNotContains(response, '<th>Grade Point</th>', html=True)

    def test_student_ca_table_uses_compact_headers_and_average_status(self):
        StudentResult.objects.create(
            student=self.student,
            assign1_absent=True,
            assign2=60,
            cat1_theory=60,
            cat2_theory=60,
            ca_approved=True,
        )
        session = self.client.session
        session['student_id'] = self.student.id
        session.save()

        response = self.client.get(reverse('student-dashboard'))

        self.assertContains(response, '<th>Module Name</th>', html=True)
        self.assertContains(response, '<th>A1</th>', html=True)
        self.assertContains(response, '<th>Average CA /40</th>', html=True)
        module = response.context['semester1_modules'][0]
        self.assertEqual(module['result']['ca_display'], 'ABS')
        self.assertIsNone(module['result']['assign1'])
        self.assertNotContains(response, 'Recent attendance activity')
        self.assertContains(response, 'Payment History')
        self.assertContains(response, 'Academic Obligations')
        self.assertContains(response, 'Theory Modules')
        theory_table = response.context['semester1_theory_modules']
        self.assertEqual(len(theory_table), 1)
        self.assertEqual(response.context['semester1_practical_modules'], [])

    def test_cat_analysis_reports_module_grades_and_total_pass_rate(self):
        StudentResult.objects.create(
            student=self.student,
            cat1_theory=80,
        )
        failing_student = Student.objects.create(
            nactvet_reg_no='REG-FAIL', name='Failing Student', module=self.module,
        )
        StudentResult.objects.create(
            student=failing_student,
            cat1_theory=40,
        )
        incomplete_student = Student.objects.create(
            nactvet_reg_no='REG-INC', name='Incomplete Student', module=self.module,
        )
        StudentResult.objects.create(student=incomplete_student)
        self.client.force_authenticate(self.admin)

        response = self.client.get(
            f'/api/results/cat-analysis/?cat=1&semester_id={self.semester.id}'
        )

        self.assertEqual(response.status_code, 200)
        row = next(r for r in response.data['rows'] if r['module_id'] == self.module.id)
        self.assertEqual(row['assessed'], 2)
        self.assertEqual(row['passed'], 1)
        self.assertEqual(row['failed'], 1)
        self.assertEqual(row['incomplete'], 1)
        self.assertEqual(row['grade_counts']['A'], 1)
        self.assertEqual(row['grade_counts']['D'], 1)
        self.assertEqual(row['pass_rate'], 50.0)
        self.assertEqual(response.data['stats']['pass_rate'], 50.0)

    def test_cat_analysis_averages_theory_and_practical_components(self):
        self.module.has_practical = True
        self.module.save(update_fields=['has_practical'])
        StudentResult.objects.create(
            student=self.student,
            cat2_theory=80,
            cat2_practical=40,
        )
        self.client.force_authenticate(self.admin)

        response = self.client.get(
            f'/api/results/cat-analysis/?cat=2&module_id={self.module.id}'
            f'&semester_id={self.semester.id}'
        )

        self.assertEqual(response.status_code, 200)
        row = response.data['rows'][0]
        self.assertEqual(row['assessed'], 1)
        self.assertEqual(row['passed'], 1)
        self.assertEqual(row['grade_counts']['C'], 1)
        self.assertEqual(row['pass_rate'], 100.0)

    def test_cat_analysis_download_is_formatted_excel_report(self):
        StudentResult.objects.create(student=self.student, cat1_theory=80)
        self.client.force_authenticate(self.admin)

        response = self.client.get(
            f'/api/results/cat-analysis/download/?cat=1&semester_id={self.semester.id}'
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        workbook = load_workbook(BytesIO(response.content))
        sheet = workbook.active
        self.assertEqual(sheet['A1'].value, 'EDUTRACK — CAT 1 MODULE ANALYSIS')
        self.assertEqual(sheet['A7'].value, 'Module Code')
        self.assertEqual(sheet.freeze_panes, 'A8')
        self.assertEqual(sheet.sheet_view.showGridLines, False)
        self.assertEqual(sheet['A8'].value, self.module.code)

    def test_cat_analysis_requires_semester_filter(self):
        self.client.force_authenticate(self.admin)

        response = self.client.get('/api/results/cat-analysis/?cat=1')

        self.assertEqual(response.status_code, 400)
        self.assertEqual(response.data['detail'], 'Select a semester for CAT analysis.')
        self.assertFalse(Session.objects.exists())

    def test_teacher_cannot_move_student_to_unassigned_module(self):
        self.client.force_authenticate(self.teacher)
        response = self.client.patch(
            f'/api/students/{self.student.id}/',
            {'module': self.other_module.id},
            format='json',
        )
        self.assertEqual(response.status_code, 403)
        self.student.refresh_from_db()
        self.assertEqual(self.student.module, self.module)

    def test_teacher_cannot_modify_class_levels_or_academic_years(self):
        self.client.force_authenticate(self.teacher)
        level_response = self.client.post(
            '/api/class-levels/', {'name': 'NTA Level 5', 'order': 5}, format='json'
        )
        year_response = self.client.post(
            '/api/academic-years/', {'name': '2026/2027'}, format='json'
        )
        self.assertEqual(level_response.status_code, 403)
        self.assertEqual(year_response.status_code, 403)

    def test_portal_pin_is_write_only_and_hashed(self):
        self.client.force_authenticate(self.teacher)
        response = self.client.patch(
            f'/api/students/{self.student.id}/',
            {'portal_pin': 'new-pin-42'},
            format='json',
        )
        self.assertEqual(response.status_code, 200)
        self.assertNotIn('portal_pin', response.data)
        self.student.refresh_from_db()
        self.assertNotEqual(self.student.portal_pin_hash, 'new-pin-42')
        self.assertTrue(self.student.check_portal_pin('new-pin-42'))

    def test_admin_can_set_password_for_filtered_students(self):
        other_level = ClassLevel.objects.create(name='NTA Level 5', order=5)
        other_module = Module.objects.create(
            name='Accounting',
            code='AC501',
            teacher='Teacher Three',
            class_level=other_level,
            semester=self.semester,
        )
        other_student = Student.objects.create(
            nactvet_reg_no='REG-002',
            name='Baraka John',
            module=other_module,
        )
        self.client.force_authenticate(self.admin)

        response = self.client.post('/api/students/bulk_set_pin/', {
            'class_level_id': self.level.id,
            'portal_pin': 'shared-123',
        }, format='json')

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data['updated'], 1)
        self.student.refresh_from_db()
        other_student.refresh_from_db()
        self.assertTrue(self.student.check_portal_pin('shared-123'))
        self.assertTrue(self.student.must_change_portal_password)
        self.assertFalse(other_student.check_portal_pin('shared-123'))

    def test_authenticated_user_can_change_own_password(self):
        self.client.force_login(self.teacher)
        response = self.client.post('/api/change-password/', {
            'current_password': 'safe-password',
            'new_password': 'new-safe-password',
        }, format='json')

        self.assertEqual(response.status_code, 200)
        self.teacher.refresh_from_db()
        self.assertTrue(self.teacher.check_password('new-safe-password'))

    def test_student_can_change_pin_for_all_enrollments(self):
        second_enrollment = Student.objects.create(
            nactvet_reg_no=self.student.nactvet_reg_no,
            name=self.student.name,
            module=self.other_module,
        )
        session = self.client.session
        session['student_id'] = self.student.id
        session.save()

        response = self.client.post('/api/change-password/', {
            'current_password': '482913',
            'new_password': 'Student@123',
        }, format='json')

        self.assertEqual(response.status_code, 200)
        self.student.refresh_from_db()
        second_enrollment.refresh_from_db()
        self.assertTrue(self.student.check_portal_pin('Student@123'))
        self.assertTrue(second_enrollment.check_portal_pin('Student@123'))
        self.assertFalse(self.student.must_change_portal_password)
        self.assertFalse(second_enrollment.must_change_portal_password)

    def test_bulk_result_save_preserves_omitted_fields(self):
        result = StudentResult.objects.create(
            student=self.student,
            assign1=80,
            assign2=75,
            cat1_theory=70,
            cat2_theory=65,
        )
        self.client.force_authenticate(self.teacher)

        response = self.client.post('/api/results/bulk_save/', [
            {'id': result.id, 'assign1': 90},
        ], format='json')

        self.assertEqual(response.status_code, 200)
        result.refresh_from_db()
        self.assertEqual(float(result.assign1), 90)
        self.assertEqual(float(result.assign2), 75)
        self.assertEqual(float(result.cat1_theory), 70)
        self.assertEqual(float(result.cat2_theory), 65)

    def test_teacher_cannot_approve_ca_results(self):
        result = StudentResult.objects.create(student=self.student, assign1=80)
        self.client.force_authenticate(self.teacher)

        response = self.client.post('/api/results/bulk_save/', [
            {'id': result.id, 'ca_approved': True},
        ], format='json')

        self.assertEqual(response.status_code, 200)
        result.refresh_from_db()
        self.assertFalse(result.ca_approved)
        self.assertTrue(response.data['errors'])

    def test_admin_can_approve_ca_results(self):
        result = StudentResult.objects.create(student=self.student, assign1=80)
        self.client.force_authenticate(self.admin)

        response = self.client.post('/api/results/bulk_save/', [
            {'id': result.id, 'ca_approved': True},
        ], format='json')

        self.assertEqual(response.status_code, 200)
        result.refresh_from_db()
        self.assertTrue(result.ca_approved)

    def test_student_dashboard_hides_ca_results_until_admin_approval(self):
        StudentResult.objects.create(
            student=self.student,
            assign1=80,
            assign2=70,
            cat1_theory=60,
            cat2_theory=50,
            ca_approved=False,
        )
        session = self.client.session
        session['student_id'] = self.student.id
        session.save()

        response = self.client.get(reverse('student-dashboard'))

        self.assertContains(response, 'CA results for Semester 1 have not been published yet.')
        self.assertNotContains(response, '80.00')

    def test_student_dashboard_shows_ca_results_after_admin_approval(self):
        StudentResult.objects.create(
            student=self.student,
            assign1=80,
            assign2=70,
            cat1_theory=60,
            cat2_theory=50,
            ca_approved=True,
        )
        session = self.client.session
        session['student_id'] = self.student.id
        session.save()

        response = self.client.get(reverse('student-dashboard'))

        self.assertContains(response, '80.00')
        self.assertNotContains(response, 'CA results for Semester 1 have not been published yet.')

    def test_student_dashboard_hides_final_results_until_admin_approval(self):
        StudentResult.objects.create(
            student=self.student,
            assign1=80,
            assign2=70,
            cat1_theory=60,
            cat2_theory=50,
            end_theory=90,
            final_approved=False,
        )
        session = self.client.session
        session['student_id'] = self.student.id
        session.save()

        response = self.client.get(reverse('student-dashboard'))

        self.assertContains(response, 'CA results for Semester 1 have not been published yet.')
        self.assertContains(response, 'CA Results · Semester 1')
        self.assertContains(response, 'Final examination marks have not been published yet.')
        self.assertNotContains(response, '80.00')
        self.assertNotContains(response, '90.00')
        self.assertNotContains(response, '78.0')

    def test_student_dashboard_shows_final_results_after_admin_approval(self):
        StudentResult.objects.create(
            student=self.student,
            assign1=80,
            assign2=70,
            cat1_theory=60,
            cat2_theory=50,
            end_theory=90,
            final_approved=True,
        )
        session = self.client.session
        session['student_id'] = self.student.id
        session.save()

        response = self.client.get(reverse('student-dashboard'))

        self.assertContains(response, '90.00')
        self.assertContains(response, '78.0')

    def test_eligibility_api_and_excel_both_require_sick_certificate(self):
        session = Session.objects.create(
            module=self.module,
            session_type=Session.THEORY,
            exam_period=Session.CAT1,
            date=date.today(),
            label='CAT preparation',
        )
        AttendanceRecord.objects.create(
            session=session,
            student=self.student,
            status=AttendanceRecord.SICK,
            certificate_submitted=False,
        )
        self.client.force_authenticate(self.admin)

        api_response = self.client.get('/api/eligibility/', {'module_id': self.module.id})
        self.assertEqual(api_response.status_code, 200)
        self.assertEqual(api_response.data['rows'][0]['cat1_attended'], 0)

        self.client.force_authenticate(user=None)
        self.client.force_login(self.admin)
        excel_response = self.client.get(
            '/api/eligibility/download/', {'module_id': self.module.id}
        )
        self.assertEqual(excel_response.status_code, 200)
        workbook = load_workbook(BytesIO(excel_response.content))
        sheet = workbook.active
        self.assertEqual(sheet.cell(row=2, column=9).value, 0)

    def test_final_eligibility_excel_has_three_sheets_per_level_and_colours(self):
        session = Session.objects.create(
            module=self.module,
            session_type=Session.THEORY,
            exam_period=Session.GENERAL,
            date=date.today(),
            label='Final eligibility session',
        )
        AttendanceRecord.objects.create(
            session=session,
            student=self.student,
            status=AttendanceRecord.PRESENT,
        )
        StudentResult.objects.create(
            student=self.student,
            assign1=60,
            assign2=60,
            cat1_theory=60,
            cat2_theory=60,
        )
        self.other_module.has_practical = True
        self.other_module.save(update_fields=['has_practical'])
        practical_student = Student.objects.create(
            nactvet_reg_no='REG-002',
            name='Baraka Juma',
            module=self.other_module,
        )
        StudentResult.objects.create(
            student=practical_student,
            assign1=70,
            assign2=70,
            cat1_theory=70,
            cat2_theory=70,
            cat1_practical=70,
            cat2_practical=70,
        )
        self.client.force_login(self.admin)

        response = self.client.get('/api/eligibility/final/download/')

        self.assertEqual(response.status_code, 200)
        workbook = load_workbook(BytesIO(response.content))
        self.assertEqual(workbook.sheetnames, [
            'NTA 4 Assignments', 'NTA 4 CATs', 'NTA 4 Eligibility',
            'NTA 5 Assignments', 'NTA 5 CATs', 'NTA 5 Eligibility',
            'NTA 6 Assignments', 'NTA 6 CATs', 'NTA 6 Eligibility',
        ])
        eligibility_sheet = workbook['NTA 4 Eligibility']
        self.assertEqual(eligibility_sheet['J8'].value, 'Eligible')
        self.assertEqual(eligibility_sheet['J8'].fill.fgColor.rgb, '00C6EFCE')
        assignments_sheet = workbook['NTA 4 Assignments']
        self.assertEqual(assignments_sheet['D7'].value, 'ASS1')
        self.assertEqual(assignments_sheet['E7'].value, 'ASS2')
        self.assertEqual(assignments_sheet['D8'].value, 60)
        self.assertEqual(assignments_sheet['E8'].value, 60)
        self.assertEqual(workbook['NTA 4 CATs']['H8'].value, 24)
        cats_sheet = workbook['NTA 4 CATs']
        self.assertNotEqual(cats_sheet['D6'].value, 'Practical')
        self.assertEqual(cats_sheet['N6'].value, 'Practical')
        self.assertEqual(cats_sheet['Q6'].value, 'Total')

    def test_accountant_can_manage_payments_but_tutor_cannot(self):
        accountant = User.objects.create_user('accountant', password='safe-password')
        AccountantProfile.objects.create(user=accountant, full_name='Finance Officer')

        self.client.force_authenticate(self.teacher)
        denied = self.client.get('/api/payment-categories/')
        self.assertEqual(denied.status_code, 403)

        self.client.force_authenticate(accountant)
        category_response = self.client.post('/api/payment-categories/', {
            'name': 'Semester 1 School Fees',
            'category_type': PaymentCategory.SCHOOL_FEES,
            'default_amount': '1000.00',
            'installment_count': 2,
            'semester': self.semester.id,
            'class_level': self.level.id,
            'is_active': True,
        })
        self.assertEqual(category_response.status_code, 201)
        self.assertEqual(category_response.data['installment_count'], 2)

        payment_response = self.client.post('/api/student-payments/', {
            'student': self.student.id,
            'category': category_response.data['id'],
            'amount_required': '1000.00',
            'amount_paid': '600.00',
            'installment_number': 2,
            'payment_date': str(date.today()),
            'reference': 'RCPT-001',
        })
        self.assertEqual(payment_response.status_code, 201)
        self.assertEqual(payment_response.data['balance'], '400.00')
        self.assertEqual(StudentPayment.objects.count(), 1)

        blocked_installment = self.client.post('/api/student-payments/', {
            'student': self.student.id,
            'category': category_response.data['id'],
            'amount_required': '1000.00',
            'amount_paid': '100.00',
            'installment_number': 3,
            'payment_date': str(date.today()),
        })
        self.assertEqual(blocked_installment.status_code, 403)

        special_category = self.client.post('/api/payment-categories/', {
            'name': 'Special Exam Fee',
            'category_type': PaymentCategory.SPECIAL_EXAM,
            'default_amount': '150.00',
            'installment_count': 1,
            'semester': self.semester.id,
            'class_level': self.level.id,
            'is_active': True,
        })
        self.assertEqual(special_category.status_code, 201)

        self.client.force_authenticate(self.admin)
        obligation_response = self.client.post('/api/finance-obligations/', {
            'student': self.student.id,
            'semester': self.semester.id,
            'obligation_type': StudentFinanceObligation.SPECIAL_EXAM,
            'category': special_category.data['id'],
            'amount_required': '150.00',
        })
        self.assertEqual(obligation_response.status_code, 201)

        self.client.force_authenticate(accountant)
        obligation_payment = self.client.post('/api/student-payments/', {
            'student': self.student.id,
            'category': special_category.data['id'],
            'obligation': obligation_response.data['id'],
            'amount_required': '150.00',
            'amount_paid': '150.00',
            'installment_number': 1,
            'payment_date': str(date.today()),
        })
        self.assertEqual(obligation_payment.status_code, 201)

        clearance_response = self.client.post('/api/finance-clearances/', {
            'student': self.student.id,
            'semester': self.semester.id,
            'period': StudentFinanceClearance.CAT1,
            'is_cleared': True,
            'note': 'Paid special exam fee.',
        })
        self.assertEqual(clearance_response.status_code, 201)
        self.assertTrue(clearance_response.data['is_cleared'])
