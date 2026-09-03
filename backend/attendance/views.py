import json
from decimal import Decimal
from functools import wraps
import re
from io import BytesIO

from django.conf import settings
from django.contrib.auth import get_user_model
from django.contrib.auth import login, logout, authenticate, update_session_auth_hash
from django.contrib.auth.decorators import login_required
from django.db import transaction
from django.db.models import Count, Q, Sum
from django.http import FileResponse, Http404, HttpResponse, HttpResponseForbidden
from django.shortcuts import get_object_or_404, render, redirect
from django.utils import timezone
from django.views.decorators.clickjacking import xframe_options_sameorigin
from django.utils.dateparse import parse_date
from django.utils.http import url_has_allowed_host_and_scheme
from rest_framework import mixins, viewsets, status
from rest_framework.decorators import action, api_view
from rest_framework.exceptions import PermissionDenied
from rest_framework.permissions import BasePermission, IsAuthenticated, SAFE_METHODS
from rest_framework.response import Response
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.worksheet.datavalidation import DataValidation

from . import analytics, evaluations, finance
from .forms import TeacherRegistrationForm, StyledAuthForm, StudentLoginForm
from .models import (
    AcademicYear, Semester, ClassLevel, Module,
    Student, Session, AttendanceRecord, TeacherProfile, AccountantProfile,
    StudentResult, PaymentCategory, StudentFinanceObligation, StudentPayment,
    StudentFinanceClearance, Announcement,
    EstateOfficerProfile, SecretaryProfile, PrincipalProfile, HeadOfDepartmentProfile,
    ResultEntryWindow,
    InventoryLocation, AssetCategory, InventoryItemType, Asset, AssetImport,
    AssetTransfer, AssetMaintenance, InventoryInspection, InventoryInspectionItem, AssetDisposal,
    ChargeType, FeeStructure, FeeInstallment, StudentProfile, StudentCharge,
    Invoice, InvoiceLine, Payment, PaymentAllocation, FinanceOverride, FinanceAuditLog,
    BankAccount, CollegeProfile,
    Form, FormSection, FormQuestion, FormAnswer, FormResponse, FormSubmissionReceipt,
)
from .serializers import (
    AcademicYearSerializer, SemesterSerializer, ClassLevelSerializer,
    ModuleSerializer, StudentSerializer,
    SessionSerializer, SessionCreateSerializer, BulkStudentSerializer,
    StudentResultSerializer, AttendanceRecordSerializer,
    FinanceStudentSerializer, PaymentCategorySerializer,
    StudentFinanceObligationSerializer, StudentPaymentSerializer,
    StudentFinanceClearanceSerializer, AnnouncementSerializer,
    InventoryLocationSerializer, AssetCategorySerializer, InventoryItemTypeSerializer, AssetSerializer,
    AssetTransferSerializer, AssetMaintenanceSerializer, InventoryInspectionSerializer,
    InventoryInspectionItemSerializer, AssetDisposalSerializer,
    BankAccountSerializer, CollegeProfileSerializer,
    ChargeTypeSerializer, FeeStructureSerializer, StudentChargeSerializer,
    InvoiceSerializer, PaymentSerializer, RecordPaymentSerializer,
    ReversePaymentSerializer, WaiveChargeSerializer, FinanceOverrideSerializer,
    FinanceAuditLogSerializer,
    FormSerializer, FormSectionSerializer, FormQuestionSerializer,
    FormResponseSerializer, StudentFormSerializer, ResultEntryWindowSerializer,
)
from .grading import grade_for_mark, gpa_classification, parse_authority_grade

User = get_user_model()

STUDENT_PASSWORD_MESSAGE = (
    'Password must be at least 8 characters and include an uppercase letter, '
    'a lowercase letter, a number, and a symbol.'
)


def student_password_is_strong(password):
    return (
        len(password) >= 8
        and re.search(r'[A-Z]', password)
        and re.search(r'[a-z]', password)
        and re.search(r'\d', password)
        and re.search(r'[^A-Za-z0-9]', password)
    )


# ── HELPERS ────────────────────────────────────────────────────────────────────

def user_modules(user):
    """The modules this account is working with.

    An admin role covers the whole college. One who also teaches can narrow it
    to their own modules — see TeachingScopeMiddleware — because a Principal
    promoted from tutoring still has classes to take attendance for, and finding
    them in a list of every module the college runs is not a promotion.

    Narrowing only. The switch cannot show anybody a module they could not
    already reach, and it changes nothing about what they are allowed to do.
    """
    if getattr(user, 'teaching_scope_only', False):
        return user.modules_taught.all()
    if user.is_staff:
        return Module.objects.all()
    return user.modules_taught.all()


def teaches_anything(user):
    return bool(user and user.is_authenticated and user.modules_taught.exists())


def is_accountant(user):
    return bool(
        user
        and user.is_authenticated
        and AccountantProfile.objects.filter(user=user, is_active=True).exists()
    )


def can_manage_finance(user):
    """Only the accountant handles money.

    `is_staff` in this system is the examination officer / HOD: they own
    results, eligibility and the academic register, and deliberately have no
    access to fees, payments or the ledger. Separating the two is the point —
    the person who decides who sits an exam is not the person who decides
    whether their money arrived.
    """
    return is_accountant(user)


def is_estate_officer(user):
    return bool(
        user and user.is_authenticated
        and EstateOfficerProfile.objects.filter(user=user, is_active=True).exists()
    )


def is_principal(user):
    return bool(
        user and user.is_authenticated
        and PrincipalProfile.objects.filter(user=user, is_active=True).exists()
    )


def is_head_of_department(user):
    return bool(
        user and user.is_authenticated
        and HeadOfDepartmentProfile.objects.filter(user=user, is_active=True).exists()
    )


def can_manage_exams(user):
    """Entering a mark, approving a result, declaring a student for an exam.

    Every admin role except the Head of Department, who reads examinations but
    does not decide them: the person who runs the department is not the person
    who decides who sits an exam. Money and college property are already outside
    every admin account — they need the accountant's or the estate officer's own
    profile.
    """
    return bool(
        user and user.is_authenticated
        and user.is_staff and not is_head_of_department(user)
    )


#: The three families of writable fields on a result. They are governed
#: separately because they close at different times and by different people.
CA_MARK_FIELDS = frozenset({
    'field_ca', 'assign1', 'assign2', 'cat1_theory', 'cat2_theory',
    'cat1_practical', 'cat2_practical',
    'assign1_absent', 'assign2_absent', 'cat1_theory_absent', 'cat2_theory_absent',
    'cat1_practical_absent', 'cat2_practical_absent',
})
END_MARK_FIELDS = frozenset({
    'end_theory', 'end_practical', 'end_theory_absent', 'end_practical_absent',
})
EXAM_OFFICE_FIELDS = frozenset({
    'ca_approved', 'final_approved', 'supplementary_mark',
    'authority_grade', 'authority_status',
})


def check_result_entry(user, module, data):
    """May this person write these marks, on this module, today?

    A tutor enters continuous assessment only — the CATs and assignments they
    set and marked themselves — and only while the examination officer has the
    books open. The end of semester examination, the supplementary and every
    approval belong to the examination officer whatever the calendar says: the
    person who marked the coursework is not the person who publishes the
    result. The Head of Department reads results and never writes them.
    """
    if can_manage_exams(user):
        return
    if is_head_of_department(user):
        raise PermissionDenied(
            'Entering and approving marks belongs to the examination officer.')

    touched = set(data)
    if touched & EXAM_OFFICE_FIELDS:
        raise PermissionDenied(
            'Only the examination officer can approve or publish results.')
    if touched & END_MARK_FIELDS:
        raise PermissionDenied(
            'End of semester examination marks are entered by the examination officer. '
            'Tutors enter continuous assessment — the CATs and assignments — only.')
    if not touched & CA_MARK_FIELDS:
        return

    semester = module.semester
    window = ResultEntryWindow.for_semester(semester, ResultEntryWindow.CA)
    if window is None:
        raise PermissionDenied(
            f'The examination officer has not opened continuous assessment marks '
            f'for {semester}.')
    if not window.is_open():
        when = (f'{window.opens_on:%d %b %Y} to {window.closes_on:%d %b %Y}'
                if window.status() == 'closed'
                else f'from {window.opens_on:%d %b %Y}')
        raise PermissionDenied(
            f'Entry for continuous assessment marks is closed for {semester} — the window '
            f'was {when}. Ask the examination officer to reopen it.')


def can_read_exams(user):
    """Seeing marks, eligibility, the exports and the performance analysis.

    Open to every admin role including the Head of Department: running a
    department means watching how its students are doing, by module and over
    time, and that is not the same power as changing a mark.
    """
    return bool(user and user.is_authenticated and user.is_staff)


def is_secretary(user):
    return bool(
        user and user.is_authenticated
        and SecretaryProfile.objects.filter(user=user, is_active=True).exists()
    )


def can_answer_requests(user):
    """Who may approve or decline what a student has asked the college for.

    The secretary's own job, and the admin's because somebody has to be able to
    do it when the secretary is away. Not the accountant and not a tutor: a
    request for leave of absence is neither a payment nor a mark.
    """
    return bool(user and user.is_authenticated and (user.is_staff or is_secretary(user)))


def active_semester():
    return Semester.objects.filter(is_active=True).select_related('academic_year').first()


def active_academic_year():
    sem = active_semester()
    if sem:
        return sem.academic_year
    return AcademicYear.objects.filter(is_active=True).first()


def attendance_is_effective(record):
    """Present counts; sick counts for eligibility only with a certificate."""
    return record.status == AttendanceRecord.PRESENT or (
        record.status == AttendanceRecord.SICK and record.certificate_submitted
    )


class IsAuthenticatedReadOnlyOrAdmin(BasePermission):
    def has_permission(self, request, view):
        return bool(
            request.user
            and request.user.is_authenticated
            and (request.method in SAFE_METHODS or request.user.is_staff)
        )


class IsFinanceUser(BasePermission):
    def has_permission(self, request, view):
        return can_manage_finance(request.user)


class IsExamOfficerOrFinance(BasePermission):
    """Exam declarations straddle the two roles.

    Declaring a student for a special, supplementary or repeat exam is an
    academic decision and belongs to the examination officer. The accountant
    only reads them, so they know what the student needs to be billed for.
    """

    def has_permission(self, request, view):
        user = request.user
        if not (user and user.is_authenticated):
            return False
        if request.method in SAFE_METHODS:
            return bool(can_read_exams(user) or can_manage_finance(user))
        return bool(can_manage_exams(user))


class IsEstateOfficer(BasePermission):
    def has_permission(self, request, view):
        return is_estate_officer(request.user)


class ReadExamsWriteExamOfficer(BasePermission):
    """Everyone who may already reach an examination endpoint keeps reading it;
    only the roles that decide examinations may change one.

    Tutors are unaffected — their own module scoping is done by the view.
    """
    message = 'Entering and approving marks belongs to the examination officer.'

    def has_permission(self, request, view):
        if not (request.user and request.user.is_authenticated):
            return False
        if request.method in SAFE_METHODS:
            return True
        return not is_head_of_department(request.user)


class CanReadExams(BasePermission):
    def has_permission(self, request, view):
        return can_read_exams(request.user)


class DeclaresExamWindows(BasePermission):
    """Any signed-in member of staff may read when the books are open — a tutor
    needs the dates before they type, not after. Only the examination officer
    says when."""
    message = 'Only the examination officer decides when marks may be entered.'

    def has_permission(self, request, view):
        if not (request.user and request.user.is_authenticated):
            return False
        return request.method in SAFE_METHODS or can_manage_exams(request.user)


class IsRequestOfficer(BasePermission):
    """Any signed-in member of staff may read the request queue; only the
    secretary or the admin may answer one."""

    def has_permission(self, request, view):
        if not (request.user and request.user.is_authenticated):
            return False
        return request.method in SAFE_METHODS or can_answer_requests(request.user)


def _make_both_semesters(year):
    """Ensure Semester 1 and 2 both exist for a given AcademicYear."""
    for num in (1, 2):
        Semester.objects.get_or_create(
            academic_year=year, number=num,
            defaults={'is_active': False},
        )


def _student_scope_for_request(request):
    qs = Student.objects.filter(module__in=user_modules(request.user))
    module_id = request.data.get('module_id') or request.query_params.get('module_id')
    class_level_id = request.data.get('class_level_id') or request.query_params.get('class_level_id')
    semester_id = request.data.get('semester_id') or request.query_params.get('semester_id')
    search = str(request.data.get('search') or request.query_params.get('search') or '').strip()
    if module_id:
        qs = qs.filter(module_id=module_id)
    if class_level_id:
        qs = qs.filter(module__class_level_id=class_level_id)
    if semester_id:
        qs = qs.filter(module__semester_id=semester_id)
    if search:
        qs = qs.filter(Q(nactvet_reg_no__icontains=search) | Q(name__icontains=search))
    return qs


# ── AUTH VIEWS ─────────────────────────────────────────────────────────────────

def login_view(request):
    if request.user.is_authenticated:
        return redirect('frontend')
    error = None
    identifier = ''

    if request.method == 'POST':
        identifier = str(request.POST.get('identifier', '')).strip()
        secret = str(request.POST.get('secret', '')).strip()

        # Try teacher/user authentication first
        user = authenticate(request, username=identifier, password=secret)
        if user is not None:
            login(request, user)
            next_url = request.GET.get('next', '')
            if next_url and url_has_allowed_host_and_scheme(
                next_url, allowed_hosts={request.get_host()}, require_https=request.is_secure()
            ):
                return redirect(next_url)
            return redirect('frontend')

        # Otherwise try student login using registration number + portal PIN.
        reg_no = identifier
        student = None
        students = Student.objects.filter(nactvet_reg_no__iexact=reg_no).select_related('module')
        for st in students:
            if st.check_portal_pin(secret):
                student = st
                break

        if student is None:
            error = 'Invalid credentials.'
        else:
            request.session['student_id'] = student.id
            request.session['student_reg_no'] = student.nactvet_reg_no
            return redirect('student-dashboard')

    return render(request, 'login.html', {'error': error, 'identifier': identifier})


def student_login_required(view_func):
    @wraps(view_func)
    def _wrapped(request, *args, **kwargs):
        if request.session.get('student_id'):
            return view_func(request, *args, **kwargs)
        return redirect('login')
    return _wrapped


def get_logged_student(request):
    student_id = request.session.get('student_id')
    if not student_id:
        return None
    return Student.objects.filter(
        id=student_id
    ).select_related(
        'module__class_level', 'module__semester__academic_year'
    ).first()


def student_logout_view(request):
    # Invalidate the complete server-side session and rotate the cookie. Merely
    # removing the student keys would leave the old session identifier reusable.
    request.session.flush()
    return redirect('login')


@api_view(['POST'])
def change_password(request):
    current_password = str(request.data.get('current_password', '')).strip()
    new_password = str(request.data.get('new_password', '')).strip()

    if request.user.is_authenticated:
        if len(new_password) < 6:
            return Response({'detail': 'New password must be at least 6 characters.'}, status=status.HTTP_400_BAD_REQUEST)
        if not request.user.check_password(current_password):
            return Response({'detail': 'Current password is incorrect.'}, status=status.HTTP_400_BAD_REQUEST)
        request.user.set_password(new_password)
        request.user.save(update_fields=['password'])
        update_session_auth_hash(request, request.user)
        return Response({'detail': 'Password updated.'})

    student = get_logged_student(request)
    if student is None:
        return Response({'detail': 'Authentication required.'}, status=status.HTTP_403_FORBIDDEN)
    if not student.check_portal_pin(current_password):
        return Response({'detail': 'Current PIN is incorrect.'}, status=status.HTTP_400_BAD_REQUEST)
    if not student_password_is_strong(new_password):
        return Response({'detail': STUDENT_PASSWORD_MESSAGE}, status=status.HTTP_400_BAD_REQUEST)
    if current_password == new_password:
        return Response({'detail': 'Choose a new password different from the generated PIN.'}, status=status.HTTP_400_BAD_REQUEST)

    updated = Student.objects.filter(nactvet_reg_no__iexact=student.nactvet_reg_no).count()
    for enrollment in Student.objects.filter(nactvet_reg_no__iexact=student.nactvet_reg_no):
        enrollment.set_portal_pin(new_password, require_change=False)
        enrollment.save(update_fields=['portal_pin_hash', 'must_change_portal_password'])
    return Response({'detail': 'Portal password updated.', 'updated': updated})


@xframe_options_sameorigin
def announcement_download(request, pk):
    """Stream an announcement PDF to any logged-in user — staff (Django auth)
    or student (session-only auth, same check as get_logged_student). Not a
    DRF view: it has to work for the session-only student portal, which never
    authenticates against request.user.

    The site sends X-Frame-Options: DENY, which is right for every page that
    can be acted on but stopped the dashboard embedding the notice it had just
    fetched — the browser reported it as "refused to connect". Relaxed to
    SAMEORIGIN on this one response: it is a read-only PDF with nothing to
    click, so framing it carries none of the risk the header exists to prevent,
    and every other view keeps DENY.
    """
    if not (request.user.is_authenticated or request.session.get('student_id')):
        return redirect('login')
    announcement = get_object_or_404(Announcement, pk=pk)
    try:
        handle = announcement.file.open('rb')
    except FileNotFoundError:
        raise Http404('That file is no longer available.')
    return FileResponse(handle, content_type='application/pdf', filename=announcement.file.name.rsplit('/', 1)[-1])


@student_login_required
def student_dashboard(request):
    student = get_logged_student(request)
    if not student:
        return redirect('login')
    if student.must_change_portal_password:
        return render(request, 'student_password_change.html', {
            'student_name': student.name,
            'registration_number': student.nactvet_reg_no,
        })

    # A form the college has made compulsory stops the portal here, the same
    # way an unchanged password does. One at a time, oldest first, so a student
    # facing three of them is not shown a wall of work.
    profile = finance.profile_for_student(student)
    pending = evaluations.pending_mandatory_forms(profile)
    if pending:
        form = pending[0]
        sections = form.sections.filter(for_office=False).prefetch_related('questions')
        # The submit script has to know how wide each table of text is to send
        # its rows back in the right shape.
        grid_widths = {
            question.id: len(question.columns)
            for section in sections for question in section.questions.all()
            if question.type == FormQuestion.GRID_TEXT
        }
        return render(request, 'student_required_form.html', {
            'student_name': student.name,
            'registration_number': student.nactvet_reg_no,
            'form': form,
            'sections': sections,
            'grid_widths': json.dumps({str(k): v for k, v in grid_widths.items()}),
            'remaining': len(pending) - 1,
        })

    active_sem = active_semester()
    enrollments = Student.objects.filter(nactvet_reg_no=student.nactvet_reg_no)
    if active_sem is not None:
        enrollments = enrollments.filter(
            module__semester__academic_year=active_sem.academic_year
        )
    enrollments = enrollments.select_related(
        'module__class_level', 'module__semester__academic_year', 'profile'
    ).prefetch_related(
        'attendance_records__session', 'result', 'module__sessions'
    ).order_by(
        'module__semester__number', 'module__name'
    )

    modules = []
    attendance_sum = 0
    attendance_count = 0
    for enrollment in enrollments:
        serializer = StudentSerializer(enrollment)
        data = serializer.data
        result = getattr(enrollment, 'result', None)
        result_data = StudentResultSerializer(result).data if result else None
        ca_approved = bool(result and result.ca_approved)
        final_approved = bool(result and result.final_approved)
        if result_data and not ca_approved:
            for field in (
                'assign1', 'assign2', 'cat1_theory', 'cat2_theory',
                'cat1_practical', 'cat2_practical',
                'assign1_w', 'assign2_w', 'cat1_theory_w', 'cat2_theory_w',
                'cat1_prac_w', 'cat2_prac_w',
                'theory_ca', 'practical_ca',
                'theory_eligible', 'practical_eligible', 'ca_eligible',
            ):
                result_data[field] = None
            if not final_approved:
                result_data['total_ca'] = None
        if result_data and not final_approved:
            for field in (
                'end_theory', 'end_practical',
                'end_theory_w', 'end_prac_w', 'end_exam_total', 'final_total',
                'supplementary_mark', 'end_exam_mark', 'supplementary_required',
                'result_status', 'grade', 'grade_point', 'grade_description',
            ):
                result_data[field] = None
        if result_data and ca_approved:
            ca_fields = [
                'assign1', 'assign2', 'cat1_theory', 'cat2_theory',
            ] + (
                ['cat1_practical', 'cat2_practical']
                if enrollment.module.has_practical else []
            )
            if any(result_data.get(f'{field}_absent') for field in ca_fields):
                result_data['ca_display'] = 'ABS'
            elif any(result_data.get(field) is None for field in ca_fields):
                result_data['ca_display'] = 'INC'
            else:
                result_data['ca_display'] = result_data.get('total_ca')
            for field in ('assign1', 'assign2', 'cat1_theory', 'cat2_theory',
                          'cat1_practical', 'cat2_practical'):
                if result_data.get(f'{field}_absent'):
                    result_data[field] = None
        if result_data and final_approved:
            for field in ('end_theory', 'end_practical'):
                if result_data.get(f'{field}_absent'):
                    result_data[field] = 'ABS'
        has_final_result = bool(
            final_approved
            and result
            and (bool(result.authority_grade)
                 or result.end_theory is not None or result.end_theory_absent
                 or result.end_practical is not None or result.end_practical_absent)
        )

        if data['sessions_total']:
            attendance_sum += data['attendance_pct']
            attendance_count += 1

        modules.append({
            'module_name': data['module_name'],
            'module_code': data['module_code'],
            'teacher': enrollment.module.teacher,
            'class_level': data['class_level_name'],
            'semester': data['semester_label'],
            'semester_number': enrollment.module.semester.number,
            'credits': enrollment.module.credits,
            'sessions_attended': data['sessions_attended'],
            'sessions_sick': data['sessions_sick'],
            'sessions_absent': data['sessions_absent'],
            'sessions_total': data['sessions_total'],
            'attendance_pct': data['attendance_pct'],
            'attendance_status': (
                'good' if data['attendance_pct'] >= 75
                else 'warning' if data['attendance_pct'] >= 50
                else 'critical'
            ) if data['sessions_total'] else 'pending',
            'result': result_data,
            'has_ca_result': bool(ca_approved and result),
            'has_final_result': has_final_result,
        })

    published_points = [
        (module['result']['grade_point'], module['credits'])
        for module in modules
        if module['has_final_result']
        and module['result']
        and module['result']['grade_point'] is not None
    ]
    gpa = (
        round(
            sum(float(points) * credits for points, credits in published_points)
            / sum(credits for _, credits in published_points),
            2,
        )
        if published_points else None
    )
    gpa_class = gpa_classification(
        gpa, student.module.class_level
    )

    overall_attendance = round(attendance_sum / attendance_count) if attendance_count else None
    total_present = sum(module['sessions_attended'] for module in modules)
    total_sick = sum(module['sessions_sick'] for module in modules)
    total_absent = sum(module['sessions_absent'] for module in modules)
    total_sessions = sum(module['sessions_total'] for module in modules)
    semester1_modules = [module for module in modules if module['semester_number'] == 1]
    semester2_modules = [module for module in modules if module['semester_number'] == 2]
    # CA is a running figure for the semester in progress, not a result. Once
    # the admin advances the semester the marks stop meaning anything to the
    # student — the end-of-semester result supersedes them — so the CA screen
    # follows the active semester and empties when the college moves on.
    # Everything published stays available under End of Semester.
    ca_modules = [
        module for module in modules
        if active_sem is not None
        and module['semester_number'] == active_sem.number
        and module['has_ca_result']
    ]
    ca_theory_modules = [
        module for module in ca_modules
        if not module['result'] or not module['result']['has_practical']
    ]
    ca_practical_modules = [
        module for module in ca_modules
        if module['result'] and module['result']['has_practical']
    ]
    active_modules = [
        module for module in modules
        if active_sem is None or module['semester_number'] == active_sem.number
    ]
    eligibility_rows = []
    clearance_cache = {}
    for enrollment in enrollments:
        # .all() (not .filter(...).count() x3) reuses the module__sessions and
        # attendance_records__session prefetches done above instead of issuing
        # 4 fresh queries per enrollment — was an N+1 that scaled with how many
        # modules a student takes.
        module_sessions = list(enrollment.module.sessions.all())
        counts = {
            Session.CAT1: sum(1 for s in module_sessions if s.exam_period == Session.CAT1),
            Session.CAT2: sum(1 for s in module_sessions if s.exam_period == Session.CAT2),
            'all': len(module_sessions),
        }
        records = list(enrollment.attendance_records.all())
        cat1_eff = sum(1 for r in records if r.session.exam_period == Session.CAT1 and attendance_is_effective(r))
        cat2_eff = sum(1 for r in records if r.session.exam_period == Session.CAT2 and attendance_is_effective(r))
        all_eff = sum(1 for r in records if attendance_is_effective(r))
        cat1_pct = round((cat1_eff / counts[Session.CAT1]) * 100) if counts[Session.CAT1] else None
        cat2_pct = round((cat2_eff / counts[Session.CAT2]) * 100) if counts[Session.CAT2] else None
        end_pct = round((all_eff / counts['all']) * 100) if counts['all'] else None
        cat1_att = (cat1_pct >= ELIGIBILITY_THRESHOLD) if counts[Session.CAT1] else None
        cat2_att = (cat2_pct >= ELIGIBILITY_THRESHOLD) if counts[Session.CAT2] else None
        end_att = (end_pct >= ELIGIBILITY_THRESHOLD) if counts['all'] else None
        total_ca, ca_ok, ca_note = ca_eligibility_for_student(enrollment)
        # Same computation the tutor's eligibility screen uses — one definition,
        # so the student and the college can never be shown different answers
        # about the same exam.
        cat1_fin = _clearance_for(enrollment, ChargeType.CAT1, clearance_cache)
        cat2_fin = _clearance_for(enrollment, ChargeType.CAT2, clearance_cache)
        end_fin = _clearance_for(enrollment, ChargeType.FINAL, clearance_cache)
        cat1_parts = [
            (cat1_att, attendance_eligibility_reason(cat1_pct, counts[Session.CAT1], 'CAT 1')),
            (cat1_fin['cleared'], cat1_fin['reason']),
        ]
        cat2_parts = [
            (cat2_att, attendance_eligibility_reason(cat2_pct, counts[Session.CAT2], 'CAT 2')),
            (cat2_fin['cleared'], cat2_fin['reason']),
        ]
        end_parts = [
            (end_att, attendance_eligibility_reason(end_pct, counts['all'], 'End-of-semester')),
            (ca_ok, ca_note),
            (end_fin['cleared'], end_fin['reason']),
        ]
        eligibility_rows.append({
            'module_code': enrollment.module.code,
            'module_name': enrollment.module.name,
            'semester': enrollment.module.semester.label,
            'cat1_pct': cat1_pct,
            'cat1_ok': combined_eligibility(cat1_parts),
            'cat1_note': combined_reason(cat1_parts),
            'cat2_pct': cat2_pct,
            'cat2_ok': combined_eligibility(cat2_parts),
            'cat2_note': combined_reason(cat2_parts),
            'ese_pct': end_pct,
            'ca_total': total_ca,
            'ese_ok': combined_eligibility(end_parts),
            'ese_note': combined_reason(end_parts),
        })
    ledger_profile = profile
    ledger_year = active_sem.academic_year if active_sem else student.module.semester.academic_year
    ledger_totals = finance.balance_for(ledger_profile, ledger_year)
    finance_charges = finance.with_balances(
        StudentCharge.objects
        .filter(profile=ledger_profile, academic_year=ledger_year)
        .select_related('charge_type', 'semester')
    ).order_by('due_date')
    finance_payments = (
        Payment.objects.filter(profile=ledger_profile)
        .prefetch_related('allocations__charge__charge_type')
        .order_by('-payment_date', '-created_at')
    )
    finance_invoices = (
        Invoice.objects.filter(profile=ledger_profile, academic_year=ledger_year)
        .select_related('bank_account', 'academic_year')
        .prefetch_related('lines__charge__charge_type')
    )
    # The template cannot call finance.invoice_paid(invoice), so the numbers a
    # student reads off the invoice table are worked out here.
    finance_invoice_rows = []
    for invoice in finance_invoices:
        paid = finance.invoice_paid(invoice)
        total = finance.money(invoice.total)
        finance_invoice_rows.append({
            'invoice': invoice,
            'total': total,
            'paid': paid,
            'outstanding': max(total - paid, Decimal('0.00')),
            'status': finance.invoice_status(invoice),
            'installments': len(invoice.lines.all()),   # prefetched
            'expires_on': invoice.expires_on,
        })
    # What the student can generate an invoice for: school fees, direct costs,
    # or one of the other payments they have been billed.
    invoice_payments = finance.invoiceable_payments(ledger_profile, ledger_year)
    # Instalment dates are no longer printed down the invoice — the invoice
    # stands all year. They surface here instead, as reminders the student sees
    # before each one falls due.
    fee_reminders = finance.installment_reminders(ledger_profile, ledger_year)
    fee_reminders_due = [r for r in fee_reminders
                         if r['urgency'] in (finance.OVERDUE, finance.DUE_SOON)]
    fee_reminders_later = [r for r in fee_reminders if r['urgency'] == finance.UPCOMING]
    finance_required = ledger_totals['billed']
    finance_paid = ledger_totals['paid']
    finance_balance = ledger_totals['balance']
    announcements = Announcement.objects.select_related('uploaded_by')[:20]
    # Evaluation forms the student may fill in. Rendered server-side so the
    # sidebar badge is right before any JavaScript runs.
    student_forms = evaluations.forms_for_student(ledger_profile, kind=Form.EVALUATION)
    forms_outstanding = sum(1 for entry in student_forms if entry['can_answer'])
    # Services the student can ask the college for, and what came of the ones
    # they have already asked for.
    service_forms = evaluations.services_for_student(ledger_profile)
    my_requests = evaluations.my_requests(ledger_profile)
    requests_pending = sum(1 for r in my_requests if r.status == FormResponse.PENDING)
    return render(request, 'student_dashboard.html', {
        'student_name': student.name,
        'registration_number': student.nactvet_reg_no,
        'announcements': announcements,
        'modules': modules,
        'module_count': len(modules),
        'active_module_count': len(active_modules),
        'semester1_modules': semester1_modules,
        'semester2_modules': semester2_modules,
        'ca_theory_modules': ca_theory_modules,
        'ca_practical_modules': ca_practical_modules,
        'has_ca_results': bool(ca_modules),
        'ca_count': len(ca_modules),
        'published_result_count': sum(1 for module in modules if module['has_final_result']),
        'has_final_results': any(module['has_final_result'] for module in modules),
        'service_forms': service_forms,
        'my_requests': my_requests,
        'requests_pending': requests_pending,
        'gpa': gpa,
        'gpa_classification': gpa_class,
        'overall_attendance': overall_attendance,
        'total_present': total_present,
        'total_sick': total_sick,
        'total_absent': total_absent,
        'total_sessions': total_sessions,
        'eligibility_rows': eligibility_rows,
        'finance_charges': finance_charges,
        'finance_payments': finance_payments,
        'finance_invoices': finance_invoices,
        'finance_invoice_rows': finance_invoice_rows,
        'invoice_payments': invoice_payments,
        'student_forms': student_forms,
        'forms_outstanding': forms_outstanding,
        'fee_reminders': fee_reminders,
        'fee_reminders_due': fee_reminders_due,
        'fee_reminders_later': fee_reminders_later,
        'academic_year_closes_on': ledger_year.closes_on if ledger_year else None,
        'finance_waived': ledger_totals['waived'],
        'finance_required': finance_required,
        'finance_paid': finance_paid,
        'finance_balance': finance_balance,
        'academic_year': (
            active_sem.academic_year.name
            if active_sem else student.module.semester.academic_year.name
        ),
        'active_semester': active_sem.label if active_sem else 'N/A',
    })


def logout_view(request):
    # Django logout flushes the complete session and rotates the session cookie.
    logout(request)
    return redirect('login')


def register_view(request):
    if not request.user.is_authenticated:
        return redirect('login')
    if not request.user.is_staff:
        return redirect('frontend')
    levels = ClassLevel.objects.prefetch_related('modules__semester__academic_year').all()
    form = TeacherRegistrationForm(request.POST or None)
    if request.method == 'POST' and form.is_valid():
        user = form.save(commit=False)
        full_name = form.cleaned_data['full_name'].strip()
        parts = full_name.split(None, 1)
        user.first_name = parts[0]
        user.last_name = parts[1] if len(parts) > 1 else ''
        user.save()
        TeacherProfile.objects.create(user=user, full_name=full_name)
        for mid in request.POST.getlist('modules'):
            try:
                Module.objects.get(id=mid).teachers.add(user)
            except Module.DoesNotExist:
                pass
        return redirect('frontend')
    return render(request, 'register.html', {'form': form, 'levels': levels})


#: What each role is, and the profile row that carries it. A user may hold
#: several: promoting a tutor to Principal is not a new account, it is the same
#: person keeping the modules they already teach and gaining an office.
ROLE_PROFILES = {
    'tutor': TeacherProfile,
    'accountant': AccountantProfile,
    'estate_officer': EstateOfficerProfile,
    'secretary': SecretaryProfile,
    'principal': PrincipalProfile,
    'hod': HeadOfDepartmentProfile,
}
#: The roles that carry Django's `is_staff`. "exam_officer" has no profile of
#: its own — it *is* is_staff with neither of the other two offices — so it is
#: named here to make it something the role editor can grant and take away
#: rather than an invisible side effect.
STAFF_ROLES = ('principal', 'hod', 'exam_officer')
#: Which title an account is listed under when it holds more than one.
ROLE_ORDER = ('principal', 'hod', 'exam_officer', 'secretary', 'accountant',
              'estate_officer', 'tutor')
ROLE_LABELS = {
    'tutor': 'Tutor', 'accountant': 'Accountant', 'estate_officer': 'Estate Officer',
    'secretary': 'Secretary', 'principal': 'Principal', 'hod': 'Head of Department',
    'exam_officer': 'Exam Officer',
}


def _held(model, user):
    """Whether this account carries this profile. TeacherProfile has no
    `is_active` switch; every other role profile does."""
    qs = model.objects.filter(user=user)
    if hasattr(model, 'is_active'):
        qs = qs.filter(is_active=True)
    return qs.exists()


def roles_for(user):
    """Every role this account holds, most senior first."""
    held = {name for name, model in ROLE_PROFILES.items() if _held(model, user)}
    if user.is_staff and not held & {'principal', 'hod'}:
        held.add('exam_officer')
    return [role for role in ROLE_ORDER if role in held]


def full_name_for(user):
    for role in ROLE_ORDER:
        model = ROLE_PROFILES.get(role)
        if model is None:
            continue
        profile = model.objects.filter(user=user).first()
        if profile:
            return profile.full_name
    return user.get_full_name() or user.username


@transaction.atomic
def set_roles(user, roles, *, full_name=None):
    """Make this account hold exactly these roles.

    Adding one never disturbs what the account already had — a tutor made Head
    of Department keeps every module they teach — and removing one takes away
    the office, not the history. Module assignments live on the module, not on
    the profile, so they survive either way.
    """
    wanted = set(roles)
    unknown = wanted - set(ROLE_PROFILES) - {'exam_officer'}
    if unknown:
        raise ValueError(f'Unknown role(s): {", ".join(sorted(unknown))}.')
    if not wanted:
        raise ValueError('An account must hold at least one role.')

    name = full_name or full_name_for(user)
    for role, model in ROLE_PROFILES.items():
        existing = model.objects.filter(user=user).first()
        if role in wanted:
            if existing is None:
                model.objects.create(user=user, full_name=name)
            elif hasattr(existing, 'is_active') and not existing.is_active:
                existing.is_active = True
                existing.save(update_fields=['is_active'])
        elif existing is not None:
            existing.delete()

    # A superuser keeps the keys to the building whatever else changes.
    staff = user.is_superuser or bool(wanted & set(STAFF_ROLES))
    if user.is_staff != staff:
        user.is_staff = staff
        user.save(update_fields=['is_staff'])
    return user


def staff_account_row(user):
    held = roles_for(user)
    return {
        'id': user.id,
        'username': user.username,
        'full_name': full_name_for(user),
        'role': held[0] if held else '',
        'roles': held,
        'role_labels': [ROLE_LABELS[role] for role in held],
        'is_staff': user.is_staff,
        'is_superuser': user.is_superuser,
        'module_ids': [m.id for m in user.modules_taught.all()] if 'tutor' in held else None,
    }


@api_view(['POST'])
@login_required
def set_module_scope(request):
    """Switch between the whole college and the modules you teach yourself.

    A view scope, not a role: a Principal working in "my modules" still holds
    every principal right. It exists because being promoted should not bury
    your own classes in a list of the college's.
    """
    scope = str(request.data.get('scope', '')).strip()
    if scope not in ('mine', 'college'):
        return Response({'detail': 'Scope must be "mine" or "college".'},
                        status=status.HTTP_400_BAD_REQUEST)
    if scope == 'mine' and not teaches_anything(request.user):
        return Response({'detail': 'You are not assigned to any module to narrow down to.'},
                        status=status.HTTP_400_BAD_REQUEST)

    request.session['teaching_scope_only'] = scope == 'mine'
    request.user.teaching_scope_only = scope == 'mine'
    return Response({'scope': scope, 'modules': user_modules(request.user).count()})


@api_view(['POST'])
@login_required
def set_staff_roles(request, user_id):
    """Move an account between roles, or give it a second one.

    The reason this exists rather than "make another account": a tutor promoted
    to Head of Department is the same person, and splitting them in two loses
    the modules they teach and gives the college two rows for one member of
    staff.
    """
    if not request.user.is_staff:
        return Response({'detail': 'Only the administrator can manage staff accounts.'},
                        status=status.HTTP_403_FORBIDDEN)

    user = get_object_or_404(User, pk=user_id)
    roles = request.data.get('roles')
    if not isinstance(roles, (list, tuple)):
        return Response({'detail': 'Send the roles as a list, e.g. ["tutor", "hod"].'},
                        status=status.HTTP_400_BAD_REQUEST)
    if user == request.user and not set(roles) & set(STAFF_ROLES):
        return Response(
            {'detail': 'That would take away your own administrator access. '
                       'Ask another administrator to make this change.'},
            status=status.HTTP_400_BAD_REQUEST)

    try:
        set_roles(user, roles, full_name=str(request.data.get('full_name', '')).strip() or None)
    except ValueError as exc:
        return Response({'detail': str(exc)}, status=status.HTTP_400_BAD_REQUEST)
    return Response(staff_account_row(user))


@api_view(['GET', 'POST'])
@login_required
def create_staff_account(request):
    if not request.user.is_staff:
        return Response({'detail': 'Only the administrator can manage staff accounts.'}, status=status.HTTP_403_FORBIDDEN)

    if request.method == 'GET':
        # An account with no profile at all used to be invisible here — which
        # meant the examination officer could not see their own row, let alone
        # give somebody else the job.
        users = User.objects.filter(
            Q(profile__isnull=False) | Q(accountant_profile__isnull=False)
            | Q(estate_officer_profile__isnull=False) | Q(secretary_profile__isnull=False)
            | Q(principal_profile__isnull=False) | Q(hod_profile__isnull=False)
            | Q(is_staff=True)
        ).distinct().select_related(
            'profile', 'accountant_profile', 'estate_officer_profile', 'secretary_profile',
            'principal_profile', 'hod_profile',
        ).prefetch_related('modules_taught').order_by('username')

        return Response([staff_account_row(user) for user in users])

    role = str(request.data.get('role', '')).strip()
    full_name = str(request.data.get('full_name', '')).strip()
    username = str(request.data.get('username', '')).strip()
    password = str(request.data.get('password', '')).strip()
    module_ids = request.data.get('module_ids') or []

    if role not in ('tutor', 'accountant', 'estate_officer', 'secretary', 'principal', 'hod'):
        return Response({'detail': 'Role must be tutor, accountant, estate officer, secretary, '
                                   'principal, or head of department.'},
                        status=status.HTTP_400_BAD_REQUEST)
    if not full_name or not username or len(password) < 6:
        return Response({'detail': 'Full name, username, and a 6+ character password are required.'}, status=status.HTTP_400_BAD_REQUEST)
    if User.objects.filter(username__iexact=username).exists():
        return Response({'detail': 'That username is already taken.'}, status=status.HTTP_400_BAD_REQUEST)

    with transaction.atomic():
        parts = full_name.split(None, 1)
        user = User.objects.create_user(
            username=username,
            password=password,
            first_name=parts[0],
            last_name=parts[1] if len(parts) > 1 else '',
        )
        # One place decides what a role means — including that the Principal and
        # the Head of Department carry admin rights. What each of them is *not*
        # allowed near (money, college property, and for the HoD examinations)
        # is decided by the checks those areas make, not by withholding is_staff.
        set_roles(user, [role], full_name=full_name)
        if role == 'tutor':
            for module in Module.objects.filter(id__in=module_ids):
                module.teachers.add(user)

    return Response(staff_account_row(user), status=status.HTTP_201_CREATED)


@api_view(['POST'])
@login_required
def set_staff_modules(request, user_id):
    """Replace an existing tutor's full module set in one call — a tutor may
    carry several modules across different class levels/semesters, so this
    is a sync (add + remove diff), not an append."""
    if not request.user.is_staff:
        return Response({'detail': 'Only the administrator can assign modules.'}, status=status.HTTP_403_FORBIDDEN)

    try:
        user = User.objects.get(pk=user_id)
    except User.DoesNotExist:
        return Response({'detail': 'Staff account not found.'}, status=status.HTTP_404_NOT_FOUND)
    if not hasattr(user, 'profile'):
        return Response({'detail': 'Only tutor accounts can be assigned modules.'}, status=status.HTTP_400_BAD_REQUEST)

    module_ids = request.data.get('module_ids') or []
    modules = Module.objects.filter(id__in=module_ids)
    user.modules_taught.set(modules)
    return Response({
        'detail': f'Updated modules for {user.profile.full_name}.',
        'module_ids': [m.id for m in modules],
    })


# ── ACADEMIC YEAR ──────────────────────────────────────────────────────────────

class AcademicYearViewSet(viewsets.ModelViewSet):
    queryset = AcademicYear.objects.all()
    serializer_class = AcademicYearSerializer
    permission_classes = [IsAuthenticatedReadOnlyOrAdmin]

    @action(detail=False, methods=['post'], url_path='advance')
    def advance(self, request):
        """
        Advance the active semester:
          Semester 1 → Semester 2 (same academic year)
          Semester 2 → Semester 1 of the next academic year
        Staff only.
        """
        if not request.user.is_staff:
            return Response({'detail': 'Staff only.'}, status=status.HTTP_403_FORBIDDEN)

        with transaction.atomic():
            cur = Semester.objects.select_for_update().filter(is_active=True).first()
            if not cur:
                return Response({'detail': 'No active semester found.'}, status=status.HTTP_400_BAD_REQUEST)

            cur.is_active = False
            cur.save()

            if cur.number == 1:
                new_sem, _ = Semester.objects.get_or_create(
                    academic_year=cur.academic_year, number=2,
                    defaults={'is_active': True}
                )
                new_sem.is_active = True
                new_sem.save()
                new_year = cur.academic_year
            else:
                cur.academic_year.is_active = False
                cur.academic_year.save()
                new_year, _ = AcademicYear.objects.get_or_create(
                    name=cur.academic_year.next_name,
                    defaults={'is_active': True}
                )
                new_year.is_active = True
                new_year.save()
                # Ensure both semesters exist for the new year
                _make_both_semesters(new_year)
                new_sem = Semester.objects.get(academic_year=new_year, number=1)
                new_sem.is_active = True
                new_sem.save()

        return Response({
            'detail': f'Advanced to {new_sem}',
            'year': new_year.name,
            'semester': new_sem.number,
            'label': new_sem.label,
        })

    @action(detail=False, methods=['get'], url_path='active')
    def active(self, request):
        sem = active_semester()
        if not sem:
            return Response(None)
        return Response(SemesterSerializer(sem).data)


class SemesterViewSet(viewsets.ReadOnlyModelViewSet):
    serializer_class = SemesterSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = Semester.objects.select_related('academic_year').all()
        if self.request.query_params.get('is_active') == 'true':
            qs = qs.filter(is_active=True)
        year_id = self.request.query_params.get('year_id')
        if year_id:
            qs = qs.filter(academic_year_id=year_id)
        return qs

    @action(detail=True, methods=['patch'], url_path='cutoffs')
    def update_cutoffs(self, request, pk=None):
        """Admin-only: set the attendance cutoff dates for CAT1, CAT2 and end of semester."""
        if not request.user.is_staff:
            raise PermissionDenied('Only the administrator can set cutoff dates.')
        semester = self.get_object()
        for field in ('cat1_cutoff', 'cat2_cutoff', 'end_cutoff'):
            if field in request.data:
                setattr(semester, field, request.data[field] or None)
        semester.save()
        return Response(SemesterSerializer(semester).data)


# ── CLASS LEVEL ────────────────────────────────────────────────────────────────

class ClassLevelViewSet(viewsets.ModelViewSet):
    queryset = ClassLevel.objects.all()
    serializer_class = ClassLevelSerializer
    permission_classes = [IsAuthenticatedReadOnlyOrAdmin]


# ── ANNOUNCEMENTS ──────────────────────────────────────────────────────────────

class AnnouncementViewSet(viewsets.ModelViewSet):
    queryset = Announcement.objects.select_related('uploaded_by').all()
    serializer_class = AnnouncementSerializer
    permission_classes = [IsAuthenticatedReadOnlyOrAdmin]

    def perform_create(self, serializer):
        serializer.save(uploaded_by=self.request.user)


# ── MODULES ────────────────────────────────────────────────────────────────────

class ModuleViewSet(viewsets.ModelViewSet):
    serializer_class = ModuleSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = user_modules(self.request.user).select_related('class_level', 'semester__academic_year')
        for param, field in [
            ('class_level_id', 'class_level_id'),
            ('semester_id', 'semester_id'),
        ]:
            val = self.request.query_params.get(param)
            if val:
                qs = qs.filter(**{field: val})
        return qs

    def perform_create(self, serializer):
        if not self.request.user.is_staff:
            raise PermissionDenied('Only the administrator can create modules.')
        module = serializer.save()
        module.teachers.add(self.request.user)

    def update(self, request, *args, **kwargs):
        if not request.user.is_staff:
            raise PermissionDenied('Only the administrator can edit modules.')
        return super().update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        if not request.user.is_staff:
            raise PermissionDenied('Only the administrator can delete modules.')
        m = self.get_object()
        name = m.name
        m.delete()
        return Response({'detail': f'Module "{name}" deleted.'}, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'], url_path='claim')
    def claim(self, request, pk=None):
        try:
            m = Module.objects.get(pk=pk)
        except Module.DoesNotExist:
            return Response({'detail': 'Not found.'}, status=status.HTTP_404_NOT_FOUND)
        # Self-service claiming is only for picking up a module nobody teaches
        # yet. Once a module has a tutor, moving/adding tutors is an admin
        # action — otherwise any teacher account could grant themselves access
        # to another module's students, attendance and CA marks at will.
        if not request.user.is_staff and m.teachers.exists() and not m.teachers.filter(pk=request.user.pk).exists():
            raise PermissionDenied('This module already has a tutor. Ask the administrator to add you.')
        m.teachers.add(request.user)
        return Response({'detail': f'Claimed: {m.name}'})

    @action(detail=True, methods=['post'], url_path='unclaim')
    def unclaim(self, request, pk=None):
        try:
            m = Module.objects.get(pk=pk)
        except Module.DoesNotExist:
            return Response({'detail': 'Not found.'}, status=status.HTTP_404_NOT_FOUND)
        m.teachers.remove(request.user)
        return Response({'detail': f'Unclaimed: {m.name}'})


# ── STUDENTS ───────────────────────────────────────────────────────────────────

class StudentViewSet(viewsets.ModelViewSet):
    serializer_class = StudentSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = Student.objects.filter(
            module__in=user_modules(self.request.user)
        ).select_related('module__class_level', 'module__semester__academic_year').prefetch_related(
            'attendance_records', 'module__sessions'
        )

        for param, field in [
            ('module_id', 'module_id'),
            ('class_level_id', 'module__class_level_id'),
            ('semester_id', 'module__semester_id'),
        ]:
            val = self.request.query_params.get(param)
            if val:
                qs = qs.filter(**{field: val})
        search = self.request.query_params.get('search', '').strip()
        if search:
            qs = qs.filter(Q(nactvet_reg_no__icontains=search) | Q(name__icontains=search))
        return qs

    def perform_create(self, serializer):
        module = serializer.validated_data.get('module')
        if not self.request.user.is_staff:
            allowed_module_ids = user_modules(self.request.user).values_list('id', flat=True)
            if module is None or module.id not in allowed_module_ids:
                raise PermissionDenied('Only tutors for this module can add students.')
        serializer.save()

    def update(self, request, *args, **kwargs):
        if not request.user.is_staff:
            obj = self.get_object()
            allowed_module_ids = set(user_modules(self.request.user).values_list('id', flat=True))
            if obj.module_id not in allowed_module_ids:
                raise PermissionDenied('Only tutors for this module can edit students.')
            target_module = request.data.get('module')
            if target_module:
                try:
                    target_module_id = int(target_module)
                except (TypeError, ValueError):
                    target_module_id = None
                if target_module_id not in allowed_module_ids:
                    raise PermissionDenied('You may only move students to modules you teach.')
        return super().update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        obj = self.get_object()
        if not request.user.is_staff:
            allowed_module_ids = user_modules(self.request.user).values_list('id', flat=True)
            if obj.module_id not in allowed_module_ids:
                raise PermissionDenied('Only tutors for this module can remove students.')
        return super().destroy(request, *args, **kwargs)

    @action(detail=False, methods=['post'], url_path='bulk_create')
    def bulk_create(self, request):
        serializer = BulkStudentSerializer(data=request.data, context={'user': request.user})
        serializer.is_valid(raise_exception=True)
        result = serializer.save()
        return Response(result, status=status.HTTP_201_CREATED)

    @action(detail=False, methods=['post'], url_path='bulk_set_pin')
    def bulk_set_pin(self, request):
        if not request.user.is_staff:
            raise PermissionDenied('Only the administrator can set student passwords in bulk.')
        portal_pin = str(request.data.get('portal_pin', '')).strip()
        if len(portal_pin) < 6:
            return Response({'detail': 'Password/PIN must be at least 6 characters.'}, status=status.HTTP_400_BAD_REQUEST)

        qs = _student_scope_for_request(request)
        updated = 0
        with transaction.atomic():
            for student in qs.select_for_update():
                student.set_portal_pin(portal_pin)
                student.save(update_fields=['portal_pin_hash', 'must_change_portal_password'])
                updated += 1
        return Response({'updated': updated})


# ── SESSIONS ───────────────────────────────────────────────────────────────────

class SessionViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = Session.objects.filter(
            module__in=user_modules(self.request.user)
        ).select_related('module__class_level', 'module__semester__academic_year').prefetch_related('records__student')

        for param, field in [
            ('module_id', 'module_id'),
            ('class_level_id', 'module__class_level_id'),
            ('semester_id', 'module__semester_id'),
        ]:
            val = self.request.query_params.get(param)
            if val:
                qs = qs.filter(**{field: val})
        return qs

    def get_serializer_class(self):
        if self.action in ('create', 'update', 'partial_update'):
            return SessionCreateSerializer
        return SessionSerializer

    @staticmethod
    def _check_attendance_cutoff(module, period):
        """Reject session writes whose exam period is past the semester cutoff date."""
        semester = module.semester
        today = timezone.localdate()

        CUTOFF_MAP = {
            Session.CAT1:    ('cat1_cutoff', 'CAT 1'),
            Session.CAT2:    ('cat2_cutoff', 'CAT 2'),
            Session.GENERAL: ('end_cutoff',  'End of Semester'),
        }
        cutoff_field, period_label = CUTOFF_MAP.get(period, ('end_cutoff', 'this period'))
        cutoff = getattr(semester, cutoff_field, None)

        if cutoff and today > cutoff:
            from rest_framework.exceptions import ValidationError
            raise ValidationError(
                f'Attendance is closed for {period_label}. '
                f'The cutoff date was {cutoff.strftime("%d %b %Y")}.'
            )

    def perform_create(self, serializer):
        module = serializer.validated_data['module']
        allowed_module_ids = user_modules(self.request.user).values_list('id', flat=True)
        if module.id not in allowed_module_ids:
            raise PermissionDenied('You may only record attendance for modules you teach.')
        period = serializer.validated_data.get('exam_period', Session.GENERAL)
        self._check_attendance_cutoff(module, period)
        serializer.save()

    def perform_update(self, serializer):
        module = serializer.validated_data.get('module', serializer.instance.module)
        allowed_module_ids = user_modules(self.request.user).values_list('id', flat=True)
        if module.id not in allowed_module_ids:
            raise PermissionDenied('You may only edit sessions for modules you teach.')
        period = serializer.validated_data.get('exam_period', serializer.instance.exam_period)
        self._check_attendance_cutoff(module, period)
        serializer.save()


# ── DASHBOARD ──────────────────────────────────────────────────────────────────

@api_view(['GET'])
@login_required
def dashboard(request):
    estate_user = is_estate_officer(request.user)
    if estate_user:
        return Response({
            'modules': 0, 'students': 0, 'sessions_today': 0, 'avg_attendance': None,
            'active_semester': None, 'recent_sessions': [], 'levels': [],
            'is_staff': False, 'is_accountant': False, 'is_estate_officer': True,
            'is_secretary': False, 'is_principal': False,
            'is_head_of_department': False, 'can_manage_exams': False,
            'teaches': False, 'module_scope': 'college',
        })
    today = timezone.localdate()
    my_modules = user_modules(request.user)
    sem = active_semester()

    subjects_count = my_modules.count()
    # Count distinct students per class level (by registration number) and sum
    # those per-class totals for an overall student count. This prevents
    # double-counting when modules in the same class have different student lists.
    students_qs = Student.objects.filter(module__in=my_modules)
    per_class = (
        students_qs
        .values('module__class_level_id')
        .annotate(student_count=Count('nactvet_reg_no', distinct=True))
    )
    students_count = int(per_class.aggregate(total=Sum('student_count'))['total'] or 0)
    sessions_today = Session.objects.filter(module__in=my_modules, date=today).count()

    all_students = (
        Student.objects.filter(module__in=my_modules)
        .prefetch_related('attendance_records').select_related('module')
    )

    # Cache session counts per module to avoid N+1
    _session_cache = {}

    def _held(module_id):
        if module_id not in _session_cache:
            _session_cache[module_id] = Session.objects.filter(module_id=module_id).count()
        return _session_cache[module_id]

    total_pct, count = 0, 0
    for st in all_students:
        held = _held(st.module_id)
        if held:
            effective = st.attendance_records.filter(status__in=['P', 'S']).count()
            total_pct += round((effective / held) * 100)
            count += 1
    avg_attendance = round(total_pct / count) if count else None

    recent = []
    for sess in (
        Session.objects.filter(module__in=my_modules)
        .select_related('module__class_level', 'module__semester__academic_year')
        .prefetch_related('records').order_by('-date', '-created_at')[:8]
    ):
        p = sess.records.filter(status='P').count()
        s = sess.records.filter(status='S').count()
        a = sess.records.filter(status='A').count()
        t = p + s + a
        recent.append({
            'id': sess.id,
            'module': sess.module.name,
            'session_type': sess.session_type,
            'session_type_display': sess.get_session_type_display(),
            'class_level': sess.module.class_level.name,
            'semester': sess.module.semester.label,
            'date': str(sess.date),
            'label': sess.label,
            'topic': sess.topic,
            'present': p, 'sick': s, 'absent': a,
            'pct': round(((p + s) / t) * 100) if t else 0,
        })

    levels = []
    for lvl in ClassLevel.objects.all():
        lvl_mods = my_modules.filter(class_level=lvl)
        if not lvl_mods.exists():
            continue
        # Distinct students in this class level by registration number
        lvl_students_count = (
            Student.objects.filter(module__in=lvl_mods)
            .values('nactvet_reg_no')
            .distinct()
            .count()
        )
        lp, lc = 0, 0
        # Use students per level for attendance pct calculation
        lvl_students = all_students.filter(module__in=lvl_mods)
        for st in lvl_students:
            held = _held(st.module_id)
            if held:
                effective = st.attendance_records.filter(status__in=['P', 'S']).count()
                lp += round((effective / held) * 100)
                lc += 1
        levels.append({
            'id': lvl.id,
            'name': lvl.name,
            'modules': lvl_mods.count(),
            'students': lvl_students_count,
            'avg_pct': round(lp / lc) if lc else None,
        })

    return Response({
        'modules': subjects_count,
        'students': students_count,
        'sessions_today': sessions_today,
        'avg_attendance': avg_attendance,
        'active_semester': SemesterSerializer(sem).data if sem else None,
        'recent_sessions': recent,
        'levels': levels,
        'is_staff': request.user.is_staff,
        'is_accountant': is_accountant(request.user),
        'is_estate_officer': False,
        'is_secretary': is_secretary(request.user),
        'is_principal': is_principal(request.user),
        'is_head_of_department': is_head_of_department(request.user),
        'can_manage_exams': can_manage_exams(request.user),
        # Whether this account has classes of its own, and which hat it is
        # currently wearing. Only somebody who both teaches and holds an admin
        # role is ever offered the switch.
        'teaches': teaches_anything(request.user),
        'module_scope': ('mine' if getattr(request.user, 'teaching_scope_only', False)
                         else 'college'),
    })


class ResultEntryWindowViewSet(viewsets.ModelViewSet):
    """When the books are open for marks.

    Any signed-in member of staff may read them — a tutor needs to know when
    they can enter, and being told after they have typed a screenful is no use.
    Only the examination officer declares one.
    """
    serializer_class = ResultEntryWindowSerializer
    permission_classes = [DeclaresExamWindows]

    def get_queryset(self):
        qs = (ResultEntryWindow.objects
              .select_related('semester__academic_year', 'declared_by'))
        semester_id = self.request.query_params.get('semester_id')
        if semester_id:
            qs = qs.filter(semester_id=semester_id)
        return qs

    def perform_create(self, serializer):
        serializer.save(declared_by=self.request.user)


# ── PERFORMANCE ───────────────────────────────────────────────────────────────

@api_view(['GET'])
@login_required
def performance(request):
    """How students are doing, by module and over time.

    The Principal, the Head of Department and the examination officer see the
    whole college; a tutor sees their own modules and nobody else's. Reading a
    summary is not the same power as changing a mark, so everyone who teaches
    gets one.
    """
    # None means "no restriction". A tutor with no modules assigned gets an
    # empty queryset, which is not the same thing and must show nothing.
    modules = (None if can_read_exams(request.user)
               and not getattr(request.user, 'teaching_scope_only', False)
               else user_modules(request.user))

    def pick(model, param):
        value = request.query_params.get(param)
        found = model.objects.filter(id=value).first() if value else None
        # A tutor asking for somebody else's module is answered as though they
        # had not filtered at all, rather than being shown it.
        if found is not None and model is Module and modules is not None:
            if not modules.filter(id=found.id).exists():
                return None
        return found

    try:
        data = analytics.summary(
            academic_year=pick(AcademicYear, 'academic_year_id'),
            semester=pick(Semester, 'semester_id'),
            class_level=pick(ClassLevel, 'class_level_id'),
            module=pick(Module, 'module_id'),
            assessment=request.query_params.get('assessment', analytics.FINAL),
            modules=modules,
        )
    except ValueError as exc:
        return Response({'detail': str(exc)}, status=status.HTTP_400_BAD_REQUEST)
    return Response({
        **data,
        'scope': 'college' if modules is None else 'my-modules',
        'options': analytics.filter_options(modules),
    })


# ── INVENTORY ─────────────────────────────────────────────────────────────────

class InventoryLocationViewSet(viewsets.ModelViewSet):
    queryset = InventoryLocation.objects.all()
    serializer_class = InventoryLocationSerializer
    permission_classes = [IsEstateOfficer]


class AssetCategoryViewSet(viewsets.ModelViewSet):
    queryset = AssetCategory.objects.all()
    serializer_class = AssetCategorySerializer
    permission_classes = [IsEstateOfficer]


class InventoryItemTypeViewSet(viewsets.ModelViewSet):
    queryset = InventoryItemType.objects.select_related('category')
    serializer_class = InventoryItemTypeSerializer
    permission_classes = [IsEstateOfficer]


class AssetViewSet(viewsets.ModelViewSet):
    serializer_class = AssetSerializer
    permission_classes = [IsEstateOfficer]

    def get_queryset(self):
        qs = Asset.objects.select_related('category', 'location')
        search = self.request.query_params.get('search', '').strip()
        if search:
            qs = qs.filter(Q(asset_tag__icontains=search) | Q(name__icontains=search) |
                           Q(description__icontains=search) | Q(responsible_office__icontains=search))
        for field in ('category', 'location', 'condition'):
            value = self.request.query_params.get(field)
            if value:
                qs = qs.filter(**{field: value})
        return qs

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user, updated_by=self.request.user)

    def perform_update(self, serializer):
        serializer.save(updated_by=self.request.user)


class AssetTransferViewSet(viewsets.ModelViewSet):
    queryset = AssetTransfer.objects.select_related('asset', 'from_location', 'to_location', 'resulting_asset')
    serializer_class = AssetTransferSerializer
    permission_classes = [IsEstateOfficer]

    def perform_create(self, serializer):
        asset = serializer.validated_data['asset']
        to_location = serializer.validated_data['to_location']
        if asset.location_id == to_location.id:
            raise PermissionDenied('Choose a different destination location.')
        with transaction.atomic():
            quantity = serializer.validated_data.get('quantity', 1)
            responsible = serializer.validated_data['new_responsible_office']
            from_location = asset.location
            if quantity < asset.quantity:
                suffix = 1
                while Asset.objects.filter(asset_tag=f'{asset.asset_tag}/{suffix:03d}').exists():
                    suffix += 1
                moved_asset = Asset.objects.create(
                    asset_tag=f'{asset.asset_tag}/{suffix:03d}', name=asset.name,
                    description=asset.description, category=asset.category, location=to_location,
                    responsible_office=responsible, quantity=quantity, condition=asset.condition,
                    created_by=self.request.user, updated_by=self.request.user,
                )
                asset.quantity -= quantity
                asset.updated_by = self.request.user
                asset.save(update_fields=['quantity', 'updated_by', 'updated_at'])
            else:
                moved_asset = asset
                asset.location = to_location
                asset.responsible_office = responsible
                asset.updated_by = self.request.user
                asset.save(update_fields=['location', 'responsible_office', 'updated_by', 'updated_at'])
            serializer.save(from_location=from_location, resulting_asset=moved_asset, recorded_by=self.request.user)


class AssetMaintenanceViewSet(viewsets.ModelViewSet):
    queryset = AssetMaintenance.objects.select_related('asset')
    serializer_class = AssetMaintenanceSerializer
    permission_classes = [IsEstateOfficer]

    def perform_create(self, serializer):
        asset = serializer.validated_data['asset']
        quantity = serializer.validated_data.get('quantity', 1)
        with transaction.atomic():
            if quantity < asset.quantity:
                suffix = 1
                while Asset.objects.filter(asset_tag=f'{asset.asset_tag}/{suffix:03d}').exists():
                    suffix += 1
                maintenance_asset = Asset.objects.create(
                    asset_tag=f'{asset.asset_tag}/{suffix:03d}', name=asset.name,
                    description=asset.description, category=asset.category, location=asset.location,
                    responsible_office=asset.responsible_office, quantity=quantity,
                    condition=asset.condition, created_by=self.request.user, updated_by=self.request.user,
                )
                asset.quantity -= quantity
                asset.updated_by = self.request.user
                asset.save(update_fields=['quantity', 'updated_by', 'updated_at'])
                serializer.save(asset=maintenance_asset, recorded_by=self.request.user)
            else:
                serializer.save(recorded_by=self.request.user)


class InventoryInspectionViewSet(viewsets.ModelViewSet):
    queryset = InventoryInspection.objects.select_related('location').prefetch_related('items')
    serializer_class = InventoryInspectionSerializer
    permission_classes = [IsEstateOfficer]

    def perform_create(self, serializer):
        serializer.save(recorded_by=self.request.user)


class InventoryInspectionItemViewSet(viewsets.ModelViewSet):
    queryset = InventoryInspectionItem.objects.select_related('inspection', 'asset')
    serializer_class = InventoryInspectionItemSerializer
    permission_classes = [IsEstateOfficer]


class AssetDisposalViewSet(viewsets.ModelViewSet):
    queryset = AssetDisposal.objects.select_related('asset')
    serializer_class = AssetDisposalSerializer
    permission_classes = [IsEstateOfficer]

    def perform_create(self, serializer):
        serializer.save(recorded_by=self.request.user)


INVENTORY_HEADERS = [
    'Office/Location *', 'Responsible Person/Office *', 'Item Type *',
    'Tag Prefix *', 'Starting Number *', 'Quantity *', 'Condition *', 'Description',
]


@api_view(['GET'])
def inventory_template(request):
    if not is_estate_officer(request.user):
        return Response({'detail': 'Estate Officer access required.'}, status=403)
    wb = Workbook()
    ws = wb.active
    ws.title = 'Assets'
    ws.append(INVENTORY_HEADERS)
    for cell in ws[1]:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = PatternFill('solid', fgColor='C0392B')
    ws.freeze_panes = 'A2'
    widths = [30, 32, 32, 22, 18, 14, 20, 38]
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + idx)].width = width
    lists = wb.create_sheet('Lists')
    categories = list(AssetCategory.objects.filter(is_active=True).values_list('name', flat=True))
    locations = list(InventoryLocation.objects.filter(is_active=True).values_list('name', flat=True))
    item_types = [f'{item.name} — {item.category.name}' for item in InventoryItemType.objects.filter(is_active=True).select_related('category')]
    conditions = [label for _, label in Asset.CONDITION_CHOICES]
    for row, value in enumerate(categories, 1): lists.cell(row, 1, value)
    for row, value in enumerate(locations, 1): lists.cell(row, 2, value)
    for row, value in enumerate(conditions, 1): lists.cell(row, 3, value)
    for row, value in enumerate(item_types, 1): lists.cell(row, 4, value)
    if locations:
        dv = DataValidation(type='list', formula1=f"'Lists'!$B$1:$B${len(locations)}")
        ws.add_data_validation(dv); dv.add('A2:A5000')
    if item_types:
        dv = DataValidation(type='list', formula1=f"'Lists'!$D$1:$D${len(item_types)}")
        ws.add_data_validation(dv); dv.add('C2:C5000')
    dv = DataValidation(type='list', formula1=f"'Lists'!$C$1:$C${len(conditions)}")
    ws.add_data_validation(dv); dv.add('G2:G5000')
    instructions = wb.create_sheet('Instructions', 0)
    instructions.append(['COLLEGE INVENTORY IMPORT TEMPLATE — Version 1'])
    instructions.append(['Every column marked * is required. Do not rename or reorder headings.'])
    instructions.append(['Choose one office and item type per row. Quantity is expanded into separate tags.'])
    instructions.append(['Example: prefix BPCH/CH, start 1, quantity 30 creates BPCH/CH/1 through BPCH/CH/30.'])
    instructions.append(['Validate the file in the system before confirming import.'])
    output = BytesIO(); wb.save(output)
    response = HttpResponse(output.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="college_inventory_template.xlsx"'
    return response


def _parse_inventory_upload(upload):
    try:
        wb = load_workbook(upload, read_only=True, data_only=True)
    except Exception:
        return [], [{'row': 0, 'field': 'File', 'message': 'Upload a valid .xlsx workbook.'}]
    if 'Assets' not in wb.sheetnames:
        return [], [{'row': 0, 'field': 'Worksheet', 'message': 'The Assets worksheet is missing.'}]
    ws = wb['Assets']
    headers = [str(c.value or '').strip() for c in ws[1]]
    if headers[:len(INVENTORY_HEADERS)] != INVENTORY_HEADERS:
        return [], [{'row': 1, 'field': 'Headings', 'message': 'Use the latest system template without changing headings.'}]
    locations = {x.name.casefold(): x for x in InventoryLocation.objects.filter(is_active=True)}
    item_types = {f'{x.name} — {x.category.name}'.casefold(): x for x in InventoryItemType.objects.filter(is_active=True).select_related('category')}
    conditions = {label.casefold(): value for value, label in Asset.CONDITION_CHOICES}
    existing = {x.casefold() for x in Asset.objects.values_list('asset_tag', flat=True)}
    seen, rows, errors = set(), [], []
    fields = ['location', 'responsible_office', 'item_type', 'tag_prefix', 'start_number', 'quantity', 'condition', 'description']
    for number, values in enumerate(ws.iter_rows(min_row=2, max_col=8, values_only=True), 2):
        if not any(v not in (None, '') for v in values): continue
        data = dict(zip(fields, values))
        for key, value in data.items():
            if key == 'description':
                continue
            if value is None or str(value).strip() == '': errors.append({'row': number, 'field': key, 'message': 'Required.'})
        location = locations.get(str(data['location'] or '').strip().casefold())
        item_type = item_types.get(str(data['item_type'] or '').strip().casefold())
        condition = conditions.get(str(data['condition'] or '').strip().casefold())
        if data['location'] and not location: errors.append({'row': number, 'field': 'location', 'message': 'Choose a location from the template list.'})
        if data['item_type'] and not item_type: errors.append({'row': number, 'field': 'item_type', 'message': 'Choose an item type from the template list.'})
        if data['condition'] and not condition: errors.append({'row': number, 'field': 'condition', 'message': 'Choose a condition from the template list.'})
        try:
            quantity = int(data['quantity'])
            start = int(data['start_number'])
            if quantity < 1 or quantity > 500 or start < 1 or float(data['quantity']) != quantity or float(data['start_number']) != start: raise ValueError
        except (TypeError, ValueError):
            quantity, start = 0, 0; errors.append({'row': number, 'field': 'quantity', 'message': 'Use quantity 1–500 and a positive starting number.'})
        prefix = str(data['tag_prefix'] or '').strip().rstrip('/')
        if item_type and location and condition and quantity:
            for sequence in range(start, start + quantity):
                tag = f'{prefix}/{sequence}'
                if tag.casefold() in existing or tag.casefold() in seen:
                    errors.append({'row': number, 'field': 'tag_prefix', 'message': f'Generated tag {tag} already exists or is duplicated.'})
                seen.add(tag.casefold())
                rows.append({'asset_tag': tag, 'name': item_type.name,
                             'description': str(data['description'] or item_type.description).strip(),
                             'category': item_type.category, 'item_type': item_type, 'location': location,
                             'responsible_office': str(data['responsible_office'] or '').strip(),
                             'quantity': 1, 'condition': condition})
    if not rows: errors.append({'row': 0, 'field': 'File', 'message': 'No item rows were found.'})
    return rows, errors


@api_view(['POST'])
def inventory_import(request):
    if not is_estate_officer(request.user):
        return Response({'detail': 'Estate Officer access required.'}, status=403)
    upload = request.FILES.get('file')
    if not upload:
        return Response({'detail': 'Choose an Excel file.'}, status=400)
    rows, errors = _parse_inventory_upload(upload)
    if errors:
        return Response({'valid': False, 'row_count': len(rows), 'errors': errors}, status=400)
    if str(request.data.get('confirm', '')).lower() not in ('1', 'true', 'yes'):
        return Response({'valid': True, 'row_count': len(rows), 'errors': []})
    with transaction.atomic():
        for row in rows:
            Asset.objects.create(**row, created_by=request.user, updated_by=request.user)
        AssetImport.objects.create(uploaded_by=request.user, file_name=upload.name, imported_rows=len(rows))
    return Response({'valid': True, 'imported': len(rows)}, status=201)


@api_view(['POST'])
def inventory_bulk_create(request):
    """Create individually tagged assets from a shared tag prefix and quantity."""
    if not is_estate_officer(request.user):
        return Response({'detail': 'Estate Officer access required.'}, status=403)

    try:
        count = int(request.data.get('count'))
        start_number = int(request.data.get('start_number', 1))
    except (TypeError, ValueError):
        return Response({'detail': 'Quantity and starting number must be positive whole numbers.'}, status=400)
    if count < 1 or count > 500 or start_number < 1:
        return Response({'detail': 'Enter a quantity from 1 to 500 and a starting number of at least 1.'}, status=400)

    prefix = str(request.data.get('asset_tag_prefix', '')).strip().rstrip('/')
    if not prefix:
        return Response({'asset_tag_prefix': ['Enter an asset tag prefix.']}, status=400)
    tags = [f'{prefix}/{number}' for number in range(start_number, start_number + count)]
    conflicts = list(Asset.objects.filter(asset_tag__in=tags).values_list('asset_tag', flat=True))
    if conflicts:
        return Response({
            'detail': 'One or more generated asset tags are already registered.',
            'conflicting_tags': conflicts[:20],
        }, status=400)

    shared = {
        'name': request.data.get('name'),
        'description': request.data.get('description', ''),
        'category': request.data.get('category'),
        'location': request.data.get('location'),
        'responsible_office': request.data.get('responsible_office'),
        'condition': request.data.get('condition'),
        'quantity': 1,
    }
    validated = []
    for tag in tags:
        serializer = AssetSerializer(data={**shared, 'asset_tag': tag})
        serializer.is_valid(raise_exception=True)
        validated.append(serializer.validated_data)

    with transaction.atomic():
        # Lock matching records so concurrent requests cannot both pass validation.
        if Asset.objects.select_for_update().filter(asset_tag__in=tags).exists():
            return Response({'detail': 'A generated tag was registered by another request. Please try again.'}, status=409)
        Asset.objects.bulk_create([
            Asset(**data, created_by=request.user, updated_by=request.user)
            for data in validated
        ])
    return Response({
        'created': count,
        'first_tag': tags[0],
        'last_tag': tags[-1],
    }, status=201)


@api_view(['POST'])
def inventory_office_register(request):
    """Register multiple item types for one office and expand each quantity into tagged assets."""
    if not is_estate_officer(request.user):
        return Response({'detail': 'Estate Officer access required.'}, status=403)
    location_id = request.data.get('location')
    responsible = str(request.data.get('responsible_office', '')).strip()
    rows = request.data.get('items')
    if not InventoryLocation.objects.filter(pk=location_id, is_active=True).exists():
        return Response({'location': ['Choose an active office/location.']}, status=400)
    if not responsible:
        return Response({'responsible_office': ['Enter the responsible person or office.']}, status=400)
    if not isinstance(rows, list) or not rows:
        return Response({'items': ['Add at least one item type.']}, status=400)

    generated, prepared, errors = [], [], []
    # Numbering is global per prefix, not per office. Work inside one transaction and
    # lock the asset rows so two simultaneous office submissions cannot reuse a tag.
    with transaction.atomic():
        existing_tags = list(Asset.objects.select_for_update().values_list('asset_tag', flat=True))
        next_numbers = {}

        def next_number_for(prefix):
            key = prefix.casefold()
            if key not in next_numbers:
                pattern = re.compile(rf'^{re.escape(prefix)}/(\d+)$', re.IGNORECASE)
                numbers = [int(match.group(1)) for tag in existing_tags if (match := pattern.match(tag))]
                next_numbers[key] = max(numbers, default=0) + 1
            number = next_numbers[key]
            next_numbers[key] += 1
            return number

        for index, row in enumerate(rows, 1):
            try:
                item_type = InventoryItemType.objects.select_related('category').get(pk=row.get('item_type'), is_active=True)
                count = int(row.get('quantity'))
            except InventoryItemType.DoesNotExist:
                errors.append({'row': index, 'field': 'item_type', 'message': 'Choose an active item type.'})
                continue
            except (TypeError, ValueError):
                errors.append({'row': index, 'field': 'quantity', 'message': 'Quantity must be a whole number.'})
                continue
            if count < 1 or count > 500:
                errors.append({'row': index, 'field': 'quantity', 'message': 'Use quantity 1–500.'})
                continue
            prefix = str(row.get('tag_prefix') or item_type.default_tag_prefix).strip().rstrip('/')
            condition = row.get('condition')
            tags = [f'{prefix}/{next_number_for(prefix)}' for _ in range(count)]
            generated.extend(tags)
            for tag in tags:
                prepared.append({
                    'asset_tag': tag, 'name': item_type.name, 'description': item_type.description,
                    'category': item_type.category_id, 'item_type': item_type.id,
                    'location': location_id, 'responsible_office': responsible,
                    'quantity': 1, 'condition': condition,
                })
        if len(generated) > 1000:
            errors.append({'row': 0, 'field': 'items', 'message': 'One office submission may create at most 1,000 asset records.'})
        if errors:
            return Response({'detail': 'Correct the office stock rows.', 'errors': errors}, status=400)

        serializers = [AssetSerializer(data=data) for data in prepared]
        for serializer in serializers:
            serializer.is_valid(raise_exception=True)
        if Asset.objects.filter(asset_tag__in=generated).exists():
            return Response({'detail': 'A generated tag was registered by another request. Please try again.'}, status=409)
        Asset.objects.bulk_create([
            Asset(**serializer.validated_data, created_by=request.user, updated_by=request.user)
            for serializer in serializers
        ])
    return Response({
        'created': len(prepared), 'item_types': len(rows),
        'first_tag': generated[0] if generated else None,
        'last_tag': generated[-1] if generated else None,
    }, status=201)


@api_view(['GET'])
def inventory_report(request):
    if not is_estate_officer(request.user):
        return Response({'detail': 'Estate Officer access required.'}, status=403)
    report_type = request.query_params.get('type', 'assets')
    wb = Workbook(); ws = wb.active
    if report_type == 'transfers':
        ws.title = 'Transfers'; ws.append(['Source Tag', 'Resulting Tag', 'Asset', 'Quantity', 'From', 'To', 'Responsible Office', 'Date', 'Reason'])
        for row in AssetTransfer.objects.select_related('asset', 'from_location', 'to_location'):
            ws.append([row.asset.asset_tag, row.resulting_asset.asset_tag if row.resulting_asset else '', row.asset.name, row.quantity, row.from_location.name, row.to_location.name,
                       row.new_responsible_office, row.transferred_at, row.reason])
    elif report_type == 'maintenance':
        ws.title = 'Maintenance'; ws.append(['Asset Tag', 'Asset', 'Quantity', 'Issue', 'Status', 'Provider', 'Cost', 'Reported', 'Completed'])
        for row in AssetMaintenance.objects.select_related('asset'):
            ws.append([row.asset.asset_tag, row.asset.name, row.quantity, row.issue, row.get_status_display(), row.provider,
                       row.cost, row.reported_date, row.completed_date])
    elif report_type == 'inspections':
        ws.title = 'Inspections'; ws.append(['Inspection Date', 'Location', 'Inspector', 'Status', 'Asset Tag', 'Asset', 'Result', 'Note'])
        for inspection in InventoryInspection.objects.select_related('location').prefetch_related('items__asset'):
            if inspection.items.exists():
                for item in inspection.items.all():
                    ws.append([inspection.inspection_date, inspection.location.name, inspection.inspector_name,
                               inspection.get_status_display(), item.asset.asset_tag, item.asset.name,
                               item.get_result_display(), item.note])
            else:
                ws.append([inspection.inspection_date, inspection.location.name, inspection.inspector_name,
                           inspection.get_status_display(), '', '', 'No items checked', ''])
    elif report_type == 'disposals':
        ws.title = 'Disposals'; ws.append(['Asset Tag', 'Asset', 'Status', 'Reason', 'Method', 'Proposed', 'Disposed', 'Reference'])
        for row in AssetDisposal.objects.select_related('asset'):
            ws.append([row.asset.asset_tag, row.asset.name, row.get_status_display(), row.reason, row.method,
                       row.proposed_date, row.disposal_date, row.reference])
    else:
        report_type = 'assets'; ws.title = 'Asset Register'
        ws.append(['Asset Tag', 'Asset Name', 'Description', 'Category', 'Location', 'Responsible Office', 'Quantity', 'Condition'])
        for row in Asset.objects.select_related('category', 'location'):
            ws.append([row.asset_tag, row.name, row.description, row.category.name, row.location.name,
                       row.responsible_office, row.quantity, row.get_condition_display()])
    for cell in ws[1]:
        cell.font = Font(bold=True, color='FFFFFF'); cell.fill = PatternFill('solid', fgColor='1F4E78')
    ws.freeze_panes = 'A2'; output = BytesIO(); wb.save(output)
    response = HttpResponse(output.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="inventory_{report_type}_report.xlsx"'
    return response


# ── FINANCE ───────────────────────────────────────────────────────────────────

class FinanceStudentViewSet(viewsets.ReadOnlyModelViewSet):
    serializer_class = FinanceStudentSerializer
    permission_classes = [IsFinanceUser]

    def get_queryset(self):
        qs = Student.objects.select_related(
            'module__class_level', 'module__semester__academic_year'
        )
        semester_id = self.request.query_params.get('semester_id')
        if not semester_id:
            sem = active_semester()
            if sem:
                qs = qs.filter(module__semester=sem)
        for param, field in [
            ('class_level_id', 'module__class_level_id'),
            ('semester_id', 'module__semester_id'),
        ]:
            val = self.request.query_params.get(param)
            if val:
                qs = qs.filter(**{field: val})
        search = str(self.request.query_params.get('search', '')).strip()
        if search:
            qs = qs.filter(
                Q(name__icontains=search)
                | Q(nactvet_reg_no__icontains=search)
                | Q(module__name__icontains=search)
                | Q(module__code__icontains=search)
            )
        return qs.order_by('module__class_level__order', 'name')


class PaymentCategoryViewSet(viewsets.ModelViewSet):
    serializer_class = PaymentCategorySerializer
    permission_classes = [IsFinanceUser]

    def get_queryset(self):
        qs = PaymentCategory.objects.select_related('created_by')
        semester_id = self.request.query_params.get('semester_id')
        if semester_id:
            qs = qs.filter(Q(semester_id=semester_id) | Q(semester__isnull=True))
        else:
            sem = active_semester()
            if sem:
                qs = qs.filter(Q(semester=sem) | Q(semester__isnull=True))
        class_level_id = self.request.query_params.get('class_level_id')
        if class_level_id:
            qs = qs.filter(Q(class_level_id=class_level_id) | Q(class_level__isnull=True))
        active = self.request.query_params.get('is_active')
        if active == 'true':
            qs = qs.filter(is_active=True)
        category_type = self.request.query_params.get('category_type')
        if category_type:
            qs = qs.filter(category_type=category_type)
        return qs.order_by('category_type', 'name')

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)


class StudentFinanceObligationViewSet(viewsets.ModelViewSet):
    serializer_class = StudentFinanceObligationSerializer
    permission_classes = [IsExamOfficerOrFinance]

    def get_queryset(self):
        qs = StudentFinanceObligation.objects.select_related(
            'student__module__class_level',
            'student__module__semester__academic_year',
            'semester__academic_year',
            'module',
            'category',
            'declared_by',
        ).prefetch_related('payments')
        semester_id = self.request.query_params.get('semester_id')
        if semester_id:
            qs = qs.filter(semester_id=semester_id)
        else:
            sem = active_semester()
            if sem:
                qs = qs.filter(semester=sem)
        for param, field in [
            ('student_id', 'student_id'),
            ('class_level_id', 'student__module__class_level_id'),
            ('obligation_type', 'obligation_type'),
        ]:
            val = self.request.query_params.get(param)
            if val:
                qs = qs.filter(**{field: val})
        return qs.order_by('-created_at')

    def perform_create(self, serializer):
        if not self.request.user.is_staff:
            raise PermissionDenied('Only the examination officer can declare special, supplementary, and repeat-module obligations.')
        serializer.save(declared_by=self.request.user)


class StudentPaymentViewSet(viewsets.ModelViewSet):
    serializer_class = StudentPaymentSerializer
    permission_classes = [IsFinanceUser]

    def get_queryset(self):
        qs = StudentPayment.objects.select_related(
            'student__module__class_level',
            'student__module__semester__academic_year',
            'category',
            'obligation',
            'recorded_by',
        )
        semester_id = self.request.query_params.get('semester_id')
        if semester_id:
            qs = qs.filter(student__module__semester_id=semester_id)
        else:
            sem = active_semester()
            if sem:
                qs = qs.filter(student__module__semester=sem)
        for param, field in [
            ('student_id', 'student_id'),
            ('category_id', 'category_id'),
            ('class_level_id', 'student__module__class_level_id'),
            ('obligation_id', 'obligation_id'),
        ]:
            val = self.request.query_params.get(param)
            if val:
                qs = qs.filter(**{field: val})
        return qs.order_by('-payment_date', '-created_at')

    def perform_create(self, serializer):
        category = serializer.validated_data.get('category')
        installment_number = serializer.validated_data.get('installment_number') or 1
        if category and installment_number > category.installment_count:
            raise PermissionDenied(f'This category allows only {category.installment_count} installment(s).')
        if serializer.validated_data.get('amount_required') in (None, 0) and category:
            serializer.validated_data['amount_required'] = category.default_amount
        serializer.save(recorded_by=self.request.user)


class StudentFinanceClearanceViewSet(viewsets.ModelViewSet):
    serializer_class = StudentFinanceClearanceSerializer
    permission_classes = [IsFinanceUser]

    def get_queryset(self):
        qs = StudentFinanceClearance.objects.select_related(
            'student__module__class_level',
            'student__module__semester__academic_year',
            'semester__academic_year',
            'approved_by',
        )
        semester_id = self.request.query_params.get('semester_id')
        if semester_id:
            qs = qs.filter(semester_id=semester_id)
        else:
            sem = active_semester()
            if sem:
                qs = qs.filter(semester=sem)
        for param, field in [
            ('student_id', 'student_id'),
            ('class_level_id', 'student__module__class_level_id'),
            ('period', 'period'),
        ]:
            val = self.request.query_params.get(param)
            if val:
                qs = qs.filter(**{field: val})
        return qs.order_by('student__name', 'period')

    def perform_create(self, serializer):
        serializer.save(approved_by=self.request.user)

    def perform_update(self, serializer):
        serializer.save(approved_by=self.request.user)


# ── REPORT ─────────────────────────────────────────────────────────────────────

@api_view(['GET'])
@login_required
def report(request):
    my_modules = user_modules(request.user)
    module_id = request.query_params.get('module_id')
    class_level_id = request.query_params.get('class_level_id')
    semester_id = request.query_params.get('semester_id')

    students = (
        Student.objects.filter(module__in=my_modules)
        .select_related('module__class_level', 'module__semester__academic_year')
        .prefetch_related('attendance_records__session')
    )
    if module_id:
        students = students.filter(module_id=module_id)
    if class_level_id:
        students = students.filter(module__class_level_id=class_level_id)
    if semester_id:
        students = students.filter(module__semester_id=semester_id)

    # Cache session counts per module {module_id: {'T': n, 'P': n, 'total': n}}
    _mod_cache = {}

    def _mcounts(mid):
        if mid not in _mod_cache:
            t = Session.objects.filter(module_id=mid, session_type=Session.THEORY).count()
            p = Session.objects.filter(module_id=mid, session_type=Session.PRACTICAL).count()
            _mod_cache[mid] = {'T': t, 'P': p, 'total': t + p}
        return _mod_cache[mid]

    rows = []
    total_eff, total_held, at_risk, critical, sick_total = 0, 0, 0, 0, 0

    for st in students:
        mc = _mcounts(st.module_id)
        sessions_held = mc['total']

        # Use prefetched attendance_records__session
        all_records = list(st.attendance_records.all())
        attended = sum(1 for r in all_records if r.status == 'P')
        sick = sum(1 for r in all_records if r.status == 'S')
        effective = attended + sick
        absent = max(sessions_held - effective, 0)

        # Theory / Practical breakdown
        theory_eff = sum(1 for r in all_records if r.session.session_type == Session.THEORY and r.status in ('P', 'S'))
        practical_eff = sum(1 for r in all_records if r.session.session_type == Session.PRACTICAL and r.status in ('P', 'S'))

        pct = round((effective / sessions_held) * 100) if sessions_held else 0
        theory_pct = round((theory_eff / mc['T']) * 100) if mc['T'] else None
        practical_pct = round((practical_eff / mc['P']) * 100) if mc['P'] else None

        total_eff += effective
        total_held += sessions_held
        sick_total += sick
        if 50 <= pct < 75:
            at_risk += 1
        elif pct < 50:
            critical += 1

        rows.append({
            'nactvet_reg_no': st.nactvet_reg_no,
            'name': st.name,
            'module': st.module.name,
            'module_code': st.module.code,
            'class_level': st.module.class_level.name,
            'semester': st.module.semester.label,
            'teacher': st.module.teacher,
            'sessions_held': sessions_held,
            'theory_held': mc['T'],
            'practical_held': mc['P'],
            'attended': attended,
            'sick': sick,
            'absent': absent,
            'theory_eff': theory_eff,
            'practical_eff': practical_eff,
            'pct': pct,
            'theory_pct': theory_pct,
            'practical_pct': practical_pct,
            'status': 'Good' if pct >= 75 else ('At Risk' if pct >= 50 else 'Critical'),
        })

    avg_pct = round((total_eff / total_held) * 100) if total_held else 0

    sessions_qs = (
        Session.objects.filter(module__in=my_modules)
        .select_related('module__class_level', 'module__semester__academic_year')
        .prefetch_related('records')
    )
    if module_id:
        sessions_qs = sessions_qs.filter(module_id=module_id)
    if class_level_id:
        sessions_qs = sessions_qs.filter(module__class_level_id=class_level_id)
    if semester_id:
        sessions_qs = sessions_qs.filter(module__semester_id=semester_id)

    history = []
    for sess in sessions_qs:
        p = sess.records.filter(status='P').count()
        s = sess.records.filter(status='S').count()
        a = sess.records.filter(status='A').count()
        t = p + s + a
        history.append({
            'id': sess.id,
            'date': str(sess.date),
            'module': sess.module.name,
            'session_type': sess.session_type,
            'session_type_display': sess.get_session_type_display(),
            'class_level': sess.module.class_level.name,
            'semester': sess.module.semester.label,
            'label': sess.label,
            'topic': sess.topic,
            'present': p, 'sick': s, 'absent': a,
            'pct': round(((p + s) / t) * 100) if t else 0,
        })

    return Response({
        'stats': {
            'students': len(rows),
            'avg_pct': avg_pct,
            'at_risk': at_risk,
            'critical': critical,
            'sick_total': sick_total,
        },
        'rows': rows,
        'session_history': history,
    })


# ── ALL MODULES (for claim/unclaim) ───────────────────────────────────────────

@api_view(['GET'])
@login_required
def all_modules(request):
    my_ids = set(request.user.modules_taught.values_list('id', flat=True))
    data = []
    for m in Module.objects.select_related('class_level', 'semester__academic_year').all():
        data.append({
            'id': m.id, 'name': m.name, 'code': m.code,
            'class_level': m.class_level.name,
            'semester': m.semester.label,
            'claimed': m.id in my_ids,
        })
    return Response(data)


# ── ELIGIBILITY ────────────────────────────────────────────────────────────────

ELIGIBILITY_THRESHOLD = 90
CA_ELIGIBILITY_THRESHOLD = 20


def ca_eligibility_for_student(student):
    result = getattr(student, 'result', None)
    if student.module.is_field_module:
        if result is None or result.field_ca is None:
            return None, None, 'Pending Field CA mark'
        total_ca = round(float(result.field_ca) * 0.4, 2)
        eligible = float(result.field_ca) >= 50
        return total_ca, eligible, (
            f'Field CA {result.field_ca}/100 meets the requirement'
            if eligible else f'Field CA {result.field_ca}/100 is below the required 50/100'
        )
    has_practical = student.module.has_practical
    required_fields = [
        ('assign1', 'Assignment 1'),
        ('assign2', 'Assignment 2'),
        ('cat1_theory', 'CAT 1 Theory'),
        ('cat2_theory', 'CAT 2 Theory'),
    ]
    if has_practical:
        required_fields.extend([
            ('cat1_practical', 'Practical Test 1'),
            ('cat2_practical', 'Practical Test 2'),
        ])

    if result is None:
        return None, None, 'Pending CA result record'

    missing = [label for field, label in required_fields
               if getattr(result, field) is None and not getattr(result, f'{field}_absent')]
    serializer = StudentResultSerializer()
    total_ca = serializer.get_total_ca(result)

    if missing:
        return total_ca, None, f"Pending CA marks: {', '.join(missing)}"
    if total_ca is None:
        return None, None, 'Pending CA marks'
    if total_ca < CA_ELIGIBILITY_THRESHOLD:
        return total_ca, False, f'Total CA {total_ca}/40 is below required {CA_ELIGIBILITY_THRESHOLD}/40'
    return total_ca, True, f'Total CA {total_ca}/40 meets required {CA_ELIGIBILITY_THRESHOLD}/40'


def attendance_eligibility_reason(pct, sessions, label):
    if not sessions:
        return f'Pending {label} attendance sessions'
    if pct is None:
        return f'Pending {label} attendance'
    if pct < ELIGIBILITY_THRESHOLD:
        return f'{label} attendance {pct}% is below required {ELIGIBILITY_THRESHOLD}%'
    return f'{label} attendance {pct}% meets required {ELIGIBILITY_THRESHOLD}%'


def finance_note(result, detailed):
    """What to say about a student's finance clearance.

    The accountant and the student themselves see the amount. The examination
    officer and tutors see only whether the student is cleared — they need that
    to run an exam, but a student's balance is not theirs to read.
    """
    if detailed:
        return result['reason']
    return 'Finance cleared' if result['cleared'] else 'Pending finance clearance'


def _clearance_for(enrollment, period, cache=None):
    """Finance clearance for the person behind this enrollment.

    Fees belong to the student, not to each module they take, so this resolves
    the enrollment to its profile first. A student with no charges raised yet
    is cleared — the ledger blocks nobody until the college has actually billed
    them.

    Pass a dict as `cache` when looping over enrollments: the answer is the
    same for every module a student takes, so without it a student in eight
    modules has their clearance recomputed eight times per period.
    """
    profile = enrollment.profile
    if profile is None:
        profile = finance.profile_for_student(enrollment)
    semester = enrollment.module.semester

    if cache is None:
        return finance.exam_clearance(
            profile, semester.academic_year, period, semester=semester)

    key = (profile.id, semester.id, period)
    if key not in cache:
        cache[key] = finance.exam_clearance(
            profile, semester.academic_year, period, semester=semester)
    return cache[key]


EXAM_PERIODS = [ChargeType.CAT1, ChargeType.CAT2, ChargeType.FINAL]


def _batch_clearance(enrollments):
    """Pre-compute clearance for every person in a list of enrollments.

    Returns a cache in the shape _clearance_for() expects, so the loop below
    reads from memory instead of issuing two queries per student per period.
    """
    by_semester = {}
    for enrollment in enrollments:
        profile = enrollment.profile or finance.profile_for_student(enrollment)
        by_semester.setdefault(enrollment.module.semester, set()).add(profile.id)

    cache = {}
    for semester, profile_ids in by_semester.items():
        results = finance.clearance_map(
            profile_ids, semester.academic_year, EXAM_PERIODS, semester=semester,
        )
        for profile_id, per_period in results.items():
            for period, result in per_period.items():
                cache[(profile_id, semester.id, period)] = result
    return cache


def combined_eligibility(parts):
    if any(value is False for value, _note in parts):
        return False
    if any(value is None for value, _note in parts):
        return None
    return True


def combined_reason(parts):
    return '; '.join(note for _value, note in parts if note)


@api_view(['GET'])
@login_required
def eligibility(request):
    my_modules = user_modules(request.user)
    module_id = request.query_params.get('module_id')
    class_level_id = request.query_params.get('class_level_id')
    semester_id = request.query_params.get('semester_id')

    students = (
        Student.objects.filter(module__in=my_modules)
        .select_related('module__class_level', 'module__semester__academic_year')
        .prefetch_related('attendance_records__session')
        .select_related('result', 'profile')
    )
    if module_id:
        students = students.filter(module_id=module_id)
    if class_level_id:
        students = students.filter(module__class_level_id=class_level_id)
    if semester_id:
        students = students.filter(module__semester_id=semester_id)

    _mod_cache = {}
    # Clearance is per person per period, not per enrollment — a student in
    # eight modules has one answer, not eight. Fetch the whole page's worth in
    # a fixed number of queries, then look each row's answer up.
    students = list(students)
    _clearance_cache = _batch_clearance(students)
    show_amounts = can_manage_finance(request.user)

    def _period_counts(mid):
        if mid not in _mod_cache:
            cat1 = Session.objects.filter(module_id=mid, exam_period=Session.CAT1).count()
            cat2 = Session.objects.filter(module_id=mid, exam_period=Session.CAT2).count()
            total = Session.objects.filter(module_id=mid).count()
            _mod_cache[mid] = {'cat1': cat1, 'cat2': cat2, 'total': total}
        return _mod_cache[mid]

    rows = []
    for st in students:
        mc = _period_counts(st.module_id)
        all_records = list(st.attendance_records.all())

        cat1_eff = sum(1 for r in all_records if r.session.exam_period == Session.CAT1 and attendance_is_effective(r))
        cat2_eff = sum(1 for r in all_records if r.session.exam_period == Session.CAT2 and attendance_is_effective(r))
        total_eff = sum(1 for r in all_records if attendance_is_effective(r))

        cat1_pct = round((cat1_eff / mc['cat1']) * 100) if mc['cat1'] else None
        cat2_pct = round((cat2_eff / mc['cat2']) * 100) if mc['cat2'] else None
        end_pct = round((total_eff / mc['total']) * 100) if mc['total'] else None
        cat1_eligible = (cat1_pct >= ELIGIBILITY_THRESHOLD) if mc['cat1'] else None
        cat2_eligible = (cat2_pct >= ELIGIBILITY_THRESHOLD) if mc['cat2'] else None
        end_eligible = (end_pct >= ELIGIBILITY_THRESHOLD) if mc['total'] else None
        total_ca, ca_eligible, ca_note = ca_eligibility_for_student(st)
        # Finance clearance is computed from the ledger, not read off a flag
        # somebody remembered to tick. A student with nothing outstanding is
        # cleared the moment their payment is recorded.
        cat1_fin = _clearance_for(st, ChargeType.CAT1, _clearance_cache)
        cat2_fin = _clearance_for(st, ChargeType.CAT2, _clearance_cache)
        end_fin = _clearance_for(st, ChargeType.FINAL, _clearance_cache)

        cat1_finance_eligible = cat1_fin['cleared']
        cat2_finance_eligible = cat2_fin['cleared']
        end_finance_eligible = end_fin['cleared']
        cat1_fin_note = finance_note(cat1_fin, show_amounts)
        cat2_fin_note = finance_note(cat2_fin, show_amounts)
        end_fin_note = finance_note(end_fin, show_amounts)
        cat1_att_note = attendance_eligibility_reason(cat1_pct, mc['cat1'], 'CAT 1')
        cat2_att_note = attendance_eligibility_reason(cat2_pct, mc['cat2'], 'CAT 2')
        end_att_note = attendance_eligibility_reason(end_pct, mc['total'], 'End-of-semester')
        cat1_exam_parts = [
            (cat1_eligible, cat1_att_note),
            (cat1_finance_eligible, cat1_fin_note),
        ]
        cat2_exam_parts = [
            (cat2_eligible, cat2_att_note),
            (cat2_finance_eligible, cat2_fin_note),
        ]
        end_exam_parts = [
            (end_eligible, end_att_note),
            (ca_eligible, ca_note),
            (end_finance_eligible, end_fin_note),
        ]

        rows.append({
            'id': st.id,
            'nactvet_reg_no': st.nactvet_reg_no,
            'name': st.name,
            'module': st.module.name,
            'module_code': st.module.code,
            'class_level': st.module.class_level.name,
            'semester': st.module.semester.label,
            'cat1_sessions': mc['cat1'],
            'cat1_attended': cat1_eff,
            'cat1_pct': cat1_pct,
            'cat1_eligible': cat1_eligible,
            'cat1_note': cat1_att_note,
            'cat1_finance_eligible': cat1_finance_eligible,
            'cat1_finance_note': cat1_fin_note,
            'cat1_exam_eligible': combined_eligibility(cat1_exam_parts),
            'cat1_exam_note': combined_reason(cat1_exam_parts),
            'cat2_sessions': mc['cat2'],
            'cat2_attended': cat2_eff,
            'cat2_pct': cat2_pct,
            'cat2_eligible': cat2_eligible,
            'cat2_note': cat2_att_note,
            'cat2_finance_eligible': cat2_finance_eligible,
            'cat2_finance_note': cat2_fin_note,
            'cat2_exam_eligible': combined_eligibility(cat2_exam_parts),
            'cat2_exam_note': combined_reason(cat2_exam_parts),
            'end_sessions': mc['total'],
            'end_attended': total_eff,
            'end_pct': end_pct,
            'end_eligible': end_eligible,
            'end_note': end_att_note,
            'end_finance_eligible': end_finance_eligible,
            'end_finance_note': end_fin_note,
            'end_exam_eligible': combined_eligibility(end_exam_parts),
            'end_exam_note': combined_reason(end_exam_parts),
            'ca_total': total_ca,
            'ca_eligible': ca_eligible,
            'ca_note': ca_note,
        })

    stats = {
        'total_students': len(rows),
        'cat1_eligible': sum(1 for r in rows if r['cat1_eligible'] is True),
        'cat1_ineligible': sum(1 for r in rows if r['cat1_eligible'] is False),
        'cat1_na': sum(1 for r in rows if r['cat1_eligible'] is None),
        'cat2_eligible': sum(1 for r in rows if r['cat2_eligible'] is True),
        'cat2_ineligible': sum(1 for r in rows if r['cat2_eligible'] is False),
        'cat2_na': sum(1 for r in rows if r['cat2_eligible'] is None),
        'end_eligible': sum(1 for r in rows if r['end_eligible'] is True),
        'end_ineligible': sum(1 for r in rows if r['end_eligible'] is False),
        'end_na': sum(1 for r in rows if r['end_eligible'] is None),
        'ca_eligible': sum(1 for r in rows if r['ca_eligible'] is True),
        'ca_ineligible': sum(1 for r in rows if r['ca_eligible'] is False),
        'ca_na': sum(1 for r in rows if r['ca_eligible'] is None),
        'cat1_exam_eligible': sum(1 for r in rows if r['cat1_exam_eligible'] is True),
        'cat1_exam_ineligible': sum(1 for r in rows if r['cat1_exam_eligible'] is False),
        'cat1_exam_na': sum(1 for r in rows if r['cat1_exam_eligible'] is None),
        'cat2_exam_eligible': sum(1 for r in rows if r['cat2_exam_eligible'] is True),
        'cat2_exam_ineligible': sum(1 for r in rows if r['cat2_exam_eligible'] is False),
        'cat2_exam_na': sum(1 for r in rows if r['cat2_exam_eligible'] is None),
        'end_exam_eligible': sum(1 for r in rows if r['end_exam_eligible'] is True),
        'end_exam_ineligible': sum(1 for r in rows if r['end_exam_eligible'] is False),
        'end_exam_na': sum(1 for r in rows if r['end_exam_eligible'] is None),
    }

    return Response({
        'stats': stats,
        'rows': rows,
        'threshold': ELIGIBILITY_THRESHOLD,
        'ca_threshold': CA_ELIGIBILITY_THRESHOLD,
    })


# ── SICK RECORDS ───────────────────────────────────────────────────────────────

@api_view(['GET'])
@login_required
def sick_records(request):
    my_modules = user_modules(request.user)
    module_id = request.query_params.get('module_id')
    semester_id = request.query_params.get('semester_id')
    class_level_id = request.query_params.get('class_level_id')

    records = (
        AttendanceRecord.objects.filter(
            status='S',
            student__module__in=my_modules,
        )
        .select_related(
            'student__module__class_level',
            'student__module__semester__academic_year',
            'session',
        )
        .order_by('-session__date')
    )

    if module_id:
        records = records.filter(student__module_id=module_id)
    if semester_id:
        records = records.filter(student__module__semester_id=semester_id)
    if class_level_id:
        records = records.filter(student__module__class_level_id=class_level_id)

    data = [
        {
            'id': r.id,
            'student_id': r.student.id,
            'student_name': r.student.name,
            'student_reg_no': r.student.nactvet_reg_no,
            'module': r.student.module.name,
            'module_code': r.student.module.code,
            'class_level': r.student.module.class_level.name,
            'semester': r.student.module.semester.label,
            'session_date': str(r.session.date),
            'session_label': r.session.label,
            'exam_period': r.session.exam_period,
            'exam_period_display': r.session.get_exam_period_display(),
            'sick_note': r.sick_note,
            'certificate_submitted': r.certificate_submitted,
        }
        for r in records
    ]

    return Response(data)


@api_view(['PATCH'])
@login_required
def update_sick_record(request, pk):
    if not request.user.is_staff:
        return Response({'detail': 'Only the administrator can update sick records.'}, status=status.HTTP_403_FORBIDDEN)

    try:
        record = AttendanceRecord.objects.select_related('student__module').get(pk=pk, status='S')
    except AttendanceRecord.DoesNotExist:
        return Response({'detail': 'Not found.'}, status=status.HTTP_404_NOT_FOUND)

    my_mod_ids = set(user_modules(request.user).values_list('id', flat=True))
    if record.student.module_id not in my_mod_ids:
        return Response({'detail': 'Forbidden.'}, status=status.HTTP_403_FORBIDDEN)

    if 'sick_note' in request.data:
        record.sick_note = str(request.data['sick_note']).strip()
    if 'certificate_submitted' in request.data:
        record.certificate_submitted = bool(request.data['certificate_submitted'])
    record.save()

    return Response({
        'id': record.id,
        'sick_note': record.sick_note,
        'certificate_submitted': record.certificate_submitted,
    })


@api_view(['PATCH'])
@login_required
def update_attendance_status(request, pk):
    """Admin-only: change the status of an attendance record.

    Allows setting status to 'P', 'A' or 'S'. When marking as 'S', an optional
    `sick_note` and `certificate_submitted` may be provided. Non-staff users
    are forbidden.
    """
    if not request.user.is_staff:
        return Response({'detail': 'Only the administrator can change attendance status.'}, status=status.HTTP_403_FORBIDDEN)

    try:
        record = AttendanceRecord.objects.select_related('student__module').get(pk=pk)
    except AttendanceRecord.DoesNotExist:
        return Response({'detail': 'Not found.'}, status=status.HTTP_404_NOT_FOUND)

    my_mod_ids = set(user_modules(request.user).values_list('id', flat=True))
    if record.student.module_id not in my_mod_ids:
        return Response({'detail': 'Forbidden.'}, status=status.HTTP_403_FORBIDDEN)

    # Validate and apply status change
    if 'status' in request.data:
        new_status = str(request.data.get('status')).upper()
        if new_status not in (AttendanceRecord.PRESENT, AttendanceRecord.ABSENT, AttendanceRecord.SICK):
            return Response({'detail': 'Invalid status.'}, status=status.HTTP_400_BAD_REQUEST)
        record.status = new_status
        if new_status != AttendanceRecord.SICK:
            # Clear sick-related fields when not sick
            record.sick_note = ''
            record.certificate_submitted = False

    if 'sick_note' in request.data and record.status == AttendanceRecord.SICK:
        record.sick_note = str(request.data['sick_note']).strip()

    if 'certificate_submitted' in request.data:
        record.certificate_submitted = bool(request.data['certificate_submitted'])

    record.save()

    return Response(AttendanceRecordSerializer(record).data)


# ── RESULTS ────────────────────────────────────────────────────────────────────

class ResultViewSet(viewsets.ModelViewSet):
    """
    Manage CA marks for students.
    Teachers can read/write results for their own modules.
    Admin can read/write all and download Excel.
    """
    serializer_class   = StudentResultSerializer
    permission_classes = [IsAuthenticated, ReadExamsWriteExamOfficer]

    def get_queryset(self):
        qs = (
            StudentResult.objects
            .filter(student__module__in=user_modules(self.request.user))
            .select_related(
                'student__module__class_level',
                'student__module__semester__academic_year',
            )
        )
        module_id = self.request.query_params.get('module_id')
        if module_id:
            qs = qs.filter(student__module_id=module_id)
        return qs

    def update(self, request, *args, **kwargs):
        check_result_entry(request.user, self.get_object().student.module, request.data)
        return super().update(request, *args, **kwargs)

    def perform_create(self, serializer):
        student = serializer.validated_data['student']
        allowed_module_ids = user_modules(self.request.user).values_list('id', flat=True)
        if student.module_id not in allowed_module_ids:
            raise PermissionDenied('You may only create results for modules you tutor.')
        # Only the fields actually carrying a value count as being written; a
        # serializer fills the rest in with None whether the tutor sent them or not.
        supplied = {field for field, value in serializer.validated_data.items()
                    if value not in (None, False)}
        check_result_entry(self.request.user, student.module, supplied)
        serializer.save()

    @action(detail=False, methods=['post'], url_path='authority-grades')
    def authority_grades(self, request):
        if not request.user.is_staff:
            raise PermissionDenied('Only the administrator can upload authority grades.')
        rows = request.data.get('rows', []) if isinstance(request.data, dict) else []
        if not rows:
            return Response({'detail': 'No grade rows supplied.'}, status=status.HTTP_400_BAD_REQUEST)

        saved, errors = 0, []
        with transaction.atomic():
            for index, row in enumerate(rows, start=2):
                reg_no = str(row.get('reg_no', '')).strip()
                module_code = str(row.get('module_code', '')).strip()
                raw = str(row.get('grade', '')).strip().upper()
                if not reg_no or not module_code or not raw:
                    continue
                student = Student.objects.filter(
                    nactvet_reg_no__iexact=reg_no, module__code__iexact=module_code,
                ).select_related('module').first()
                if student is None:
                    errors.append(f'Row {index}: no enrollment for {reg_no} / {module_code}')
                    continue
                try:
                    parsed = parse_authority_grade(raw, student.module.class_level)
                except ValueError as exc:
                    errors.append(f'Row {index}: {exc} for {reg_no} / {module_code}')
                    continue
                result, _ = StudentResult.objects.get_or_create(student=student)
                result.authority_grade = parsed['raw']
                result.authority_status = parsed['status']
                result.final_approved = True
                result.save(update_fields=['authority_grade', 'authority_status', 'final_approved', 'updated_at'])
                saved += 1
        return Response({'saved': saved, 'errors': errors})

    # ── Get or create results for every student in a module ────────────────────
    @action(detail=False, methods=['get'], url_path='module')
    def module_results(self, request):
        module_id = request.query_params.get('module_id')
        if not module_id:
            return Response({'detail': 'module_id required.'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            module = Module.objects.select_related('class_level', 'semester__academic_year').get(pk=module_id)
        except Module.DoesNotExist:
            return Response({'detail': 'Module not found.'}, status=status.HTTP_404_NOT_FOUND)

        mod_ids = set(user_modules(request.user).values_list('id', flat=True))
        if module.id not in mod_ids:
            raise PermissionDenied('You do not tutor this module.')

        students = Student.objects.filter(module=module).order_by('name')
        with transaction.atomic():
            results = [StudentResult.objects.get_or_create(student=st)[0] for st in students]

        return Response({
            'module': {
                'id':           module.id,
                'name':         module.name,
                'code':         module.code,
                'has_practical': module.has_practical,
                'is_field_module': module.is_field_module,
                'credits':       module.credits,
                'class_level':  module.class_level.name,
                'semester_id':   module.semester_id,
                'semester':     module.semester.label,
            },
            'results': StudentResultSerializer(results, many=True).data,
        })

    @action(detail=False, methods=['get'], url_path='cat-analysis')
    def cat_analysis(self, request):
        cat = str(request.query_params.get('cat', '1'))
        if cat not in ('1', '2'):
            return Response(
                {'detail': 'cat must be 1 or 2.'},
                status=status.HTTP_400_BAD_REQUEST,
            )
        semester_id = request.query_params.get('semester_id')
        if not semester_id:
            return Response(
                {'detail': 'Select a semester for CAT analysis.'},
                status=status.HTTP_400_BAD_REQUEST,
            )
        semester = Semester.objects.select_related('academic_year').filter(
            pk=semester_id
        ).first()
        if semester is None:
            return Response(
                {'detail': 'Semester not found.'},
                status=status.HTTP_404_NOT_FOUND,
            )

        qs = (
            StudentResult.objects
            .filter(student__module__in=user_modules(request.user))
            .select_related(
                'student__module__class_level',
                'student__module__semester__academic_year',
            )
            .order_by('student__module__class_level__order', 'student__module__code')
        )
        for param, field in (
            ('module_id', 'student__module_id'),
            ('semester_id', 'student__module__semester_id'),
            ('class_level_id', 'student__module__class_level_id'),
        ):
            value = request.query_params.get(param)
            if value:
                qs = qs.filter(**{field: value})

        theory_field = f'cat{cat}_theory'
        practical_field = f'cat{cat}_practical'
        modules = {}
        all_grade_counts = {}
        overall_assessed = overall_passed = overall_failed = overall_incomplete = 0

        for result in qs:
            module = result.student.module
            row = modules.setdefault(module.id, {
                'module_id': module.id,
                'module_code': module.code,
                'module_name': module.name,
                'class_level': module.class_level.name,
                'semester': module.semester.label,
                'has_practical': module.has_practical,
                'assessed': 0, 'passed': 0, 'failed': 0, 'incomplete': 0,
                'grade_counts': {},
            })

            theory_complete = (
                getattr(result, theory_field) is not None
                or getattr(result, f'{theory_field}_absent')
            )
            practical_complete = (
                not module.has_practical
                or getattr(result, practical_field) is not None
                or getattr(result, f'{practical_field}_absent')
            )
            if not theory_complete or not practical_complete:
                row['incomplete'] += 1
                overall_incomplete += 1
                continue

            theory = (
                0.0 if getattr(result, f'{theory_field}_absent')
                else float(getattr(result, theory_field))
            )
            if module.has_practical:
                practical = (
                    0.0 if getattr(result, f'{practical_field}_absent')
                    else float(getattr(result, practical_field))
                )
                mark = round((theory + practical) / 2, 2)
            else:
                mark = theory

            grade, _, _ = grade_for_mark(mark, module.class_level)
            row['assessed'] += 1
            row['grade_counts'][grade] = row['grade_counts'].get(grade, 0) + 1
            all_grade_counts[grade] = all_grade_counts.get(grade, 0) + 1
            overall_assessed += 1
            if mark >= 50:
                row['passed'] += 1
                overall_passed += 1
            else:
                row['failed'] += 1
                overall_failed += 1

        rows = list(modules.values())
        for row in rows:
            row['pass_rate'] = (
                round(row['passed'] / row['assessed'] * 100, 1)
                if row['assessed'] else 0
            )

        return Response({
            'cat': int(cat),
            'semester': semester.label,
            'stats': {
                'modules': len(rows),
                'assessed': overall_assessed,
                'passed': overall_passed,
                'failed': overall_failed,
                'incomplete': overall_incomplete,
                'pass_rate': (
                    round(overall_passed / overall_assessed * 100, 1)
                    if overall_assessed else 0
                ),
                'grade_counts': all_grade_counts,
            },
            'rows': rows,
        })

    @action(detail=False, methods=['get'], url_path='cat-analysis/download')
    def download_cat_analysis(self, request):
        analysis_response = self.cat_analysis(request)
        if analysis_response.status_code != status.HTTP_200_OK:
            return analysis_response
        data = analysis_response.data

        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = f'CAT {data["cat"]} Analysis'
        ws.sheet_view.showGridLines = False
        ws.freeze_panes = 'A8'
        ws.page_setup.orientation = 'landscape'
        ws.page_setup.fitToWidth = 1
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.print_title_rows = '7:7'

        navy = '1E2D78'
        blue = 'DCE6F1'
        green = 'DCFCE7'
        red = 'FEE2E2'
        yellow = 'FEF3C7'
        white = 'FFFFFF'
        thin = Side(style='thin', color='B8C0CC')
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        headers = [
            'Module Code', 'Module Name', 'Level', 'Assessed',
            'A', 'B+', 'B', 'C', 'D', 'F',
            'Passed', 'Failed', 'Incomplete', 'Pass Rate',
        ]
        last_col = len(headers)
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=last_col)
        title = ws.cell(1, 1, f'EDUTRACK — CAT {data["cat"]} MODULE ANALYSIS')
        title.font = Font(bold=True, size=16, color=white)
        title.fill = PatternFill('solid', fgColor=navy)
        title.alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 30

        filters = []
        for param, label, model in (
            ('semester_id', 'Semester', Semester),
            ('class_level_id', 'Class Level', ClassLevel),
            ('module_id', 'Module', Module),
        ):
            value = request.query_params.get(param)
            if value:
                obj = model.objects.filter(pk=value).first()
                filters.append(f'{label}: {obj or value}')
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=last_col)
        ws.cell(2, 1, ' | '.join(filters) if filters else 'Scope: All accessible modules').alignment = Alignment(horizontal='center')

        stats = data['stats']
        summary = [
            ('Modules', stats['modules']), ('Assessed', stats['assessed']),
            ('Passed', stats['passed']), ('Failed', stats['failed']),
            ('Incomplete', stats['incomplete']), ('Total Pass Rate', f'{stats["pass_rate"]}%'),
        ]
        for index, (label, value) in enumerate(summary, 1):
            column = 1 + (index - 1) * 2
            ws.merge_cells(start_row=4, start_column=column, end_row=4, end_column=column + 1)
            ws.merge_cells(start_row=5, start_column=column, end_row=5, end_column=column + 1)
            ws.cell(4, column, label).font = Font(bold=True, color=navy)
            ws.cell(4, column).alignment = Alignment(horizontal='center')
            ws.cell(5, column, value).font = Font(bold=True, size=13)
            ws.cell(5, column).alignment = Alignment(horizontal='center')
            for row_number in (4, 5):
                for col in range(column, column + 2):
                    ws.cell(row_number, col).fill = PatternFill('solid', fgColor=blue)
                    ws.cell(row_number, col).border = border

        for col, header in enumerate(headers, 1):
            cell = ws.cell(7, col, header)
            cell.font = Font(bold=True, color=white)
            cell.fill = PatternFill('solid', fgColor=navy)
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = border

        grades = ['A', 'B+', 'B', 'C', 'D', 'F']
        for row_number, row in enumerate(data['rows'], 8):
            values = [
                row['module_code'], row['module_name'], row['class_level'],
                row['assessed'],
                *[row['grade_counts'].get(grade, 0) for grade in grades],
                row['passed'], row['failed'], row['incomplete'],
                row['pass_rate'] / 100,
            ]
            for col, value in enumerate(values, 1):
                cell = ws.cell(row_number, col, value)
                cell.border = border
                cell.alignment = Alignment(
                    horizontal='left' if col in (1, 2, 3) else 'center',
                    vertical='center',
                    wrap_text=True,
                )
                if row_number % 2 == 0:
                    cell.fill = PatternFill('solid', fgColor='F8FAFC')
            ws.cell(row_number, 11).fill = PatternFill('solid', fgColor=green)
            ws.cell(row_number, 12).fill = PatternFill('solid', fgColor=red)
            ws.cell(row_number, 13).fill = PatternFill('solid', fgColor=yellow)
            ws.cell(row_number, 14).number_format = '0.0%'

        widths = [14, 30, 18, 10, 7, 7, 7, 7, 7, 7, 10, 10, 12, 12]
        for col, width in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(col)].width = width
        ws.auto_filter.ref = f'A7:{get_column_letter(last_col)}{max(7, 7 + len(data["rows"]))}'
        ws.print_area = f'A1:{get_column_letter(last_col)}{max(7, 7 + len(data["rows"]))}'

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = (
            f'attachment; filename="cat_{data["cat"]}_module_analysis_{timezone.localdate()}.xlsx"'
        )
        wb.save(response)
        return response

    # ── Bulk-save marks submitted from the frontend ────────────────────────────
    @action(detail=False, methods=['post'], url_path='bulk_save')
    def bulk_save(self, request):
        updates = request.data if isinstance(request.data, list) else []
        if not updates:
            return Response({'detail': 'Empty list.'}, status=status.HTTP_400_BAD_REQUEST)

        mod_ids   = set(user_modules(request.user).values_list('id', flat=True))
        CA_MARKS = ['assign1', 'assign2', 'cat1_theory', 'cat2_theory', 'cat1_practical', 'cat2_practical']
        CA_FIELDS = ['field_ca'] + CA_MARKS + [f'{field}_absent' for field in CA_MARKS]
        END_MARKS = ['end_theory', 'end_practical', 'supplementary_mark']
        END_FIELDS = END_MARKS + [f'{field}_absent' for field in END_MARKS] + ['ca_approved', 'final_approved']
        FIELDS     = CA_FIELDS + (END_FIELDS if request.user.is_staff else [])
        saved, errors = 0, []

        for item in updates:
            if not request.user.is_staff:
                restricted = {'end_theory', 'end_practical', 'supplementary_mark', 'end_theory_absent',
                              'end_practical_absent', 'ca_approved', 'final_approved'}
                blocked = restricted.intersection(item.keys())
                if blocked:
                    errors.append(f'Result {item.get("id")}: administrator approval required')
                    continue

            try:
                result = (
                    StudentResult.objects
                    .select_related('student__module')
                    .get(pk=item.get('id'))
                )
            except StudentResult.DoesNotExist:
                errors.append(f'Result {item.get("id")} not found')
                continue

            if result.student.module_id not in mod_ids:
                errors.append(f'Result {item.get("id")}: permission denied')
                continue

            for field in FIELDS:
                if field not in item:
                    continue
                raw = item[field]
                if field in ('ca_approved', 'final_approved') or field.endswith('_absent'):
                    setattr(result, field, bool(raw))
                    if field.endswith('_absent') and bool(raw):
                        setattr(result, field[:-7], None)
                    continue
                if raw == '' or raw is None:
                    setattr(result, field, None)
                else:
                    try:
                        v = float(raw)
                        if not (0 <= v <= 100):
                            errors.append(f'Result {item.get("id")}: {field} must be 0–100')
                            continue
                        setattr(result, field, v)
                        absent_field = f'{field}_absent'
                        if hasattr(result, absent_field):
                            setattr(result, absent_field, False)
                    except (TypeError, ValueError):
                        errors.append(f'Result {item.get("id")}: invalid value for {field}')
                        continue
            result.save()
            saved += 1

        return Response({'saved': saved, 'errors': errors})


# ── RESULTS EXCEL DOWNLOAD (admin only) ────────────────────────────────────────

@login_required
def download_ca_signoff(request):
    """Create a module CA acknowledgement sheet for student signatures."""
    module_id = request.GET.get('module_id')
    if not module_id:
        return HttpResponse('module_id is required.', status=400)

    try:
        module = (
            user_modules(request.user)
            .select_related('class_level', 'semester__academic_year')
            .get(pk=module_id)
        )
    except (Module.DoesNotExist, ValueError):
        return HttpResponseForbidden('You do not have access to this module.')

    students = list(
        Student.objects.filter(module=module)
        .select_related('result')
        .order_by('name', 'nactvet_reg_no')
    )

    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = Inches(0.55)
    section.left_margin = section.right_margin = Inches(0.55)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CONTINUOUS ASSESSMENT RESULTS ACKNOWLEDGEMENT')
    run.bold = True
    run.font.size = Pt(14)
    details = document.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run(
        f'{module.code} — {module.name}\n'
        f'{module.class_level.name} · {module.semester.label} · '
        f'{module.semester.academic_year.name}'
    ).bold = True
    document.add_paragraph(
        'Each student should verify the CA shown below, then sign in the Student Signature column.'
    )

    headers = [
        '#', 'NACTVET Reg. No.', 'Student Name',
        'A1 /100', 'A2 /100', 'CAT 1 /100', 'CAT 2 /100',
    ]
    if module.has_practical:
        headers += ['Practical 1 /100', 'Practical 2 /100']
    headers += ['Total CA Average /40', 'Student Signature']
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    serializer = StudentResultSerializer()

    def mark_text(result, field):
        if not result:
            return '—'
        if getattr(result, f'{field}_absent'):
            return 'ABS'
        value = getattr(result, field)
        return '—' if value is None else f'{value:.2f}'

    for number, student in enumerate(students, 1):
        result = getattr(student, 'result', None)
        total_ca = serializer.get_total_ca(result) if result else None
        row = table.add_row()
        values = [
            number, student.nactvet_reg_no, student.name,
            mark_text(result, 'assign1'), mark_text(result, 'assign2'),
            mark_text(result, 'cat1_theory'), mark_text(result, 'cat2_theory'),
        ]
        if module.has_practical:
            values += [
                mark_text(result, 'cat1_practical'), mark_text(result, 'cat2_practical'),
            ]
        values += ['—' if total_ca is None else f'{total_ca:.2f}', '']
        for index, value in enumerate(values):
            row.cells[index].text = str(value)
            row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row.cells[index].paragraphs[0].paragraph_format.space_after = Pt(0)
            row.cells[index].paragraphs[0].runs[0].font.size = Pt(7.5)
        row.height = Inches(0.42)

    widths = [0.3, 1.15, 1.7] + [0.62] * (len(headers) - 4) + [1.35]
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)

    document.add_paragraph()
    footer = document.add_paragraph('Tutor/Verifier Name: __________________________   Signature: __________________')
    footer.runs[0].bold = True

    safe_code = re.sub(r'[^A-Za-z0-9_-]+', '_', module.code).strip('_') or 'module'
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{safe_code}_ca_signoff.docx"'
    document.save(response)
    return response


@login_required
def download_results(request):
    if not can_read_exams(request.user):
        return HttpResponseForbidden('Administrator access required.')

    module_id       = request.GET.get('module_id')
    semester_id     = request.GET.get('semester_id')
    class_level_id  = request.GET.get('class_level_id')

    qs = (
        StudentResult.objects
        .filter(student__module__in=user_modules(request.user))
        .select_related(
            'student__module__class_level',
            'student__module__semester__academic_year',
        )
        .order_by('student__module__class_level__order', 'student__module__name', 'student__name')
    )
    if module_id:      qs = qs.filter(student__module_id=module_id)
    if semester_id:    qs = qs.filter(student__module__semester_id=semester_id)
    if class_level_id: qs = qs.filter(student__module__class_level_id=class_level_id)

    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = 'CA Results'

    HDR_FILL = PatternFill('solid', fgColor='1E2D78')
    HDR_FONT = Font(bold=True, color='FFFFFF', size=10)
    CENTER   = Alignment(horizontal='center', vertical='center', wrap_text=True)
    GREEN_F  = Font(bold=True, color='16A34A')   # text only
    RED_F    = Font(bold=True, color='DC2626')
    BLACK_F  = Font(color='000000')

    headers = [
        '#', 'NACTVET Reg No', 'Student Name', 'Module', 'Code', 'Level', 'Semester', 'Type',
        'A1 /100', 'A2 /100', 'CAT1-T /100', 'CAT2-T /100', 'P1 /100', 'P2 /100',
        'A1 wt', 'A2 wt', 'CAT1-T wt', 'CAT2-T wt', 'P1 wt', 'P2 wt',
        'Theory CA', 'Practical CA', 'Total CA /40',
        'T-Eligible', 'P-Eligible', 'CA Eligible',
    ]
    ws.append(headers)
    for ci in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=ci)
        cell.fill = HDR_FILL
        cell.font = HDR_FONT
        cell.alignment = CENTER
    ws.row_dimensions[1].height = 30

    def wt(raw, weight):
        return round(float(raw) / 100 * weight, 2) if raw is not None else None

    def yn(v):
        if v is None: return 'Pending'
        return 'YES' if v else 'NO'

    def fmt(v):
        return float(v) if v is not None else ''

    def nz(v):
        """Display substitution that keeps a real 0 distinct from 'not entered'."""
        return v if v is not None else ''

    for rn, res in enumerate(qs, 2):
        m  = res.student.module
        hp = m.has_practical
        a1, a2   = res.assign1,       res.assign2
        ct1, ct2 = res.cat1_theory,   res.cat2_theory
        cp1, cp2 = res.cat1_practical, res.cat2_practical

        if hp:
            a1w, a2w   = wt(a1, 2),   wt(a2, 2)
            ct1w, ct2w = wt(ct1, 8),  wt(ct2, 8)
            cp1w, cp2w = wt(cp1, 10), wt(cp2, 10)
            filled_t = [v for v in [a1w, a2w, ct1w, ct2w] if v is not None]
            filled_p = [v for v in [cp1w, cp2w]           if v is not None]
            t_ca     = round(sum(filled_t), 2) if filled_t else None
            p_ca     = round(sum(filled_p), 2) if filled_p else None
            tot      = round((t_ca or 0) + (p_ca or 0), 2) if (t_ca is not None or p_ca is not None) else None
            all_t    = all(v is not None for v in [a1, a2, ct1, ct2])
            all_p    = all(v is not None for v in [cp1, cp2])
            t_elig   = (t_ca >= 10) if (all_t and t_ca is not None) else None
            p_elig   = (p_ca >= 10) if (all_p and p_ca is not None) else None
            ca_elig  = (t_elig and p_elig) if (t_elig is not None and p_elig is not None) else None
        else:
            a1w, a2w   = wt(a1, 5),  wt(a2, 5)
            ct1w, ct2w = wt(ct1, 15), wt(ct2, 15)
            cp1w = cp2w = p_ca = None
            filled_t = [v for v in [a1w, a2w, ct1w, ct2w] if v is not None]
            t_ca     = round(sum(filled_t), 2) if filled_t else None
            tot      = t_ca
            all_done = all(v is not None for v in [a1, a2, ct1, ct2])
            t_elig = p_elig = None
            ca_elig  = (t_ca >= 20) if (all_done and t_ca is not None) else None

        row = [
            rn - 1,
            res.student.nactvet_reg_no, res.student.name,
            m.name, m.code, m.class_level.name, m.semester.label,
            'Theory + Practical' if hp else 'Theory Only',
            fmt(a1), fmt(a2), fmt(ct1), fmt(ct2), fmt(cp1), fmt(cp2),
            nz(a1w), nz(a2w), nz(ct1w), nz(ct2w), nz(cp1w), nz(cp2w),
            nz(t_ca), nz(p_ca), nz(tot),
            yn(t_elig) if hp else 'N/A',
            yn(p_elig) if hp else 'N/A',
            yn(ca_elig),
        ]
        ws.append(row)
        # Color only the status cell text; all fills remain white
        elig_col = len(headers)
        status_font = GREEN_F if ca_elig is True else (RED_F if ca_elig is False else BLACK_F)
        ws.cell(row=rn, column=elig_col).font = status_font

    for col in ws.columns:
        width = max(len(str(cell.value or '')) for cell in col)
        ws.column_dimensions[get_column_letter(col[0].column)].width = min(width + 3, 28)

    fname = f'ca_results_{timezone.localdate()}.xlsx'
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="{fname}"'
    wb.save(response)
    return response


# ── FIELD RESULTS EXCEL DOWNLOAD (admin only) ──────────────────────────────────

@login_required
def download_field_results(request):
    """Export one field module in a compact CA 40% + report 60% workbook."""
    if not can_read_exams(request.user):
        return HttpResponseForbidden('Administrator access required.')

    module_id = request.GET.get('module_id')
    if not module_id:
        return HttpResponse('module_id is required.', status=400)
    try:
        module = (
            user_modules(request.user)
            .select_related('class_level', 'semester__academic_year')
            .get(pk=module_id, is_field_module=True)
        )
    except (Module.DoesNotExist, ValueError):
        return HttpResponseForbidden('Select a valid field results module.')

    students = (
        Student.objects.filter(module=module)
        .select_related('result')
        .order_by('name', 'nactvet_reg_no')
    )
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = 'Field Results'
    navy, blue = '1E2D78', 'D9EAF7'
    thin = Side(style='thin', color='808080')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)

    title_rows = (
        'BLUE PHARMA COLLEGE OF HEALTH', 'FIELD RESULTS',
        f'{module.code} — {module.name}',
        f'{module.class_level.name} · {module.semester.label} · {module.semester.academic_year.name}',
    )
    for row_number, value in enumerate(title_rows, 1):
        ws.merge_cells(start_row=row_number, start_column=1, end_row=row_number, end_column=9)
        cell = ws.cell(row=row_number, column=1, value=value)
        cell.font = Font(bold=True, size=14 if row_number == 1 else 11, color=navy)
        cell.alignment = center

    headers = (
        'SN', 'NACTVET Registration Number', 'Student Name',
        'PPB/LOGBOOK', 'AV', 'REPORT', 'AV', 'TOTAL', 'STATUS',
    )
    for column, heading in enumerate(headers, 1):
        cell = ws.cell(row=6, column=column, value=heading)
        cell.fill = PatternFill('solid', fgColor=navy)
        cell.font = Font(bold=True, color='FFFFFF')
        cell.alignment = center
        cell.border = border
    for column, heading in enumerate(('Raw /100', '40%', 'Raw /100', '60%'), 4):
        cell = ws.cell(row=5, column=column, value=heading)
        cell.fill = PatternFill('solid', fgColor=blue)
        cell.font = Font(bold=True, color='000000')
        cell.alignment = center
        cell.border = border

    serializer = StudentResultSerializer()
    for number, student in enumerate(students, 1):
        result = getattr(student, 'result', None)
        values = (
            number, student.nactvet_reg_no, student.name,
            result.field_ca if result else None,
            serializer.get_total_ca(result) if result else None,
            result.end_theory if result else None,
            serializer.get_end_theory_w(result) if result else None,
            serializer.get_final_total(result) if result else None,
            serializer.get_result_status(result) if result else 'INCOMPLETE',
        )
        for column, value in enumerate(values, 1):
            cell = ws.cell(row=6 + number, column=column, value='' if value is None else value)
            cell.alignment = Alignment(horizontal='left' if column == 3 else 'center', vertical='center')
            cell.border = border

    for column, width in enumerate((7, 26, 28, 16, 11, 14, 11, 11, 16), 1):
        ws.column_dimensions[get_column_letter(column)].width = width
    ws.freeze_panes = 'D7'
    ws.auto_filter.ref = f'A6:I{max(ws.max_row, 6)}'
    ws.sheet_view.showGridLines = False
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.fitToWidth = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    safe_code = re.sub(r'[^A-Za-z0-9_-]+', '_', module.code).strip('_') or 'field'
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = (
        f'attachment; filename="{safe_code}_field_results_{timezone.localdate()}.xlsx"'
    )
    wb.save(response)
    return response


# ── FINAL RESULTS EXCEL DOWNLOAD (admin only) ──────────────────────────────────

@login_required
def download_final_results(request):
    if not can_read_exams(request.user):
        return HttpResponseForbidden('Administrator access required.')

    module_id      = request.GET.get('module_id')
    semester_id    = request.GET.get('semester_id')
    class_level_id = request.GET.get('class_level_id')

    qs = (
        StudentResult.objects
        .filter(student__module__in=user_modules(request.user))
        .select_related(
            'student__module__class_level',
            'student__module__semester__academic_year',
        )
        .order_by('student__module__class_level__order', 'student__module__name', 'student__name')
    )
    if module_id:      qs = qs.filter(student__module_id=module_id)
    if semester_id:    qs = qs.filter(student__module__semester_id=semester_id)
    if class_level_id: qs = qs.filter(student__module__class_level_id=class_level_id)

    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = 'Final Results'

    HDR_FILL = PatternFill('solid', fgColor='1E2D78')
    HDR_FONT = Font(bold=True, color='FFFFFF', size=10)
    CENTER   = Alignment(horizontal='center', vertical='center', wrap_text=True)
    GREEN_F  = Font(bold=True, color='16A34A')
    RED_F    = Font(bold=True, color='DC2626')
    BLACK_F  = Font(color='000000')

    headers = [
        '#', 'NACTVET Reg No', 'Student Name', 'Module', 'Code', 'Level', 'Semester', 'Type',
        # CA marks
        'A1 /100', 'A2 /100', 'CAT1-T /100', 'CAT2-T /100', 'P1 /100', 'P2 /100',
        'Theory CA /20(or/40)', 'Practical CA /20', 'Total CA /40', 'CA Eligible',
        # End of semester
        'End Theory /100', 'End Practical /100',
        'End Theory wt', 'End Practical wt',
        'End Exam Total',
        # Grand total
        'Final Total /100', 'Supplementary /100', 'Grade', 'Grade Point', 'Result Status',
    ]
    ws.append(headers)
    for ci in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=ci)
        cell.fill = HDR_FILL
        cell.font = HDR_FONT
        cell.alignment = CENTER
    ws.row_dimensions[1].height = 36

    def wt(raw, weight):
        return round(float(raw) / 100 * weight, 2) if raw is not None else None

    def yn(v):
        if v is None: return 'Pending'
        return 'YES' if v else 'NO'

    def fmt(v):
        return float(v) if v is not None else ''

    def nz(v):
        """Display substitution that keeps a real 0 distinct from 'not entered'."""
        return v if v is not None else ''

    for rn, res in enumerate(qs, 2):
        m  = res.student.module
        hp = m.has_practical
        a1, a2   = res.assign1,       res.assign2
        ct1, ct2 = res.cat1_theory,   res.cat2_theory
        cp1, cp2 = res.cat1_practical, res.cat2_practical
        et       = res.end_theory
        ep       = res.end_practical

        # CA
        if hp:
            a1w, a2w   = wt(a1, 2),   wt(a2, 2)
            ct1w, ct2w = wt(ct1, 8),  wt(ct2, 8)
            cp1w, cp2w = wt(cp1, 10), wt(cp2, 10)
            filled_t   = [v for v in [a1w, a2w, ct1w, ct2w] if v is not None]
            filled_p   = [v for v in [cp1w, cp2w]           if v is not None]
            t_ca       = round(sum(filled_t), 2) if filled_t else None
            p_ca       = round(sum(filled_p), 2) if filled_p else None
            tot_ca     = round((t_ca or 0) + (p_ca or 0), 2) if (t_ca is not None or p_ca is not None) else None
            all_t      = all(v is not None for v in [a1, a2, ct1, ct2])
            all_p      = all(v is not None for v in [cp1, cp2])
            t_elig     = (t_ca >= 10) if (all_t and t_ca is not None) else None
            p_elig     = (p_ca >= 10) if (all_p and p_ca is not None) else None
            ca_elig    = (t_elig and p_elig) if (t_elig is not None and p_elig is not None) else None
            # End exam
            etw  = wt(et, 30)
            epw  = wt(ep, 30)
        else:
            a1w, a2w   = wt(a1, 5),  wt(a2, 5)
            ct1w, ct2w = wt(ct1, 15), wt(ct2, 15)
            cp1w = cp2w = p_ca = None
            filled_t   = [v for v in [a1w, a2w, ct1w, ct2w] if v is not None]
            t_ca       = round(sum(filled_t), 2) if filled_t else None
            tot_ca     = t_ca
            all_done   = all(v is not None for v in [a1, a2, ct1, ct2])
            t_elig = p_elig = None
            ca_elig    = (t_ca >= 20) if (all_done and t_ca is not None) else None
            # End exam
            etw  = wt(et, 60)
            epw  = None

        end_exam_total = round((etw or 0) + (epw or 0), 2) if (etw is not None or epw is not None) else None
        final = round((tot_ca or 0) + (end_exam_total or 0), 2) if (tot_ca is not None or end_exam_total is not None) else None
        outcome = StudentResultSerializer(res).data
        pass_fail = outcome['result_status']

        row = [
            rn - 1,
            res.student.nactvet_reg_no, res.student.name,
            m.name, m.code, m.class_level.name, m.semester.label,
            'Theory + Practical' if hp else 'Theory Only',
            fmt(a1), fmt(a2), fmt(ct1), fmt(ct2), fmt(cp1), fmt(cp2),
            nz(t_ca), nz(p_ca), nz(tot_ca), yn(ca_elig),
            fmt(et), fmt(ep) if hp else 'N/A',
            nz(etw), nz(epw) if hp else 'N/A',
            nz(end_exam_total),
            nz(final), fmt(res.supplementary_mark), outcome['grade'] or '',
            outcome['grade_point'] if outcome['grade_point'] is not None else '',
            pass_fail,
        ]
        ws.append(row)

        # Color only the Pass/Fail cell text; all fills remain white
        pf_col = len(headers)
        ws.cell(row=rn, column=pf_col).font = GREEN_F if pass_fail == 'PASS' else (RED_F if pass_fail == 'FAIL' else BLACK_F)

    for col in ws.columns:
        width = max(len(str(cell.value or '')) for cell in col)
        ws.column_dimensions[get_column_letter(col[0].column)].width = min(width + 3, 30)

    fname = f'final_results_{timezone.localdate()}.xlsx'
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="{fname}"'
    wb.save(response)
    return response


# ── ELIGIBILITY EXCEL DOWNLOAD (admin only) ────────────────────────────────────

@login_required
def download_eligibility_excel(request):
    if not can_read_exams(request.user):
        return HttpResponseForbidden('Administrator access required.')

    module_id      = request.GET.get('module_id')
    semester_id    = request.GET.get('semester_id')
    class_level_id = request.GET.get('class_level_id')
    basis           = request.GET.get('basis', 'all')

    my_modules = user_modules(request.user)
    students = (
        Student.objects.filter(module__in=my_modules)
        .select_related('module__class_level', 'module__semester__academic_year', 'result')
        .prefetch_related('attendance_records__session')
        .order_by('module__class_level__order', 'module__name', 'name')
    )
    if module_id:      students = students.filter(module_id=module_id)
    if semester_id:    students = students.filter(module__semester_id=semester_id)
    if class_level_id: students = students.filter(module__class_level_id=class_level_id)

    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = 'Eligibility'

    HDR_FILL = PatternFill('solid', fgColor='1E2D78')
    HDR_FONT = Font(bold=True, color='FFFFFF', size=10)
    CENTER   = Alignment(horizontal='center', vertical='center', wrap_text=True)
    GREEN_F  = Font(bold=True, color='16A34A')
    RED_F    = Font(bold=True, color='DC2626')
    BLACK_F  = Font(color='000000')

    if basis == 'performance':
        from collections import defaultdict

        enrollments = list(students)
        levels = defaultdict(list)
        for student in enrollments:
            levels[student.module.class_level.name].append(student)

        wb.remove(ws)
        used_titles = set()
        for level_name, level_students in levels.items():
            base_title = ''.join(c for c in level_name if c not in '[]:*?/\\')[:31] or 'Class'
            title = base_title
            suffix = 2
            while title in used_titles:
                title = f'{base_title[:27]} {suffix}'
                suffix += 1
            used_titles.add(title)
            level_ws = wb.create_sheet(title)
            module_codes = sorted({student.module.code for student in level_students})
            headers = ['Student Reg. No.', 'Name', *module_codes]
            level_ws.append(headers)
            for cell in level_ws[1]:
                cell.fill = HDR_FILL
                cell.font = HDR_FONT
                cell.alignment = CENTER

            student_rows = defaultdict(dict)
            student_names = {}
            module_summary = {code: {'eligible': 0, 'ineligible': 0} for code in module_codes}
            for student in level_students:
                _total_ca, eligible, _note = ca_eligibility_for_student(student)
                is_eligible = eligible is True
                student_names[student.nactvet_reg_no] = student.name
                student_rows[student.nactvet_reg_no][student.module.code] = is_eligible
                module_summary[student.module.code]['eligible' if is_eligible else 'ineligible'] += 1

            for reg_no in sorted(student_rows, key=lambda key: student_names[key].lower()):
                values = [reg_no, student_names[reg_no]]
                for code in module_codes:
                    values.append(
                        'Eligible' if student_rows[reg_no].get(code) is True
                        else ('Ineligible' if code in student_rows[reg_no] else '')
                    )
                level_ws.append(values)
                row_number = level_ws.max_row
                for column, code in enumerate(module_codes, 3):
                    cell = level_ws.cell(row=row_number, column=column)
                    if cell.value == 'Eligible':
                        cell.font = GREEN_F
                    elif cell.value == 'Ineligible':
                        cell.font = RED_F
                    cell.alignment = CENTER

            summary_start = level_ws.max_row + 3
            level_ws.cell(row=summary_start, column=1, value='Summary by Module').font = Font(bold=True, size=12)
            summary_headers = ['Module Code', 'Eligible', 'Ineligible', 'Total']
            for column, value in enumerate(summary_headers, 1):
                cell = level_ws.cell(row=summary_start + 1, column=column, value=value)
                cell.fill = HDR_FILL
                cell.font = HDR_FONT
                cell.alignment = CENTER

            class_eligible = class_ineligible = 0
            for offset, code in enumerate(module_codes, summary_start + 2):
                counts = module_summary[code]
                class_eligible += counts['eligible']
                class_ineligible += counts['ineligible']
                level_ws.append([code, counts['eligible'], counts['ineligible'], counts['eligible'] + counts['ineligible']])
                level_ws.cell(row=offset, column=2).font = GREEN_F
                level_ws.cell(row=offset, column=3).font = RED_F

            class_row = level_ws.max_row + 2
            level_ws.cell(row=class_row, column=1, value=f'{level_name} Total').font = Font(bold=True)
            level_ws.cell(row=class_row, column=2, value=class_eligible).font = GREEN_F
            level_ws.cell(row=class_row, column=3, value=class_ineligible).font = RED_F
            level_ws.cell(row=class_row, column=4, value=class_eligible + class_ineligible).font = Font(bold=True)
            level_ws.freeze_panes = 'C2'

            for column_cells in level_ws.columns:
                width = max(len(str(cell.value or '')) for cell in column_cells)
                level_ws.column_dimensions[get_column_letter(column_cells[0].column)].width = min(max(width + 3, 12), 28)

        if not wb.worksheets:
            empty_ws = wb.create_sheet('Eligibility')
            empty_ws.append(['No eligibility records found'])

        fname = f'module_eligibility_{timezone.localdate()}.xlsx'
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{fname}"'
        wb.save(response)
        return response

    headers = [
        '#', 'NACTVET Reg No', 'Student Name', 'Module', 'Code', 'Level', 'Semester',
        'CAT1 Sessions', 'CAT1 Attended', 'CAT1 %', 'CAT1 Eligible',
        'CAT2 Sessions', 'CAT2 Attended', 'CAT2 %', 'CAT2 Eligible',
        'End Sessions', 'End Attended', 'End %', 'End Eligible',
    ]
    ws.append(headers)
    for ci in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=ci)
        cell.fill = HDR_FILL
        cell.font = HDR_FONT
        cell.alignment = CENTER
    ws.row_dimensions[1].height = 30

    _mod_cache = {}
    def _period_counts(mid):
        if mid not in _mod_cache:
            cat1  = Session.objects.filter(module_id=mid, exam_period=Session.CAT1).count()
            cat2  = Session.objects.filter(module_id=mid, exam_period=Session.CAT2).count()
            total = Session.objects.filter(module_id=mid).count()
            _mod_cache[mid] = {'cat1': cat1, 'cat2': cat2, 'total': total}
        return _mod_cache[mid]

    def yn(v):
        if v is None: return 'N/A'
        return 'YES' if v else 'NO'

    for rn, st in enumerate(students, 2):
        mc   = _period_counts(st.module_id)
        recs = list(st.attendance_records.all())

        cat1_eff  = sum(1 for r in recs if r.session.exam_period == Session.CAT1 and attendance_is_effective(r))
        cat2_eff  = sum(1 for r in recs if r.session.exam_period == Session.CAT2 and attendance_is_effective(r))
        total_eff = sum(1 for r in recs if attendance_is_effective(r))

        cat1_pct  = round((cat1_eff / mc['cat1']) * 100) if mc['cat1'] else None
        cat2_pct  = round((cat2_eff / mc['cat2']) * 100) if mc['cat2'] else None
        end_pct   = round((total_eff / mc['total']) * 100) if mc['total'] else None

        cat1_el = (cat1_pct >= ELIGIBILITY_THRESHOLD) if mc['cat1'] else None
        cat2_el = (cat2_pct >= ELIGIBILITY_THRESHOLD) if mc['cat2'] else None
        end_el  = (end_pct  >= ELIGIBILITY_THRESHOLD) if mc['total'] else None

        row = [
            rn - 1,
            st.nactvet_reg_no, st.name,
            st.module.name, st.module.code, st.module.class_level.name, st.module.semester.label,
            mc['cat1'], cat1_eff, cat1_pct if cat1_pct is not None else '', yn(cat1_el),
            mc['cat2'], cat2_eff, cat2_pct if cat2_pct is not None else '', yn(cat2_el),
            mc['total'], total_eff, end_pct if end_pct is not None else '', yn(end_el),
        ]
        ws.append(row)

        # Color only the three eligibility text cells
        for col_offset, elig in [(11, cat1_el), (15, cat2_el), (19, end_el)]:
            ws.cell(row=rn, column=col_offset).font = (
                GREEN_F if elig is True else (RED_F if elig is False else BLACK_F)
            )

    for col in ws.columns:
        width = max(len(str(cell.value or '')) for cell in col)
        ws.column_dimensions[get_column_letter(col[0].column)].width = min(width + 3, 24)

    fname = f'eligibility_{timezone.localdate()}.xlsx'
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="{fname}"'
    wb.save(response)
    return response


@login_required
def download_final_eligibility_excel(request):
    """Export attendance, CA marks and combined eligibility for NTA 4–6."""
    if not can_read_exams(request.user):
        return HttpResponseForbidden('Administrator access required.')

    from collections import defaultdict
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    semester_id = request.GET.get('semester_id')
    scope = request.GET.get('scope', '').strip().lower()
    modules = (
        user_modules(request.user)
        .filter(class_level__order__in=(4, 5, 6))
        .select_related('class_level', 'semester__academic_year')
        .order_by('class_level__order', 'code')
    )
    if semester_id:
        modules = modules.filter(semester_id=semester_id)
    if scope == 'semester_1':
        modules = modules.filter(semester__number=Semester.SEM1)
    elif scope == 'semester_2':
        modules = modules.filter(semester__number=Semester.SEM2)
    elif scope == 'field':
        modules = modules.filter(is_field_module=True)
    modules = list(modules)

    enrollments = list(
        Student.objects.filter(module__in=modules)
        .select_related('module', 'result')
        .prefetch_related('attendance_records__session')
        .order_by('name', 'nactvet_reg_no', 'module__code')
    )

    module_session_counts = {
        module.id: Session.objects.filter(module=module).count()
        for module in modules
    }
    enrollment_data = {}
    students_by_level = defaultdict(dict)
    for student in enrollments:
        records = list(student.attendance_records.all())
        total_sessions = module_session_counts[student.module_id]
        attended = sum(1 for record in records if attendance_is_effective(record))
        attendance_pct = round(attended / total_sessions * 100, 1) if total_sessions else None
        ca_total, ca_eligible, _ca_note = ca_eligibility_for_student(student)
        attendance_eligible = (
            attendance_pct >= ELIGIBILITY_THRESHOLD
            if attendance_pct is not None else None
        )
        final_eligible = (
            None if ca_eligible is None or attendance_eligible is None
            else ca_eligible and attendance_eligible
        )
        enrollment_data[(student.nactvet_reg_no, student.module_id)] = {
            'student': student,
            'sessions': total_sessions,
            'attended': attended,
            'attendance_pct': attendance_pct,
            'ca_total': ca_total,
            'ca_eligible': ca_eligible,
            'final_eligible': final_eligible,
        }
        students_by_level[student.module.class_level.order][student.nactvet_reg_no] = student.name

    level_modules = defaultdict(list)
    for module in modules:
        level_modules[module.class_level.order].append(module)

    navy = '1E2D78'
    blue = 'D9EAF7'
    green = 'C6EFCE'
    green_text = '006100'
    red = 'FFC7CE'
    red_text = '9C0006'
    yellow = 'FFF2CC'
    yellow_text = '7F6000'
    white = 'FFFFFF'
    thin = Side(style='thin', color='808080')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)

    def style_header(cell, fill=navy):
        cell.fill = PatternFill('solid', fgColor=fill)
        cell.font = Font(bold=True, color=white)
        cell.alignment = center
        cell.border = border

    def style_status(cell, status):
        if status is True or status == 'Eligible':
            cell.fill = PatternFill('solid', fgColor=green)
            cell.font = Font(bold=True, color=green_text)
        elif status is False or status == 'Ineligible':
            cell.fill = PatternFill('solid', fgColor=red)
            cell.font = Font(bold=True, color=red_text)
        else:
            cell.fill = PatternFill('solid', fgColor=yellow)
            cell.font = Font(bold=True, color=yellow_text)
        cell.alignment = center
        cell.border = border

    def finish_sheet(ws):
        ws.freeze_panes = 'D8'
        ws.auto_filter.ref = f'A7:{get_column_letter(ws.max_column)}{ws.max_row}'
        ws.sheet_view.showGridLines = False
        ws.page_setup.orientation = 'landscape'
        ws.page_setup.fitToWidth = 1
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        for column in range(1, ws.max_column + 1):
            width = max(len(str(ws.cell(row=row, column=column).value or '')) for row in range(1, ws.max_row + 1))
            ws.column_dimensions[get_column_letter(column)].width = min(max(width + 2, 11), 24)
        ws.column_dimensions['B'].width = max(ws.column_dimensions['B'].width, 24)
        ws.column_dimensions['C'].width = max(ws.column_dimensions['C'].width, 23)

    def prepare_sheet(ws, level, title, last_column):
        last_letter = get_column_letter(max(last_column, 3))
        headings = [
            'BLUE PHARMA COLLEGE OF HEALTH',
            'DEPARTMENT OF PHARMACEUTICAL SCIENCES',
            f'NTA LEVEL {level}',
            title,
        ]
        for row, heading in enumerate(headings, 1):
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=max(last_column, 3))
            cell = ws.cell(row=row, column=1, value=heading)
            cell.font = Font(bold=True, size=14 if row == 1 else 11, color=navy)
            cell.alignment = center
        for column, value in enumerate(('SN', 'Name', 'NACTVET Registration Number'), 1):
            style_header(ws.cell(row=5, column=column, value=value))
            ws.merge_cells(start_row=5, start_column=column, end_row=7, end_column=column)
        ws.row_dimensions[5].height = 36
        ws.print_title_rows = '1:7'
        ws.print_area = f'A1:{last_letter}{max(ws.max_row, 5)}'

    def result_values(data):
        result = getattr(data['student'], 'result', None) if data else None
        if result is None:
            return result, ('', '', '', ''), '', '', ''
        serializer = StudentResultSerializer()
        theory = serializer.get_theory_ca(result)
        practical = serializer.get_practical_ca(result)
        return result, (
            result.cat1_theory, result.cat2_theory,
            result.assign1, result.assign2,
        ), theory, practical, data['ca_total']

    def field_result_values(data):
        result = getattr(data['student'], 'result', None) if data else None
        if result is None:
            return '', '', '', ''
        serializer = StudentResultSerializer()
        return (
            result.field_ca,
            serializer.get_total_ca(result),
            result.end_theory,
            serializer.get_end_theory_w(result),
        )

    def write_mark(cell, value, pass_mark=50):
        cell.value = value if value is not None else ''
        cell.alignment = center
        cell.border = border
        if value is not None and value != '':
            style_status(cell, float(value) >= pass_mark)

    wb = Workbook()
    wb.remove(wb.active)

    for level in (4, 5, 6):
        level_module_list = level_modules[level]
        student_map = students_by_level[level]

        assignments_ws = wb.create_sheet(f'NTA {level} Assignments')
        assignment_width = 2 * len(level_module_list)
        prepare_sheet(assignments_ws, level, 'ASSIGNMENT RESULTS', 3 + assignment_width)
        assignment_starts = {}
        column = 4
        for module in level_module_list:
            assignment_starts[module.id] = column
            assignments_ws.merge_cells(start_row=5, start_column=column, end_row=5, end_column=column + 1)
            style_header(assignments_ws.cell(row=5, column=column, value=f'{module.name} ({module.code})'))
            assignments_ws.merge_cells(start_row=6, start_column=column, end_row=6, end_column=column + 1)
            style_header(assignments_ws.cell(row=6, column=column, value='Assignments'), blue)
            assignments_ws.cell(row=6, column=column).font = Font(bold=True, color='000000')
            for offset, heading in enumerate(('ASS1', 'ASS2')):
                style_header(assignments_ws.cell(row=7, column=column + offset, value=heading), blue)
                assignments_ws.cell(row=7, column=column + offset).font = Font(bold=True, color='000000')
            column += 2
        for row_number, (reg_no, name) in enumerate(sorted(student_map.items(), key=lambda item: item[1].lower()), 8):
            assignments_ws.append([row_number - 7, name, reg_no])
            for module in level_module_list:
                data = enrollment_data.get((reg_no, module.id))
                result = getattr(data['student'], 'result', None) if data else None
                a1 = result.assign1 if result else None
                a2 = result.assign2 if result else None
                start = assignment_starts[module.id]
                for offset, value in enumerate((a1, a2)):
                    write_mark(assignments_ws.cell(row=row_number, column=start + offset), value)
        finish_sheet(assignments_ws)

        cats_ws = wb.create_sheet(f'NTA {level} CATs')
        cat_width = sum(
            2 if module.is_field_module else (9 if module.has_practical else 5)
            for module in level_module_list
        )
        prepare_sheet(cats_ws, level, 'CONTINUOUS ASSESSMENT RESULTS', 3 + cat_width)
        cat_starts = {}
        column = 4
        for module in level_module_list:
            width = 2 if module.is_field_module else (9 if module.has_practical else 5)
            cat_starts[module.id] = column
            cats_ws.merge_cells(start_row=5, start_column=column, end_row=5, end_column=column + width - 1)
            style_header(cats_ws.cell(row=5, column=column, value=f'{module.name} ({module.code})'))
            if module.is_field_module:
                for offset, heading in enumerate(('PPB/LOGBOOK', 'AV')):
                    style_header(cats_ws.cell(row=6, column=column + offset, value=heading), blue)
                    cats_ws.cell(row=6, column=column + offset).font = Font(bold=True, color='000000')
                    cats_ws.merge_cells(start_row=6, start_column=column + offset, end_row=7, end_column=column + offset)
                column += width
                continue
            cats_ws.merge_cells(start_row=6, start_column=column, end_row=6, end_column=column + 4)
            style_header(cats_ws.cell(row=6, column=column, value='Theory'), blue)
            cats_ws.cell(row=6, column=column).font = Font(bold=True, color='000000')
            for offset, heading in enumerate(('WR1', 'WR2', 'AS1', 'ASS2', 'AV')):
                style_header(cats_ws.cell(row=7, column=column + offset, value=heading), blue)
                cats_ws.cell(row=7, column=column + offset).font = Font(bold=True, color='000000')
            if module.has_practical:
                cats_ws.merge_cells(start_row=6, start_column=column + 5, end_row=6, end_column=column + 7)
                style_header(cats_ws.cell(row=6, column=column + 5, value='Practical'), blue)
                cats_ws.cell(row=6, column=column + 5).font = Font(bold=True, color='000000')
                for offset, heading in enumerate(('PRAC1', 'PRAC2', 'AV'), 5):
                    style_header(cats_ws.cell(row=7, column=column + offset, value=heading), blue)
                    cats_ws.cell(row=7, column=column + offset).font = Font(bold=True, color='000000')
                style_header(cats_ws.cell(row=6, column=column + 8, value='Total'), blue)
                cats_ws.cell(row=6, column=column + 8).font = Font(bold=True, color='000000')
                style_header(cats_ws.cell(row=7, column=column + 8, value='AV'), blue)
                cats_ws.cell(row=7, column=column + 8).font = Font(bold=True, color='000000')
            column += width
        for row_number, (reg_no, name) in enumerate(sorted(student_map.items(), key=lambda item: item[1].lower()), 8):
            cats_ws.append([row_number - 7, name, reg_no])
            for module in level_module_list:
                data = enrollment_data.get((reg_no, module.id))
                if module.is_field_module:
                    field_ca, field_ca_av, _report, _report_av = field_result_values(data)
                    start = cat_starts[module.id]
                    write_mark(cats_ws.cell(row=row_number, column=start), field_ca)
                    write_mark(cats_ws.cell(row=row_number, column=start + 1), field_ca_av, 20)
                    continue
                result, theory_raw, theory_av, practical_av, total_av = result_values(data)
                start = cat_starts[module.id]
                for offset, value in enumerate(theory_raw):
                    write_mark(cats_ws.cell(row=row_number, column=start + offset), value)
                write_mark(cats_ws.cell(row=row_number, column=start + 4), theory_av, 10 if module.has_practical else 20)
                if module.has_practical:
                    practical_raw = (result.cat1_practical, result.cat2_practical) if result else ('', '')
                    write_mark(cats_ws.cell(row=row_number, column=start + 5), practical_raw[0])
                    write_mark(cats_ws.cell(row=row_number, column=start + 6), practical_raw[1])
                    write_mark(cats_ws.cell(row=row_number, column=start + 7), practical_av, 10)
                    write_mark(cats_ws.cell(row=row_number, column=start + 8), total_av, CA_ELIGIBILITY_THRESHOLD)
        finish_sheet(cats_ws)

        eligibility_ws = wb.create_sheet(f'NTA {level} Eligibility')
        eligibility_width = sum(
            6 if module.is_field_module else (11 if module.has_practical else 7)
            for module in level_module_list
        )
        prepare_sheet(eligibility_ws, level, 'FINAL ELIGIBILITY TO END-OF-SEMESTER EXAMINATION', 3 + eligibility_width)
        eligibility_starts = {}
        column = 4
        for module in level_module_list:
            if module.is_field_module:
                mark_width = 4
                width = 6
                eligibility_starts[module.id] = column
                eligibility_ws.merge_cells(start_row=5, start_column=column, end_row=5, end_column=column + width - 1)
                style_header(eligibility_ws.cell(row=5, column=column, value=f'{module.name} ({module.code})'))
                for offset, heading in enumerate(('PPB/LOGBOOK', 'AV', 'REPORT', 'AV')):
                    style_header(eligibility_ws.cell(row=6, column=column + offset, value=heading), blue)
                    eligibility_ws.cell(row=6, column=column + offset).font = Font(bold=True, color='000000')
                    eligibility_ws.merge_cells(start_row=6, start_column=column + offset, end_row=7, end_column=column + offset)
                for offset, heading in enumerate(('Attendance (%)', 'Final Eligibility'), mark_width):
                    style_header(eligibility_ws.cell(row=6, column=column + offset, value=heading), blue)
                    eligibility_ws.cell(row=6, column=column + offset).font = Font(bold=True, color='000000')
                    eligibility_ws.merge_cells(start_row=6, start_column=column + offset, end_row=7, end_column=column + offset)
                column += width
                continue
            mark_width = 9 if module.has_practical else 5
            width = mark_width + 2
            eligibility_starts[module.id] = column
            eligibility_ws.merge_cells(start_row=5, start_column=column, end_row=5, end_column=column + width - 1)
            style_header(eligibility_ws.cell(row=5, column=column, value=f'{module.name} ({module.code})'))
            eligibility_ws.merge_cells(start_row=6, start_column=column, end_row=6, end_column=column + 4)
            style_header(eligibility_ws.cell(row=6, column=column, value='Theory'), blue)
            eligibility_ws.cell(row=6, column=column).font = Font(bold=True, color='000000')
            for offset, heading in enumerate(('WR1', 'WR2', 'AS1', 'ASS2', 'AV')):
                style_header(eligibility_ws.cell(row=7, column=column + offset, value=heading), blue)
                eligibility_ws.cell(row=7, column=column + offset).font = Font(bold=True, color='000000')
            if module.has_practical:
                eligibility_ws.merge_cells(start_row=6, start_column=column + 5, end_row=6, end_column=column + 7)
                style_header(eligibility_ws.cell(row=6, column=column + 5, value='Practical'), blue)
                eligibility_ws.cell(row=6, column=column + 5).font = Font(bold=True, color='000000')
                for offset, heading in enumerate(('PRAC1', 'PRAC2', 'AV'), 5):
                    style_header(eligibility_ws.cell(row=7, column=column + offset, value=heading), blue)
                    eligibility_ws.cell(row=7, column=column + offset).font = Font(bold=True, color='000000')
                style_header(eligibility_ws.cell(row=6, column=column + 8, value='Total'), blue)
                eligibility_ws.cell(row=6, column=column + 8).font = Font(bold=True, color='000000')
                style_header(eligibility_ws.cell(row=7, column=column + 8, value='AV'), blue)
                eligibility_ws.cell(row=7, column=column + 8).font = Font(bold=True, color='000000')
            for offset, heading in enumerate(('Attendance (%)', 'Final Eligibility'), mark_width):
                style_header(eligibility_ws.cell(row=6, column=column + offset, value=heading), blue)
                eligibility_ws.cell(row=6, column=column + offset).font = Font(bold=True, color='000000')
                eligibility_ws.merge_cells(start_row=6, start_column=column + offset, end_row=7, end_column=column + offset)
            column += width
        for row_number, (reg_no, name) in enumerate(sorted(student_map.items(), key=lambda item: item[1].lower()), 8):
            eligibility_ws.append([row_number - 7, name, reg_no])
            for module in level_module_list:
                data = enrollment_data.get((reg_no, module.id))
                start = eligibility_starts[module.id]
                if module.is_field_module:
                    mark_width = 4
                    field_marks = field_result_values(data)
                    for offset, value in enumerate(field_marks):
                        pass_mark = 20 if offset == 1 else (30 if offset == 3 else 50)
                        write_mark(eligibility_ws.cell(row=row_number, column=start + offset), value, pass_mark)
                else:
                    result, theory_raw, theory_av, practical_av, total_av = result_values(data)
                    for offset, value in enumerate(theory_raw):
                        write_mark(eligibility_ws.cell(row=row_number, column=start + offset), value)
                    write_mark(eligibility_ws.cell(row=row_number, column=start + 4), theory_av, 10 if module.has_practical else 20)
                    mark_width = 5
                    if module.has_practical:
                        practical_raw = (result.cat1_practical, result.cat2_practical) if result else ('', '')
                        write_mark(eligibility_ws.cell(row=row_number, column=start + 5), practical_raw[0])
                        write_mark(eligibility_ws.cell(row=row_number, column=start + 6), practical_raw[1])
                        write_mark(eligibility_ws.cell(row=row_number, column=start + 7), practical_av, 10)
                        write_mark(eligibility_ws.cell(row=row_number, column=start + 8), total_av, CA_ELIGIBILITY_THRESHOLD)
                        mark_width = 9
                attendance_cell = eligibility_ws.cell(row=row_number, column=start + mark_width, value=data['attendance_pct'] if data and data['attendance_pct'] is not None else '')
                style_status(attendance_cell, None if not data or data['attendance_pct'] is None else data['attendance_pct'] >= ELIGIBILITY_THRESHOLD)
                status = 'INCOMPLETE' if not data or data['final_eligible'] is None else ('Eligible' if data['final_eligible'] else 'Ineligible')
                eligibility_cell = eligibility_ws.cell(row=row_number, column=start + mark_width + 1, value=status)
                style_status(eligibility_cell, data['final_eligible'] if data else None)
        summary_start = eligibility_ws.max_row + 2
        summary_labels = ('Number of Eligible Students', 'Number of Ineligible Students', 'Number of Incomplete Students')
        for offset, label in enumerate(summary_labels):
            cell = eligibility_ws.cell(row=summary_start + offset, column=2, value=label)
            cell.font = Font(bold=True)
            cell.border = border
        for module in level_module_list:
            statuses = [
                enrollment_data[(reg_no, module.id)]['final_eligible']
                for reg_no in student_map
                if (reg_no, module.id) in enrollment_data
            ]
            counts = (
                sum(status is True for status in statuses),
                sum(status is False for status in statuses),
                sum(status is None for status in statuses),
            )
            mark_width = 4 if module.is_field_module else (9 if module.has_practical else 5)
            status_column = eligibility_starts[module.id] + mark_width + 1
            for offset, count in enumerate(counts):
                cell = eligibility_ws.cell(row=summary_start + offset, column=status_column, value=count)
                style_status(cell, True if offset == 0 else (False if offset == 1 else None))
        finish_sheet(eligibility_ws)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="final_eligibility_{timezone.localdate()}.xlsx"'
    wb.save(response)
    return response


# ── FEES LEDGER ───────────────────────────────────────────────────────────────
#
# Replaces the FinanceStudent/PaymentCategory/StudentPayment/Clearance viewsets
# above. The difference that matters: money is never written directly through a
# serializer here. Every mutation goes through attendance.finance, which is the
# only place that allocates payments and writes the audit trail.


class ChargeTypeViewSet(viewsets.ModelViewSet):
    """The catalogue: what the college can charge for, and what each blocks."""
    serializer_class = ChargeTypeSerializer
    permission_classes = [IsFinanceUser]

    def get_queryset(self):
        qs = ChargeType.objects.prefetch_related('charges')
        family = self.request.query_params.get('family')
        if family:
            qs = qs.filter(family=family)
        if self.request.query_params.get('is_active') == 'true':
            qs = qs.filter(is_active=True)
        return qs.order_by('family', 'name')

    def perform_create(self, serializer):
        charge_type = serializer.save()
        finance.audit('charge_type.create', 'ChargeType', actor=self.request.user,
                      entity_id=charge_type.id, summary=f'Created {charge_type.name}',
                      after=serializer.data, ip=finance.client_ip(self.request))

    def perform_update(self, serializer):
        before = ChargeTypeSerializer(serializer.instance).data
        charge_type = serializer.save()
        finance.audit('charge_type.update', 'ChargeType', actor=self.request.user,
                      entity_id=charge_type.id, summary=f'Updated {charge_type.name}',
                      before=before, after=serializer.data, ip=finance.client_ip(self.request))

    def perform_destroy(self, instance):
        if instance.charges.exists():
            raise PermissionDenied(
                'Students have already been charged under this type. Deactivate it instead '
                'so the existing charges keep their meaning.'
            )
        instance.delete()


class FeeStructureViewSet(viewsets.ModelViewSet):
    """The grid: how much, and over how many installments, per NTA level."""
    serializer_class = FeeStructureSerializer
    permission_classes = [IsFinanceUser]

    def get_queryset(self):
        qs = FeeStructure.objects.select_related(
            'charge_type', 'class_level', 'academic_year'
        ).prefetch_related('installment_schedule')
        year_id = self.request.query_params.get('academic_year_id')
        if year_id:
            qs = qs.filter(academic_year_id=year_id)
        level_id = self.request.query_params.get('class_level_id')
        if level_id:
            qs = qs.filter(class_level_id=level_id)
        return qs.order_by('class_level__order', 'charge_type__family', 'charge_type__name')

    def _save_with_schedule(self, serializer):
        due_dates = serializer.validated_data.pop('due_dates', None)
        structure = serializer.save()
        if due_dates:
            finance.set_installment_schedule(structure, due_dates)
        return structure

    def perform_create(self, serializer):
        structure = self._save_with_schedule(serializer)
        finance.audit('fee_structure.create', 'FeeStructure', actor=self.request.user,
                      entity_id=structure.id,
                      summary=f'{structure.charge_type} · {structure.class_level} = {structure.amount}',
                      ip=finance.client_ip(self.request))

    def perform_update(self, serializer):
        before = {
            'amount': str(serializer.instance.amount),
            'installments': serializer.instance.installments,
        }
        structure = self._save_with_schedule(serializer)
        finance.audit('fee_structure.update', 'FeeStructure', actor=self.request.user,
                      entity_id=structure.id,
                      summary=f'{structure.charge_type} · {structure.class_level} = {structure.amount}',
                      before=before,
                      after={'amount': str(structure.amount), 'installments': structure.installments},
                      ip=finance.client_ip(self.request))

    @action(detail=False, methods=['get'])
    def grid(self, request):
        """The whole fee structure as one sheet — charge types down, class
        levels across. Accountants read a grid far faster than a list, and this
        is the screen that replaces the Excel workbook."""
        year_id = request.query_params.get('academic_year_id')
        year = AcademicYear.objects.filter(id=year_id).first() or active_academic_year()
        if year is None:
            return Response({'detail': 'No academic year selected.'}, status=status.HTTP_400_BAD_REQUEST)

        levels = list(ClassLevel.objects.order_by('order', 'name'))
        types = list(ChargeType.objects.filter(is_active=True).order_by('family', 'name'))
        cells = {
            (s.charge_type_id, s.class_level_id): s
            for s in FeeStructure.objects.filter(academic_year=year)
            .prefetch_related('installment_schedule')
        }

        rows = []
        for charge_type in types:
            row = {
                'charge_type_id': charge_type.id,
                'charge_type_name': charge_type.name,
                'family': charge_type.family,
                'family_display': charge_type.get_family_display(),
                'applies': charge_type.applies,
                'cells': [],
            }
            for level in levels:
                structure = cells.get((charge_type.id, level.id))
                row['cells'].append({
                    'class_level_id': level.id,
                    'structure_id': structure.id if structure else None,
                    'amount': str(structure.amount) if structure else None,
                    'installments': structure.installments if structure else None,
                    'billing_period': structure.billing_period if structure else None,
                    'due_dates': [str(i.due_date) for i in structure.installment_schedule.all()] if structure else [],
                })
            rows.append(row)

        return Response({
            'academic_year': {'id': year.id, 'name': year.name},
            'class_levels': [{'id': lv.id, 'name': lv.name} for lv in levels],
            'rows': rows,
        })


class StudentChargeViewSet(viewsets.ReadOnlyModelViewSet):
    serializer_class = StudentChargeSerializer
    permission_classes = [IsFinanceUser]

    def get_queryset(self):
        qs = StudentCharge.objects.select_related(
            'profile', 'charge_type', 'academic_year', 'semester'
        ).prefetch_related('allocations')
        for param, field in [
            ('profile_id', 'profile_id'),
            ('academic_year_id', 'academic_year_id'),
            ('charge_type_id', 'charge_type_id'),
        ]:
            value = self.request.query_params.get(param)
            if value:
                qs = qs.filter(**{field: value})
        if self.request.query_params.get('outstanding') == 'true':
            qs = [c for c in qs if finance.charge_balance(c) > 0]
        return qs

    @action(detail=True, methods=['post'])
    def waive(self, request, pk=None):
        """Reduce a debt without pretending money arrived, so a bursary never
        shows up in the collections report as income."""
        if not can_manage_finance(request.user):
            raise PermissionDenied('Only the accountant can waive a charge.')
        charge = get_object_or_404(StudentCharge, pk=pk)
        serializer = WaiveChargeSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        if serializer.validated_data['amount'] > charge.amount:
            return Response({'detail': 'A waiver cannot exceed the charge.'},
                            status=status.HTTP_400_BAD_REQUEST)
        finance.waive_charge(charge, serializer.validated_data['amount'],
                             serializer.validated_data['reason'], actor=request.user)
        return Response(StudentChargeSerializer(charge).data)


class InvoiceViewSet(viewsets.ReadOnlyModelViewSet):
    serializer_class = InvoiceSerializer
    permission_classes = [IsFinanceUser]

    def get_queryset(self):
        qs = Invoice.objects.select_related('profile', 'academic_year').prefetch_related(
            'lines__charge__charge_type', 'payments'
        )
        profile_id = self.request.query_params.get('profile_id')
        if profile_id:
            qs = qs.filter(profile_id=profile_id)
        reference = self.request.query_params.get('reference')
        if reference:
            qs = qs.filter(reference__iexact=reference.strip())
        return qs

    @action(detail=False, methods=['get'], url_path='lookup')
    def lookup(self, request):
        """Find an invoice by the reference written on the slip. This is the
        accountant's entry point at the counter — they type what is on the
        paper and the system says whose it is and what it was for."""
        reference = str(request.query_params.get('reference', '')).strip().upper()
        if not reference.startswith('BPH-'):
            reference = f'BPH-{reference}'
        if not finance.reference_is_valid(reference):
            return Response(
                {'detail': 'That reference is not valid — check the digits on the slip.'},
                status=status.HTTP_400_BAD_REQUEST,
            )
        invoice = self.get_queryset().filter(reference=reference).first()
        if invoice is None:
            return Response({'detail': f'No invoice found for {reference}.'},
                            status=status.HTTP_404_NOT_FOUND)
        return Response(InvoiceSerializer(invoice).data)

    @action(detail=True, methods=['post'])
    def cancel(self, request, pk=None):
        invoice = get_object_or_404(Invoice, pk=pk)
        if invoice.payments.exists():
            return Response({'detail': 'That invoice has payments against it. Reverse them first.'},
                            status=status.HTTP_400_BAD_REQUEST)
        invoice.cancelled = True
        invoice.cancelled_reason = str(request.data.get('reason', ''))[:300]
        invoice.save(update_fields=['cancelled', 'cancelled_reason'])
        finance.audit('invoice.cancel', 'Invoice', actor=request.user, entity_id=invoice.id,
                      profile=invoice.profile, summary=f'Cancelled {invoice.reference}',
                      ip=finance.client_ip(request))
        return Response(InvoiceSerializer(invoice).data)


class PaymentViewSet(viewsets.ReadOnlyModelViewSet):
    """Append-only. There is no update and no delete — a correction is a
    reversal, so the record of what happened survives the correction."""
    serializer_class = PaymentSerializer
    permission_classes = [IsFinanceUser]

    def get_queryset(self):
        qs = Payment.objects.select_related(
            'profile', 'invoice', 'recorded_by'
        ).prefetch_related('allocations__charge__charge_type', 'reversal')
        for param, field in [
            ('profile_id', 'profile_id'),
            ('invoice_id', 'invoice_id'),
            ('channel', 'channel'),
        ]:
            value = self.request.query_params.get(param)
            if value:
                qs = qs.filter(**{field: value})
        date_from = self.request.query_params.get('from')
        if date_from:
            qs = qs.filter(payment_date__gte=date_from)
        date_to = self.request.query_params.get('to')
        if date_to:
            qs = qs.filter(payment_date__lte=date_to)
        return qs

    @action(detail=False, methods=['post'], url_path='record')
    def record(self, request):
        """Record money the accountant is holding proof of."""
        serializer = RecordPaymentSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = serializer.validated_data
        try:
            payment = finance.record_payment(
                data['profile'], data['amount'], data['payment_date'],
                recorded_by=request.user, invoice=data.get('invoice'),
                channel=data['channel'], bank_reference=data['bank_reference'],
                efd_receipt_no=data['efd_receipt_no'], payer_name=data['payer_name'],
                payer_relation=data['payer_relation'], proof=data.get('proof'),
                note=data['note'], request=request,
            )
        except ValueError as exc:
            return Response({'detail': str(exc)}, status=status.HTTP_400_BAD_REQUEST)
        return Response(PaymentSerializer(payment).data, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['post'])
    def reverse(self, request, pk=None):
        payment = get_object_or_404(Payment, pk=pk)
        serializer = ReversePaymentSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        try:
            reversal = finance.reverse_payment(
                payment, serializer.validated_data['reason'], actor=request.user, request=request,
            )
        except ValueError as exc:
            return Response({'detail': str(exc)}, status=status.HTTP_400_BAD_REQUEST)
        return Response(PaymentSerializer(reversal).data, status=status.HTTP_201_CREATED)


class FinanceOverrideViewSet(viewsets.ModelViewSet):
    """The human decision that beats the arithmetic — bursary, sponsor delay,
    hardship, or a hold placed for a reason outside the ledger."""
    serializer_class = FinanceOverrideSerializer
    permission_classes = [IsFinanceUser]

    def get_queryset(self):
        qs = FinanceOverride.objects.select_related('profile', 'academic_year', 'approved_by')
        for param, field in [
            ('profile_id', 'profile_id'),
            ('academic_year_id', 'academic_year_id'),
            ('period', 'period'),
        ]:
            value = self.request.query_params.get(param)
            if value:
                qs = qs.filter(**{field: value})
        return qs

    def perform_create(self, serializer):
        override = serializer.save(approved_by=self.request.user)
        finance.audit('override.create', 'FinanceOverride', actor=self.request.user,
                      entity_id=override.id, profile=override.profile,
                      summary=f'{override.get_period_display()} {override.status}: {override.reason}',
                      after=serializer.data, ip=finance.client_ip(self.request))

    def perform_update(self, serializer):
        before = FinanceOverrideSerializer(serializer.instance).data
        override = serializer.save(approved_by=self.request.user)
        finance.audit('override.update', 'FinanceOverride', actor=self.request.user,
                      entity_id=override.id, profile=override.profile,
                      summary=f'{override.get_period_display()} {override.status}: {override.reason}',
                      before=before, after=serializer.data, ip=finance.client_ip(self.request))

    def perform_destroy(self, instance):
        finance.audit('override.revoke', 'FinanceOverride', actor=self.request.user,
                      entity_id=instance.id, profile=instance.profile,
                      summary=f'Revoked {instance.get_period_display()} override',
                      before=FinanceOverrideSerializer(instance).data,
                      ip=finance.client_ip(self.request))
        instance.is_active = False
        instance.save(update_fields=['is_active'])


class FinanceAuditLogViewSet(viewsets.ReadOnlyModelViewSet):
    """Read-only, forever, and never writable through the API."""
    serializer_class = FinanceAuditLogSerializer
    permission_classes = [IsFinanceUser]

    def get_queryset(self):
        qs = FinanceAuditLog.objects.select_related('actor', 'profile')
        profile_id = self.request.query_params.get('profile_id')
        if profile_id:
            qs = qs.filter(profile_id=profile_id)
        action_name = self.request.query_params.get('action')
        if action_name:
            qs = qs.filter(action=action_name)
        return qs[:500]


# ── FEES LEDGER · summary and actions ────────────────────────────────────────

CLEARANCE_PERIODS = [
    ChargeType.REGISTRATION, ChargeType.CAT1, ChargeType.CAT2,
    ChargeType.FINAL, ChargeType.RESULTS,
]


def _finance_rows(profiles, year, semester):
    """One row per person — the row the old model could not produce.

    Balances and clearance are fetched for the whole list at once; doing it per
    student cost about fourteen queries each, which does not survive a class
    list.
    """
    profiles = list(profiles)
    totals = finance.balance_map(profiles, year)
    clearance = finance.clearance_map(profiles, year, CLEARANCE_PERIODS, semester=semester)

    rows = []
    for profile in profiles:
        # The enrollments are already prefetched, so this costs no query.
        enrollment = next(iter(profile.enrollments.all()), None)
        level = enrollment.module.class_level if enrollment else None
        mine = totals.get(profile.id, {})
        rows.append({
            'profile_id': profile.id,
            'nactvet_reg_no': profile.nactvet_reg_no,
            'name': profile.name,
            'class_level': level.name if level else None,
            'class_level_id': level.id if level else None,
            'billed': str(mine.get('billed', Decimal('0.00'))),
            'waived': str(mine.get('waived', Decimal('0.00'))),
            'paid': str(mine.get('paid', Decimal('0.00'))),
            'balance': str(mine.get('balance', Decimal('0.00'))),
            'clearance': {
                period: result['cleared']
                for period, result in clearance.get(profile.id, {}).items()
            },
        })
    return rows


@api_view(['GET'])
def finance_students(request):
    """Every student's position for the year: billed, paid, outstanding, and
    what they are cleared for. This is the debtors list."""
    if not can_manage_finance(request.user):
        raise PermissionDenied('Finance access required.')

    year = AcademicYear.objects.filter(
        id=request.query_params.get('academic_year_id')
    ).first() or active_academic_year()
    if year is None:
        return Response({'detail': 'No academic year is active.'}, status=status.HTTP_400_BAD_REQUEST)
    semester = active_semester()

    profiles = StudentProfile.objects.prefetch_related(
        'enrollments__module__class_level', 'charges__charge_type', 'charges__allocations',
    )
    search = str(request.query_params.get('search', '')).strip()
    if search:
        profiles = profiles.filter(
            Q(name__icontains=search) | Q(nactvet_reg_no__icontains=search)
        )

    rows = _finance_rows(profiles, year, semester)

    level_id = request.query_params.get('class_level_id')
    if level_id:
        rows = [r for r in rows if str(r['class_level_id']) == str(level_id)]
    if request.query_params.get('owing') == 'true':
        rows = [r for r in rows if Decimal(r['balance']) > 0]

    rows.sort(key=lambda r: (-Decimal(r['balance']), r['name']))
    return Response({
        'academic_year': {'id': year.id, 'name': year.name},
        'totals': {
            'students': len(rows),
            'billed': str(sum(Decimal(r['billed']) for r in rows)),
            'paid': str(sum(Decimal(r['paid']) for r in rows)),
            'outstanding': str(sum(Decimal(r['balance']) for r in rows)),
            'owing': sum(1 for r in rows if Decimal(r['balance']) > 0),
        },
        'rows': rows,
    })


@api_view(['GET'])
def finance_statement(request, profile_id):
    """One student's complete ledger — every charge, every payment, the running
    balance, and what each exam period is blocked on.

    A dispute is settled by looking at this rather than by arguing.
    """
    profile = get_object_or_404(StudentProfile, pk=profile_id)
    if not can_manage_finance(request.user):
        raise PermissionDenied('Finance access required.')
    year = AcademicYear.objects.filter(
        id=request.query_params.get('academic_year_id')
    ).first() or active_academic_year()
    return Response(_statement_payload(profile, year, active_semester()))


def _clearance_payload(result):
    """exam_clearance returns live StudentCharge objects so callers can inspect
    them; over the wire they need serialising."""
    return {
        'status': result['status'],
        'cleared': result['cleared'],
        'balance': str(result['balance']),
        'reason': result['reason'],
        'overridden': result['overridden'],
        'charges': StudentChargeSerializer(result['charges'], many=True).data,
    }


def _statement_payload(profile, year, semester):
    charges = (
        StudentCharge.objects.filter(profile=profile, academic_year=year)
        .select_related('charge_type', 'semester').prefetch_related('allocations')
        .order_by('due_date')
    )
    payments = (
        Payment.objects.filter(profile=profile)
        .select_related('invoice', 'recorded_by')
        .prefetch_related('allocations__charge__charge_type', 'reversal')
        .order_by('-payment_date', '-created_at')
    )
    invoices = (
        Invoice.objects.filter(profile=profile, academic_year=year)
        .prefetch_related('lines__charge__charge_type', 'payments')
    )
    totals = finance.balance_for(profile, year)
    return {
        'profile': {
            'id': profile.id,
            'nactvet_reg_no': profile.nactvet_reg_no,
            'name': profile.name,
        },
        'academic_year': {'id': year.id, 'name': year.name} if year else None,
        'totals': {k: str(v) for k, v in totals.items()},
        'charges': StudentChargeSerializer(charges, many=True).data,
        'payments': PaymentSerializer(payments, many=True).data,
        'invoices': InvoiceSerializer(invoices, many=True).data,
        'reminders': [
            {
                'name': reminder['name'],
                'group': reminder['group'],
                'installment_number': reminder['installment_number'],
                'installments_total': reminder['installments_total'],
                'amount': str(reminder['amount']),
                'due_date': reminder['due_date'],
                'days': reminder['days'],
                'days_label': reminder['days_label'],
                'urgency': reminder['urgency'],
                'reference': reminder['reference'],
            }
            for reminder in finance.installment_reminders(profile, year)
        ] if year else [],
        'clearance': {
            period: _clearance_payload(
                finance.exam_clearance(profile, year, period, semester=semester))
            for period in CLEARANCE_PERIODS
        } if year else {},
    }


@api_view(['POST'])
def finance_generate_charges(request):
    """Raise the year's automatic charges for a whole class level at once.

    Idempotent — running it again after adding students bills only the new
    ones, because a charge is unique per student, type, year and installment.
    """
    if not can_manage_finance(request.user):
        raise PermissionDenied('Only the accountant can generate charges.')

    year = get_object_or_404(AcademicYear, pk=request.data.get('academic_year'))
    level = ClassLevel.objects.filter(id=request.data.get('class_level')).first()

    profiles = StudentProfile.objects.prefetch_related('enrollments__module__class_level')
    if level:
        profiles = profiles.filter(enrollments__module__class_level=level).distinct()

    raised, billed, failures = 0, 0, []
    for profile in profiles:
        try:
            created = finance.generate_charges(profile, year, actor=request.user, class_level=level)
        except ValueError as exc:
            failures.append({'student': profile.nactvet_reg_no, 'detail': str(exc)})
            continue
        if created:
            billed += 1
            raised += len(created)

    return Response({
        'students_billed': billed,
        'charges_raised': raised,
        'failures': failures,
        'detail': f'Raised {raised} charge(s) across {billed} student(s).',
    }, status=status.HTTP_200_OK if not failures else status.HTTP_207_MULTI_STATUS)


@api_view(['POST'])
def finance_raise_charge(request):
    """Bill one student for something outside the automatic structure — a
    hostel place, a supplementary exam, a repeat module."""
    if not can_manage_finance(request.user):
        raise PermissionDenied('Finance access required.')
    profile = get_object_or_404(StudentProfile, pk=request.data.get('profile'))
    charge_type = get_object_or_404(ChargeType, pk=request.data.get('charge_type'))
    year = get_object_or_404(AcademicYear, pk=request.data.get('academic_year'))
    semester = Semester.objects.filter(id=request.data.get('semester')).first()

    try:
        amount = Decimal(str(request.data.get('amount')))
    except Exception:
        return Response({'detail': 'Give a valid amount.'}, status=status.HTTP_400_BAD_REQUEST)
    if amount <= 0:
        return Response({'detail': 'A charge must be greater than zero.'},
                        status=status.HTTP_400_BAD_REQUEST)
    # parse_date rather than handing the raw string to the model: Django would
    # coerce it on save, but the in-memory instance would still hold a string
    # and every date comparison on it (is_overdue, clearance) would blow up.
    due_date = parse_date(str(request.data.get('due_date') or ''))
    if due_date is None:
        return Response({'detail': 'Give a due date as YYYY-MM-DD.'},
                        status=status.HTTP_400_BAD_REQUEST)

    charge = finance.raise_charge(
        profile, charge_type, year, amount, due_date, semester=semester,
        actor=request.user, note=str(request.data.get('note', ''))[:300],
    )
    return Response(StudentChargeSerializer(charge).data, status=status.HTTP_201_CREATED)


@api_view(['POST'])
def finance_issue_invoice(request):
    """Turn outstanding charges into one payment instruction with one
    reference. Raised at the office; students use the portal endpoint."""
    if not can_manage_finance(request.user):
        raise PermissionDenied('Finance access required.')
    profile = get_object_or_404(StudentProfile, pk=request.data.get('profile'))
    year = AcademicYear.objects.filter(
        id=request.data.get('academic_year')
    ).first() or active_academic_year()

    charge_ids = request.data.get('charges') or []
    charges = list(StudentCharge.objects.filter(id__in=charge_ids, profile=profile))
    if not charges:
        charges = finance.outstanding_charges(profile, year)
    try:
        invoices = finance.issue_invoices(
            profile, charges, year, source=Invoice.OFFICE, actor=request.user,
        )
    except ValueError as exc:
        return Response({'detail': str(exc)}, status=status.HTTP_400_BAD_REQUEST)
    return Response(InvoiceSerializer(invoices, many=True).data, status=status.HTTP_201_CREATED)


@api_view(['GET'])
def finance_collections(request):
    """What came in over a period, for the cash book — and the EFD numbers to
    reconcile it against the machine."""
    if not can_manage_finance(request.user):
        raise PermissionDenied('Finance access required.')

    today = timezone.localdate()
    date_from = request.query_params.get('from') or str(today.replace(day=1))
    date_to = request.query_params.get('to') or str(today)

    payments = (
        Payment.objects.filter(payment_date__gte=date_from, payment_date__lte=date_to)
        .select_related('profile', 'recorded_by')
        .prefetch_related('allocations__charge__charge_type')
        .order_by('-payment_date', '-created_at')
    )

    by_channel, by_family = {}, {}
    for payment in payments:
        by_channel[payment.channel] = by_channel.get(payment.channel, Decimal('0.00')) + payment.amount
        for allocation in payment.allocations.all():
            family = allocation.charge.charge_type.get_family_display()
            by_family[family] = by_family.get(family, Decimal('0.00')) + allocation.amount

    total = sum((p.amount for p in payments), Decimal('0.00'))
    return Response({
        'from': date_from,
        'to': date_to,
        'total': str(total),
        'count': payments.count(),
        'missing_efd': payments.filter(efd_receipt_no='', reverses__isnull=True).count(),
        'by_channel': {k: str(v) for k, v in by_channel.items()},
        'by_family': {k: str(v) for k, v in by_family.items()},
        'payments': PaymentSerializer(payments, many=True).data,
    })


# ── FEES LEDGER · the student's own view ─────────────────────────────────────
#
# The student portal authenticates by session (`student_id`), never against
# request.user, so these are plain Django views rather than DRF ones — same
# reason announcement_download is.


def _student_profile(request):
    student = get_logged_student(request)
    if student is None:
        return None
    return finance.profile_for_student(student)


@api_view(['GET'])
def my_fees(request):
    """What I owe, what I have paid, and what I can sit for."""
    profile = _student_profile(request)
    if profile is None:
        return Response({'detail': 'Authentication required.'}, status=status.HTTP_403_FORBIDDEN)
    year = active_academic_year()
    if year is None:
        return Response({'detail': 'No academic year is active.'}, status=status.HTTP_400_BAD_REQUEST)
    return Response(_statement_payload(profile, year, active_semester()))


def _invoiceable_payload(payment):
    """One payment the student can be invoiced for, as the portal shows it."""
    invoice = payment['invoice']
    return {
        'family': payment['family'],
        'family_display': payment['family_display'],
        'group': payment['group'],
        'installments': payment['installments'],
        'billed': str(payment['billed']),
        'paid': str(payment['paid']),
        'outstanding': str(payment['outstanding']),
        'bank_account': (
            {
                'bank_name': payment['bank_account'].bank_name,
                'account_number': payment['bank_account'].account_number,
                'purpose': payment['bank_account'].purpose,
            } if payment['bank_account'] else None
        ),
        'items': [
            {
                'name': charge.charge_type.name,
                'installment_number': charge.installment_number,
                'installments_total': charge.installments_total,
                'due_date': charge.due_date,
                'amount': str(finance.money(charge.payable)),
                'balance': str(finance.charge_balance(charge)),
            }
            for charge in payment['charges']
        ],
        'invoice': InvoiceSerializer(invoice).data if invoice else None,
    }


@api_view(['GET'])
def my_invoice_options(request):
    """What I can raise an invoice for.

    The student picks a payment — school fees, direct costs, or one of the
    other payments such as accommodation or a supplementary exam — so the
    portal needs to know which ones they are billed for, what each is worth,
    and whether one already has a reference.
    """
    profile = _student_profile(request)
    if profile is None:
        return Response({'detail': 'Authentication required.'}, status=status.HTTP_403_FORBIDDEN)
    year = active_academic_year()
    if year is None:
        return Response({'detail': 'No academic year is active.'}, status=status.HTTP_400_BAD_REQUEST)

    payments = finance.invoiceable_payments(profile, year)
    return Response({
        'academic_year': {'id': year.id, 'name': year.name, 'closes_on': year.closes_on},
        'payments': [_invoiceable_payload(payment) for payment in payments],
    })


@api_view(['POST'])
def my_invoice(request):
    """Generate the invoice for a payment I intend to make.

    The student names the payment — `family` for school fees, direct costs or
    other payments, or `group` for one of them specifically. The invoice that
    comes back covers every instalment of it and expires when the academic year
    does, so the same reference is quoted on every slip that payment takes to
    the bank.

    That reference is what whoever actually pays quotes, so a parent's deposit
    still identifies the student it belongs to.
    """
    profile = _student_profile(request)
    if profile is None:
        return Response({'detail': 'Authentication required.'}, status=status.HTTP_403_FORBIDDEN)
    year = active_academic_year()
    if year is None:
        return Response({'detail': 'No academic year is active.'}, status=status.HTTP_400_BAD_REQUEST)

    family = (request.data.get('family') or '').strip()
    group = (request.data.get('group') or '').strip()
    charge_ids = request.data.get('charges') or []

    if family and family not in dict(ChargeType.FAMILY_CHOICES):
        return Response({'detail': f'{family} is not a kind of payment.'},
                        status=status.HTTP_400_BAD_REQUEST)

    charges = []
    if not family:
        # `charges` and `group` both name a payment; the invoice still covers
        # every instalment of it, because one payment carries one reference.
        if charge_ids:
            charges = list(StudentCharge.objects.filter(id__in=charge_ids, profile=profile)
                           .select_related('charge_type'))
            if not charges:
                return Response({'detail': 'Those charges are not yours.'},
                                status=status.HTTP_400_BAD_REQUEST)
        else:
            year_wide = finance.year_charges(profile, year)
            charges = [c for c in year_wide
                       if not group or c.charge_type.group_label == group]
            if group and not charges:
                return Response({'detail': f'You are not billed for {group}.'},
                                status=status.HTTP_400_BAD_REQUEST)
            charges = [c for c in charges if finance.charge_balance(c) > Decimal('0.00')]

    if not (family or charges):
        # "Nothing outstanding" is misleading when the real reason is that the
        # college has not billed anyone yet — the student reads it as "I am
        # paid up" and stops asking.
        billed = StudentCharge.objects.filter(profile=profile, academic_year=year).exists()
        detail = (
            'You are fully paid up for this year — there is nothing to invoice.'
            if billed else
            'Your fees for this year have not been set up yet. '
            'The Accounts Office will bill you before you need to pay.'
        )
        return Response({'detail': detail}, status=status.HTTP_400_BAD_REQUEST)

    try:
        # One invoice per payment: tuition and other charges are banked
        # separately, so a student settling both makes two deposits and needs
        # two references.
        if family:
            invoices = finance.issue_invoices_for_family(
                profile, year, family, source=Invoice.STUDENT)
        else:
            invoices = finance.issue_invoices(profile, charges, year, source=Invoice.STUDENT)
    except ValueError as exc:
        return Response({'detail': str(exc)}, status=status.HTTP_400_BAD_REQUEST)
    return Response(InvoiceSerializer(invoices, many=True).data, status=status.HTTP_201_CREATED)


def invoice_print(request, reference):
    """The invoice as a printable page, and as the student's download.

    Deliberately HTML rather than a generated PDF: the project has no PDF
    library, and the browser's print dialogue produces one anyway — which works
    from a phone, which is what most students have. `?download=1` opens that
    dialogue on load, so "Download" on the portal lands the student straight in
    Save as PDF instead of on a page they then have to work out how to save.
    """
    viewer_profile = None
    if request.session.get('student_id'):
        viewer_profile = _student_profile(request)
    elif not (request.user.is_authenticated and can_manage_finance(request.user)):
        return redirect('login')

    invoice = get_object_or_404(
        Invoice.objects.select_related('profile', 'academic_year', 'bank_account')
        .prefetch_related('lines__charge__charge_type', 'lines__charge__fee_structure'),
        reference__iexact=reference,
    )
    # A student may only ever print their own invoice.
    if viewer_profile is not None and invoice.profile_id != viewer_profile.id:
        raise Http404('No such invoice.')

    level = finance.class_level_for(invoice.profile, invoice.academic_year)
    college = CollegeProfile.get()
    total = finance.money(invoice.total)
    paid = finance.invoice_paid(invoice)

    # One item paid in instalments needs no item table — the amount and the
    # instalment count say it all. Several items must be listed, or the student
    # cannot tell what the total is made of.
    components = finance.invoice_components(invoice)
    installments = max((c['installments'] for c in components), default=0)

    return render(request, 'invoice_print.html', {
        'invoice': invoice,
        'components': components,
        'itemised': len(components) > 1,
        'installments': installments,
        'transactions': finance.invoice_transactions(invoice),
        'total': total,
        'paid': paid,
        'outstanding': max(total - paid, Decimal('0.00')),
        'invoice_status': finance.invoice_status(invoice),
        'expires_on': invoice.expires_on,
        'class_level': level.name if level else '',
        'college': college,
        'terms': [t.strip() for t in (college.invoice_terms or '').splitlines() if t.strip()],
        'auto_print': request.GET.get('download') in ('1', 'true', 'yes'),
    })


class BankAccountViewSet(viewsets.ModelViewSet):
    """The college's collection accounts.

    Tuition and other charges are banked separately, so which account a charge
    belongs to decides which invoice it lands on.
    """
    serializer_class = BankAccountSerializer
    permission_classes = [IsFinanceUser]

    def get_queryset(self):
        qs = BankAccount.objects.all()
        if self.request.query_params.get('is_active') == 'true':
            qs = qs.filter(is_active=True)
        return qs

    def perform_create(self, serializer):
        account = serializer.save()
        finance.audit('bank_account.create', 'BankAccount', actor=self.request.user,
                      entity_id=account.id, summary=str(account),
                      ip=finance.client_ip(self.request))

    def perform_update(self, serializer):
        before = BankAccountSerializer(serializer.instance).data
        account = serializer.save()
        finance.audit('bank_account.update', 'BankAccount', actor=self.request.user,
                      entity_id=account.id, summary=str(account),
                      before=before, after=serializer.data,
                      ip=finance.client_ip(self.request))

    def perform_destroy(self, instance):
        if instance.charge_types.exists() or instance.invoices.exists():
            raise PermissionDenied(
                'Charges or invoices already point at this account. Deactivate it instead, '
                'so existing invoices keep showing where the money went.'
            )
        instance.delete()


@api_view(['GET', 'PATCH'])
def college_profile(request):
    """The college's own details, as they print on an invoice."""
    if not can_manage_finance(request.user):
        raise PermissionDenied('Finance access required.')
    profile = CollegeProfile.objects.first() or CollegeProfile.objects.create()
    if request.method == 'PATCH':
        serializer = CollegeProfileSerializer(profile, data=request.data, partial=True)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        finance.audit('college.update', 'CollegeProfile', actor=request.user,
                      entity_id=profile.id, summary='Updated college details on invoices',
                      ip=finance.client_ip(request))
        return Response(serializer.data)
    return Response(CollegeProfileSerializer(profile).data)


# ── EVALUATION FORMS · the admin's side ───────────────────────────────────────

class FormViewSet(viewsets.ModelViewSet):
    """The forms the college publishes, and what came back.

    Read is open to any signed-in member of staff; only the admin
    (`is_staff` — the exam officer / HOD) may create, edit or publish one.
    """
    serializer_class = FormSerializer
    permission_classes = [IsAuthenticatedReadOnlyOrAdmin]

    def get_queryset(self):
        qs = (
            Form.objects.select_related('academic_year', 'created_by')
            .prefetch_related('sections__questions', 'levels')
        )
        year_id = self.request.query_params.get('academic_year_id')
        if year_id:
            qs = qs.filter(academic_year_id=year_id)
        if self.request.query_params.get('is_active') == 'true':
            qs = qs.filter(is_active=True)
        return qs

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

    def destroy(self, request, *args, **kwargs):
        """Delete a form the college no longer runs.

        Deleting one takes its responses with it, so a form that has been
        answered needs `?confirm=yes` — one deliberate extra step between a
        mis-click and a year of evaluation data. Export first; the refusal
        message says so.
        """
        form = self.get_object()
        answered = form.responses.count()
        confirmed = str(request.query_params.get('confirm', '')).lower() in ('yes', 'true', '1')

        if answered and not confirmed:
            return Response(
                {'detail': f'"{form.title}" has {answered} response'
                           f'{"" if answered == 1 else "s"}. Deleting it deletes them too. '
                           f'Download the Excel first if you need them, then confirm.',
                 'responses': answered, 'requires_confirmation': True},
                status=status.HTTP_409_CONFLICT,
            )

        finance.audit('form.delete', 'Form', actor=request.user, entity_id=form.id,
                      summary=f'Deleted "{form.title}" and {answered} response(s)',
                      before={'title': form.title, 'slug': form.slug,
                              'responses': answered},
                      ip=finance.client_ip(request))
        form.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

    @action(detail=True, methods=['post'])
    def submit(self, request, pk=None):
        """A member of staff filling in a staff form — a mentor assessing a
        student, say. Student forms are answered on the portal, not here."""
        form = self.get_object()
        if form.audience != Form.STAFF:
            return Response(
                {'detail': 'That form is answered by students on their portal.'},
                status=status.HTTP_400_BAD_REQUEST)
        if not request.user.is_staff:
            raise PermissionDenied('Only the admin can fill this form in.')

        answers = request.data.get('answers')
        if not isinstance(answers, dict):
            return Response({'detail': 'Send the answers as {question id: answer}.'},
                            status=status.HTTP_400_BAD_REQUEST)
        try:
            evaluations.submit(form, answers, submitted_by=request.user,
                               academic_year=form.academic_year or active_academic_year())
        except ValueError as exc:
            return Response({'detail': str(exc)}, status=status.HTTP_400_BAD_REQUEST)
        return Response({'detail': 'Recorded.'}, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['get'])
    def responses(self, request, pk=None):
        form = self.get_object()
        rows = (
            form.responses.select_related('profile', 'class_level')
            .prefetch_related('answers__question')
        )
        return Response(FormResponseSerializer(rows, many=True).data)

    @action(detail=True, methods=['get'])
    def summary(self, request, pk=None):
        """Counts and averages per question — what the charts draw."""
        return Response(evaluations.summarise(self.get_object()))

    @action(detail=True, methods=['get'])
    def export(self, request, pk=None):
        form = self.get_object()
        workbook = evaluations.export_workbook(form)
        buffer = BytesIO()
        workbook.save(buffer)
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        response['Content-Disposition'] = f'attachment; filename="{form.slug}_responses.xlsx"'
        return response

    @action(detail=True, methods=['get'], url_path='who-responded')
    def who_responded(self, request, pk=None):
        """Who has answered — never what they said.

        The admin still has to chase the students who have not filled a form
        in, and that must stay possible without breaking anonymity, so this
        reads the receipts rather than the responses.
        """
        form = self.get_object()
        answered = (
            FormSubmissionReceipt.objects.filter(form=form)
            .select_related('profile').order_by('profile__name')
        )
        return Response({
            'answered': [
                {'name': r.profile.name, 'reg_no': r.profile.nactvet_reg_no,
                 'submitted_at': r.submitted_at}
                for r in answered
            ],
            'count': answered.count(),
        })


class FormSectionViewSet(viewsets.ModelViewSet):
    serializer_class = FormSectionSerializer
    permission_classes = [IsAuthenticatedReadOnlyOrAdmin]

    def get_queryset(self):
        qs = FormSection.objects.select_related('form').prefetch_related('questions')
        form_id = self.request.query_params.get('form_id')
        return qs.filter(form_id=form_id) if form_id else qs

    def perform_destroy(self, instance):
        # A section takes its questions, and its questions take their answers.
        answered = FormAnswer.objects.filter(question__section=instance).exists()
        if answered:
            raise PermissionDenied(
                'Students have already answered questions in this section. Deleting it would '
                'delete their answers — edit the questions instead, or delete the whole form '
                'if the college has finished with it.'
            )
        instance.delete()


class FormQuestionViewSet(viewsets.ModelViewSet):
    serializer_class = FormQuestionSerializer
    permission_classes = [IsAuthenticatedReadOnlyOrAdmin]

    def get_queryset(self):
        qs = FormQuestion.objects.select_related('section__form')
        form_id = self.request.query_params.get('form_id')
        if form_id:
            qs = qs.filter(section__form_id=form_id)
        section_id = self.request.query_params.get('section_id')
        if section_id:
            qs = qs.filter(section_id=section_id)
        return qs

    def perform_destroy(self, instance):
        if instance.answers.exists():
            raise PermissionDenied(
                'Students have already answered this question. Deleting it would delete '
                'their answers with it — edit the wording instead, or deactivate the form.'
            )
        instance.delete()


# ── EVALUATION FORMS · the student's side ─────────────────────────────────────
#
# Session-authenticated like the fees endpoints, not request.user.

class ServiceRequestViewSet(mixins.RetrieveModelMixin,
                            mixins.ListModelMixin,
                            viewsets.GenericViewSet):
    """The queue of things students have asked the college for.

    Read is open to any signed-in member of staff; only the admin decides one.
    A request is the only thing on the portal the college owes an answer to, so
    it is a queue and not a report — filter it by status and work it down.
    """
    serializer_class = FormResponseSerializer
    permission_classes = [IsRequestOfficer]

    def get_queryset(self):
        return evaluations.request_queue(self.request.query_params.get('status'))

    @action(detail=True, methods=['post'])
    def decide(self, request, pk=None):
        if not can_answer_requests(request.user):
            raise PermissionDenied('Only the secretary or the admin can answer a service request.')
        service_request = self.get_object()
        try:
            evaluations.decide(
                service_request,
                status=request.data.get('status'),
                note=request.data.get('note', ''),
                by=request.user,
            )
        except ValueError as exc:
            return Response({'detail': str(exc)}, status=status.HTTP_400_BAD_REQUEST)

        finance.audit('service_request.decide', 'FormResponse', actor=request.user,
                      entity_id=service_request.id,
                      summary=f'{service_request.form.title} — '
                              f'{service_request.get_status_display().lower()}',
                      ip=finance.client_ip(request))
        return Response(FormResponseSerializer(service_request).data)


def request_print(request, pk):
    """An approved request as the college's own paper form, ready to print.

    HTML rather than a generated PDF, for the same reason the invoice is: the
    browser's print dialogue makes a PDF anyway, and it works from the phone
    most students actually have. `?download=1` opens that dialogue on load.

    Only an approved request prints. A pending one is not a document yet, and a
    declined one is not a document at all — printing either would put the
    college's letterhead on something nobody has agreed to.
    """
    viewer_profile = None
    if request.session.get('student_id'):
        viewer_profile = _student_profile(request)
    elif not can_answer_requests(request.user):
        return redirect('login')

    service_request = get_object_or_404(
        FormResponse.objects.select_related('form', 'profile', 'class_level',
                                            'academic_year', 'decided_by')
        .prefetch_related('answers__question', 'form__sections__questions'),
        pk=pk, form__kind=Form.REQUEST,
    )
    # A student may only ever print their own.
    if viewer_profile is not None and service_request.profile_id != viewer_profile.id:
        raise Http404('No such request.')
    if service_request.status != FormResponse.APPROVED:
        raise Http404('That request has not been approved.')

    answers = {answer.question_id: answer.value for answer in service_request.answers.all()}
    sections = []
    for section in service_request.form.sections.all():
        questions = list(section.questions.all())
        sections.append({
            'section': section,
            'for_office': section.for_office,
            # An office part is printed blank; a student part is printed with
            # what they wrote, and a question they left out is simply omitted.
            'rows': [{'question': question,
                      'display': evaluations.answer_text(question, answers.get(question.id))}
                     for question in questions
                     if section.for_office or answers.get(question.id) is not None],
        })

    return render(request, 'request_print.html', {
        'request_row': service_request,
        'form': service_request.form,
        'sections': sections,
        'college': CollegeProfile.get(),
        'notes': [line for line in service_request.form.print_note.splitlines() if line.strip()],
        'auto_print': request.GET.get('download') == '1',
    })


@api_view(['GET'])
def my_forms(request):
    """The forms I can fill in, and which I have already done."""
    profile = _student_profile(request)
    if profile is None:
        return Response({'detail': 'Authentication required.'}, status=status.HTTP_403_FORBIDDEN)
    return Response([
        {
            'id': entry['form'].id,
            'slug': entry['form'].slug,
            'title': entry['form'].title,
            'intro': entry['form'].intro,
            'is_anonymous': entry['form'].is_anonymous,
            'closes_on': entry['form'].closes_on,
            'answered': entry['answered'],
            'can_answer': entry['can_answer'],
        }
        for entry in evaluations.forms_for_student(profile)
    ])


@api_view(['GET'])
def my_form(request, slug):
    """One open form, with its questions, ready to fill in."""
    profile = _student_profile(request)
    if profile is None:
        return Response({'detail': 'Authentication required.'}, status=status.HTTP_403_FORBIDDEN)
    form = get_object_or_404(
        Form.objects.prefetch_related('sections__questions', 'levels'),
        slug=slug, audience=Form.STUDENT)
    # Same answer for closed and for aimed-at-another-level: a student has no
    # business learning which of the two it is.
    if not form.is_open() or not form.applies_to(evaluations.level_of(profile)):
        return Response({'detail': 'That form is not open for responses.'},
                        status=status.HTTP_404_NOT_FOUND)
    return Response({
        **StudentFormSerializer(form).data,
        'answered': form.id in evaluations.answered_form_ids(profile),
    })


@api_view(['POST'])
def submit_my_form(request, slug):
    """Send in my answers."""
    profile = _student_profile(request)
    if profile is None:
        return Response({'detail': 'Authentication required.'}, status=status.HTTP_403_FORBIDDEN)
    form = get_object_or_404(Form, slug=slug, audience=Form.STUDENT)

    answers = request.data.get('answers')
    if not isinstance(answers, dict):
        return Response({'detail': 'Send your answers as {question id: answer}.'},
                        status=status.HTTP_400_BAD_REQUEST)

    year = active_academic_year()
    try:
        evaluations.submit(
            form, answers, profile=profile,
            class_level=finance.class_level_for(profile, year),
            academic_year=year,
        )
    except ValueError as exc:
        return Response({'detail': str(exc)}, status=status.HTTP_400_BAD_REQUEST)
    detail = ('Your request has been sent. Follow it under Services → My Requests.'
              if form.kind == Form.REQUEST
              else 'Thank you — your response has been recorded.')
    return Response({'detail': detail}, status=status.HTTP_201_CREATED)
